"""CR-18 (Eliminate "Foam Family" Terminology Across Rigid Foam Edition,
Charlie's instruction, 2026-08-13) regression coverage.

Charlie's CR found that "Foam Family"/"Foam family"/"foam family" - a
Flexible Foam Edition term - had leaked into several customer-facing
surfaces of the Rigid Foam app, most visibly the shared "Analyze by:
Product Grade / Foam Family" control (helpers.analysis_unit_picker(),
consumed verbatim by views/16_Trend_Analysis.py, views/
17_Process_Property_Correlation.py, and views/
19_Machine_Settings_Optimization.py), each of those three pages' own
action text and pooling captions, the same "Foam scope" pattern on
views/5_Physical_Property_Result.py and views/6_Quality_Observation.py,
the PI3-subject-context `subject_desc` string duplicated in those three
pages AND three report-building functions in reports.py, and a report
link-type label dict in reports.py's Expert Notes aggregate report. The
fix standardizes every one of those customer-facing occurrences to
"Product Family"/"Product family"/"product family" - the term CR-18
required. R1-WP3 (2026-08-21) renamed that term again to "PU Material
Family"; the assertions below were updated to the new wording, and the
CR-18 requirement they encode (no Flexible-Foam "foam family" on any
customer-facing surface) is unchanged. While leaving every
internal identifier (`mode: "family"`, `link_type: "pu_material_family"`,
`FoamGrade`, `foam_grade_id`, comments/docstrings describing internal
behavior) untouched, per CR-18's own Internal Compatibility Boundary.

This file is the regression-scanning test CR-18 section 9 requires. It
covers, in order:

  1. A repository-wide case-insensitive scan for "foam family" - every
     hit is classified: it must live in one of a fixed, reviewed set of
     files/line numbers, and it must NOT be inside a Streamlit widget
     call, docx-writing call, or other customer-visible string
     construction (i.e. it really is the internal
     comment/docstring prose it claims to be).
  2. Live-rendered AppTest scans of the three Industrial Intelligence
     pages plus the two Quality pages for the shared control/warning/
     caption/action-text wording.
  3. A direct Process Parameter Optimization (page 19) check of its own
     Action text and the shared Analyze-by radio - the exact leak CR-18
     section 2 quotes.
  4. helpers.analysis_unit_picker()'s own family-mode behavior, driven
     live through Trend Analysis: the customer-facing control/warning/
     label text reads "PU Material Family" (R1-WP3; "Product family"
     before it) while the internal `mode`/
     `link_type` dict values and the pooled `grade_ids` list (business
     logic) are unchanged from before CR-18.
  5. Generated Word report text (reports.build_trend_analysis_report_data
     + render_trend_analysis_report_docx, and the Expert Notes aggregate
     report's link-type label) - the PI3-subject-context `subject_desc`
     string and the link-type label are both clean of "Foam Family".

Deliberately NOT re-scanned here: views/20_Expert_Notes.py's own
customer-facing text, which CR-15 already fixed and this CR's inventory
confirmed has no new leaks (its 3 remaining "foam family" mentions are
all comments/docstrings documenting CR-15's own completed fix).

Usage: python -m pytest tests/test_cr18_product_family_terminology.py -v
"""
import datetime as dt
import io
import os
import sys
import uuid

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
os.environ.setdefault("DATABASE_URL", "sqlite://")

import pytest
from docx import Document
from streamlit.testing.v1 import AppTest

import db
import helpers
import reports
import tenant_scope

APP_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PAGE16 = os.path.join(APP_DIR, "views", "16_Trend_Analysis.py")
PAGE17 = os.path.join(APP_DIR, "views", "17_Process_Property_Correlation.py")
PAGE19 = os.path.join(APP_DIR, "views", "19_Machine_Settings_Optimization.py")
PAGE5 = os.path.join(APP_DIR, "views", "5_Physical_Property_Result.py")
PAGE6 = os.path.join(APP_DIR, "views", "6_Quality_Observation.py")


def _reset_schema():
    db.Base.metadata.drop_all(db.ENGINE)
    db.Base.metadata.create_all(db.ENGINE)


def _extract_docx_text(docx_bytes):
    doc = Document(io.BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# 1. Repository-wide "foam family" scan (CR-18 section 9, item 1) - every
#    hit must be one of these exact, reviewed (file, line) pairs, and must
#    not look like a customer-visible string construction.
# ---------------------------------------------------------------------------

ALLOWED_FOAM_FAMILY_HITS = {
    # analytics.py/db.py line numbers shifted again by the WP7 Phase 1
    # Closeout Correction (2026-08-14, Charlie's review items 2.1-2.3):
    # db.py gained new imports (CheckConstraint, Index, event, func,
    # select, validates) plus the applicability same-scope unique index,
    # controlled-vocabulary @validates hooks, and the ProcessParameterValue
    # controlled-UOM-derivation event listener; analytics.py gained a
    # deterministic ORDER BY + comment in eligible_process_settings() -
    # same unchanged comment/docstring text elsewhere, new line positions.
    # db.py shifted again by the WP7 Phase 2 Closeout Correction (same
    # date, Charlie's Material Gap 2): PRODUCTION_RUN_STATUSES constant +
    # ProductionRun.run_start/run_end/status/order_item_reference columns
    # and their CheckConstraint/@validates pair. Shifted a third time by
    # the same Closeout Correction's Material Gap 3 fix:
    # ProductionMethod.uses_cycle_shot_operation and
    # Machine.cycle_shot_operation_override columns (config-driven
    # cycle/shot declaration, never inferred from a name). Shifted a
    # fourth time by the WP7 Phase 4 hybrid reader (2026-08-14, per
    # Stefan's direction to flag Charlie and go hybrid in the meantime):
    # analytics.py gained the ProcessParameterValue import,
    # dynamic_process_setting_field_key()/_dynamic_process_setting_fields()
    # helpers, and additive dynamic-field wiring in
    # eligible_phase_setting_fields()/run_settings_dataframe()/
    # rank_setting_correlations()/rank_setting_optimization() - same
    # unchanged comment/docstring text elsewhere, new line positions.
    # Shifted a fifth time (same date) when Charlie's Architecture
    # Clarification and Direction to JC rejected that hybrid outright:
    # eligible_phase_setting_fields()/run_settings_dataframe()/
    # rank_setting_correlations()/rank_setting_optimization() were reverted
    # to pure pre-hybrid legacy behavior, and the hybrid was replaced with
    # the new pure shared reader (production_run_process_parameters()/
    # production_run_parameter_dataframe(), plus the retained
    # dynamic_process_setting_field_key() helper) - net addition of ~180
    # lines of new docstring/code. Shifted a sixth time (same date) by the
    # WP7 Phase 4 Overview/output KPI cutover: analytics.py gained the
    # ProductionOutputSummary/UnitOfMeasure imports plus
    # production_run_output_summary()/production_output_totals(). Shifted a
    # seventh time (same date) by the WP7 Phase 4 Process-Property
    # Correlation (#978) / Process Parameter Optimization (#979) cutover:
    # merged_run_property_dataframe()/rank_setting_correlations()/
    # rank_setting_optimization() rewritten onto production_run_parameter_
    # dataframe() (new docstring paragraphs + tuple-return/data_type
    # branching logic), and pages 17/19 updated to consume the returned
    # field_labels/data_type maps instead of PHASE_SETTING_LABELS/
    # BOOLEAN_SETTING_FIELDS - same unchanged comment/docstring text
    # elsewhere, new line positions. The set below is the verified,
    # current, exact hit list, not an arithmetic shift guess.
    # Shifted an eighth time (same date), when the Correlation/Optimization
    # cutover's controllable/analytics_eligible filter (a real gap found by
    # the WP7 Phase 4 required-test-gate 7 audit, "Optimization
    # eligibility") added ~15 lines of docstring to merged_run_property_
    # dataframe() explaining the fix.
    # Shifted a ninth time (2026-08-14), by the WP7 Phase 4 targeted-
    # completion correction (Charlie's Closeout Review Return to JC,
    # Material Completion Item 1): production_run_process_parameters()
    # gained a 12-line docstring paragraph plus 4 additive dict keys
    # (min_value/max_value/min_value_override/max_value_override) - net
    # +18 lines, pushing every later hit down by 18.
    # Shifted a tenth time (2026-08-14), by WP7 Phase 4 targeted-completion
    # Item 3 (Trend Analysis method-aware parameter path):
    # process_parameter_definitions_for_trend()/process_parameter_run_
    # series() were added to analytics.py between production_run_parameter_
    # dataframe() and production_run_output_summary() - the former's own
    # docstring legitimately uses the same established "(a pooled foam
    # family) - see _grade_id_list" phrasing as every other foam_grade_id
    # docstring in this file (new hit, analytics.py:652), and every later
    # pre-existing hit shifted down by the ~130 new lines. views/
    # 16_Trend_Analysis.py also gained ~110 lines (the new "What to trend"
    # radio branch and its setup-section restructuring), shifting its own
    # two pre-existing hits down in turn.
    # Shifted an eleventh time (2026-08-14), by the WP7 Phase 4 targeted
    # closure gate's re-run dependency scan (Charlie's Closeout Review
    # Return to JC, targeted closure gate instruction 2): the scan's
    # direct-model-read pass (not just fixed-symbol) found analytics.
    # actual_usage_dataframe() still routed through a located
    # ProductionPhase to reach ComponentStreamReading - the same class of
    # gap Item 1.3 already fixed for Batch Release. Its docstring gained a
    # ~20-line correction paragraph explaining the fix (net +10 lines after
    # also trimming the now-stale "Finalized phase"/N+1 wording it
    # replaced), pushing only the one hit below it (previously
    # analytics.py:1996) down to analytics.py:2006. Every hit above this
    # function's docstring is unaffected.
    # analytics.py shifted again by WP7 Phase 5 (Legacy Retirement,
    # 2026-08-15): compute_runtime_output(), the PHASE_SETTING_FIELDS/
    # LABELS/BOOLEAN_SETTING_FIELDS/PHASE1_RIGID_INELIGIBLE_SETTINGS block,
    # and eligible_phase_setting_fields() were all removed (replaced with
    # shorter retirement comments), and run_settings_dataframe() was
    # simplified to drop its ProductionPhase query and field loop - net
    # fewer lines, so every "foam family" hit below moved up; one former
    # hit (in the removed PHASE_SETTING_FIELDS block's own docstring) no
    # longer exists at all, so the count dropped from 10 to 9.
    ("analytics.py", 15), ("analytics.py", 138), ("analytics.py", 233),
    ("analytics.py", 541), ("analytics.py", 792), ("analytics.py", 1043),
    ("analytics.py", 1181), ("analytics.py", 1702), ("analytics.py", 1877),
    # db.py shifted +132 lines by Phase 8 Decision 2 (2026-08-19): the
    # MachineStreamConfiguration / MachineStreamAssignment models, their
    # explanatory block comment, and ProductionRun.machine_stream_
    # configuration_id all landed above this line. Same unchanged comment
    # text, new position - 2046 -> 2178. Shifted again (+68) by Phase 8
    # Decision 3 the same day: RecipeComponent gained chemical_role,
    # chemical_role_source_id, chemical_role_source_location, their two
    # CheckConstraints and the block comment explaining why a role is never
    # inferred - 2178 -> 2246. Shifted once more by the Decision 3 review
    # correction (same day): the constraint comment explaining why IS NOT NULL
    # is required alongside trim() - 2246 -> 2256. Shifted a fourth time by
    # R-PRE-WP1 (2026-08-20): Machine.material_delivery_mode and the block
    # comment explaining why NULL means "applicable" - 2256 -> 2272. Comment
    # text unchanged throughout; only its position has ever moved. This
    # allowlist is deliberately position-based rather than text-based, so
    # every addition above line 2272 in db.py will move it again; that is the
    # cost of the check being exact, and the fix is always to re-point the
    # line, never to soften the scan. Fifth move, R1-WP3 (2026-08-21): the
    # PU Material Family model gained explanatory comments for the controlled
    # name and for the two columns that moved off the record - 2272 -> 2287.
    ("db.py", 2287),
    ("helpers.py", 86), ("helpers.py", 87), ("helpers.py", 157),
    ("views/16_Trend_Analysis.py", 228), ("views/16_Trend_Analysis.py", 411),
    ("views/17_Process_Property_Correlation.py", 109),
    ("views/19_Machine_Settings_Optimization.py", 105),
    ("views/20_Expert_Notes.py", 6), ("views/20_Expert_Notes.py", 38),
    ("views/20_Expert_Notes.py", 147),
    # Line numbers shifted again by the CR-22 / F22-01, F22-02 Product
    # scope rename + reorder edits, and again by the CR-22 focused closeout
    # correction (2026-08-16) which added an explanatory comment above
    # render_function_action_intro() on this page.
    ("views/5_Physical_Property_Result.py", 642),
    # Line number shifted by the CR-22 correction (2026-08-16, Charlie's
    # focused closeout return on F22-06): "Record against" was reordered
    # before the issue-type picker and a production_method_controlled_id
    # resolution block was added, both above this comment; shifted again
    # by an added explanatory comment above render_function_action_intro();
    # shifted again by the Phase 8 Wave A (2026-08-17) P8-D01 fix to the
    # "Breakdown by issue" chart, which replaced a single-line taxonomy
    # lookup with a multi-line quality_issue_registry.lookup() call above
    # this comment.
    ("views/6_Quality_Observation.py", 631),
}

# A hit whose line contains any of these is a live customer-facing string
# construction, not comment/docstring prose - CR-18 forbids "foam family"
# here regardless of which file it's in.
_CUSTOMER_FACING_MARKERS = [
    "st.radio(", "st.selectbox(", "st.warning(", "st.caption(", "st.error(",
    "st.info(", "st.success(", "st.markdown(", "st.write(", "st.title(",
    "st.subheader(", "cellPara(", "add_run(", "add_paragraph(",
]
FOAM_FAMILY_RE_SOURCE = "foam family|foam_family"


def _repo_python_files():
    import re
    pattern = re.compile(FOAM_FAMILY_RE_SOURCE, re.IGNORECASE)
    for root, dirs, files in os.walk(APP_DIR):
        dirs[:] = [d for d in dirs if d not in (".git", "__pycache__", "tests")]
        for fname in files:
            if not fname.endswith(".py"):
                continue
            if fname == "version.py":
                continue  # append-only changelog, never edited/scanned (established convention)
            path = os.path.join(root, fname)
            relpath = os.path.relpath(path, APP_DIR).replace(os.sep, "/")
            with open(path, encoding="utf-8") as f:
                for lineno, line in enumerate(f, start=1):
                    if pattern.search(line):
                        yield relpath, lineno, line


def test_repo_wide_foam_family_scan_only_hits_known_internal_lines():
    unexpected = []
    seen = set()
    for relpath, lineno, line in _repo_python_files():
        seen.add((relpath, lineno))
        if (relpath, lineno) not in ALLOWED_FOAM_FAMILY_HITS:
            unexpected.append((relpath, lineno, line.strip()))
    assert not unexpected, (
        f"Unreviewed 'foam family' hit(s) found outside the CR-18 internal-comment "
        f"allowlist - these need a customer-facing fix or an explicit allowlist "
        f"update: {unexpected}"
    )
    # Also assert the allowlist itself isn't stale (every entry still present) -
    # catches the allowlist silently going out of sync with the code.
    missing = ALLOWED_FOAM_FAMILY_HITS - seen
    assert not missing, f"Allowlisted 'foam family' line(s) no longer present in source: {missing}"


def test_allowed_foam_family_lines_are_not_customer_facing_strings():
    """Every allowlisted hit must be internal comment/docstring prose, not a
    live widget/docx string construction - guards against the allowlist ever
    being used to paper over a real leak."""
    offenders = []
    for relpath, lineno, line in _repo_python_files():
        if (relpath, lineno) in ALLOWED_FOAM_FAMILY_HITS:
            if any(marker in line for marker in _CUSTOMER_FACING_MARKERS):
                offenders.append((relpath, lineno, line.strip()))
    assert not offenders, f"Allowlisted line(s) actually construct a customer-facing string: {offenders}"


# ---------------------------------------------------------------------------
# 2 & 3. Live-rendered AppTest scans - Industrial Intelligence pages +
#    Quality pages. No fixture data needed: the leaked strings CR-18 found
#    were all static copy (action text, radio options, warning/caption
#    text), rendered on first load before any grade/family is selected.
# ---------------------------------------------------------------------------

def _all_widget_text(at):
    parts = []
    for kind in ("markdown", "caption", "title", "subheader", "warning", "info", "error"):
        parts.extend(w.value for w in getattr(at, kind))
    for r in at.radio:
        parts.extend(str(o) for o in r.options)
    for sb in at.selectbox:
        parts.extend(str(o) for o in sb.options)
    return "\n".join(parts)


@pytest.mark.parametrize("page_path,page_name", [
    (PAGE16, "Trend Analysis"),
    (PAGE17, "Process-Property Correlation"),
    (PAGE19, "Process Parameter Optimization"),
])
def test_industrial_intelligence_page_has_no_foam_family_leak(page_path, page_name):
    db.init_db()
    _reset_schema()
    at = AppTest.from_file(page_path, default_timeout=30)
    at.secrets["AUTH_DISABLED"] = True
    at.run()
    assert not at.exception, f"Unhandled exception loading {page_name}: {at.exception}"

    text = _all_widget_text(at)
    assert "foam family" not in text.lower(), f"'foam family' leaked on {page_name}: {text!r}"
    assert "foam family" not in " ".join(str(o) for r in at.radio for o in r.options).lower()


def test_process_parameter_optimization_action_text_and_radio_say_pu_material_family():
    """Direct check of the exact leak CR-18 section 2 quotes: page 19's
    Action text and the shared Analyze-by radio."""
    db.init_db()
    _reset_schema()
    at = AppTest.from_file(PAGE19, default_timeout=30)
    at.secrets["AUTH_DISABLED"] = True
    at.run()
    assert not at.exception, f"Unhandled exception loading Process Parameter Optimization: {at.exception}"

    action_text = "\n".join(w.value for w in at.markdown)
    assert "a PU Material Family to pool several grades together" in action_text
    assert "a foam family to pool several grades together" not in action_text


def test_quality_pages_product_scope_control_says_pu_material_family():
    """views/5 and views/6's "Product scope" radio (renamed from "Foam
    scope" and reordered per CR-22 / F22-01, F22-02, AF22-01) has options
    'All product grades / PU Material Family / Product grade' - Product family
    before Product grade, matching the hierarchy order used everywhere else
    (e.g. helpers.analysis_unit_picker()). Both pages gate their filterable
    table (and this radio) behind "at least one production run/trial
    exists" - reaching it live would require a full run/trial fixture for
    no extra assurance, so this is a direct source check of the exact
    widget-construction lines instead, the same convention tests/
    test_cr17_nav_restore.py uses for nav-list source assertions."""
    for page_path, page_name in ((PAGE5, "Quality Test Result"), (PAGE6, "Quality Issue")):
        with open(page_path, encoding="utf-8") as f:
            source = f.read()
        assert '"Product scope", ["All product grades", "PU Material Family", "Product grade"]' in source, (
            f"{page_name}'s Product scope radio options are not "
            f"'All product grades / PU Material Family / Product grade'"
        )
        assert '"Foam scope"' not in source, (
            f"{page_name} still has a 'Foam scope' widget label - should be 'Product scope' (CR-22 / F22-01)"
        )
        assert 'st.caption("No PU Material Family available for these grades yet.")' in source, (
            f"{page_name}'s empty-state caption is not 'No PU Material Family available...'"
        )
        assert '"PU Material Family", families, format_func=pu_material_family_label' in source, (
            f"{page_name}'s family selectbox label is not 'PU Material Family'"
        )
        assert '"Foam family"' not in source, f"{page_name} still constructs a 'Foam family' string"
        assert 'No foam family available' not in source, f"{page_name} still has the old empty-state caption"


# ---------------------------------------------------------------------------
# 4. helpers.analysis_unit_picker() family-mode regression, driven live
#    through Trend Analysis: customer-facing text reads "PU Material Family"
#    while internal identifiers and the pooling business logic (which
#    grade_ids get pooled) are unchanged.
# ---------------------------------------------------------------------------

def _seed_family_with_two_grades():
    u = uuid.uuid4().hex[:8]
    session = db.get_session()
    company = db.Company(name=f"CR18 Co {u}", is_platform_owner=True)
    session.add(company); session.flush()
    plant = db.Plant(company_id=company.id, name=f"CR18 Plant {u}")
    session.add(plant); session.flush()
    family = db.PUMaterialFamily(plant_id=plant.id, name=f"CR18 Family {u}")
    session.add(family); session.flush()
    machine = db.Machine(plant_id=plant.id, name=f"CR18 Machine {u}")
    session.add(machine); session.flush()

    grade_ids = []
    grade_names = []
    for i in range(2):
        grade = db.FoamGrade(pu_material_family_id=family.id, grade_name=f"CR18-G{i}-{u}")
        session.add(grade); session.flush()
        recipe = db.RecipeVersion(foam_grade_id=grade.id, version_label="v1", approval_status="Draft", is_active=True)
        session.add(recipe); session.flush()
        run = db.ProductionRun(
            plant_id=plant.id, foam_grade_id=grade.id, machine_id=machine.id, recipe_version_id=recipe.id,
            run_date=dt.date(2026, 8, 1 + i), batch_reference=f"CR18-B{i}-{u}",
        )
        session.add(run); session.flush()
        result = db.PhysicalPropertyResult(
            production_run_id=run.id, property_name="Density", actual_value=32.0 + i,
            unit="kg/m3", tested_at=run.run_date,
        )
        session.add(result); session.flush()
        grade_ids.append(grade.id)
        grade_names.append(grade.grade_name)
    session.commit()
    ids = {"company_id": company.id, "grade_ids": sorted(grade_ids), "grade_names": sorted(grade_names)}
    session.close()
    return ids


@pytest.fixture()
def family_fixture():
    db.init_db()
    _reset_schema()
    return _seed_family_with_two_grades()


def test_analysis_unit_picker_family_mode_uses_pu_material_family_wording_and_pools_correctly(family_fixture):
    tenant_scope.clear_scope_cache()
    at = AppTest.from_file(PAGE16, default_timeout=30)
    at.secrets["AUTH_DISABLED"] = True
    at.session_state["is_platform_owner"] = True
    at.session_state["company_id"] = family_fixture["company_id"]
    at.run()
    assert not at.exception, f"Unhandled exception loading Trend Analysis: {at.exception}"

    # Option order is Product family before Product grade (CR-22 / F22-02,
    # AF22-01) - helpers.analysis_unit_picker() was reordered so the
    # hierarchy order (All product grades -> Product family -> Product
    # grade) is consistent everywhere it appears.
    unit_mode = next(r for r in at.radio if r.key == "trend_unit_mode")
    assert unit_mode.options == ["PU Material Family", "Product grade"], (
        f"Analyze-by control's own option text is not 'PU Material Family'/'Product grade': {unit_mode.options}"
    )
    unit_mode.set_value("PU Material Family")
    at.run()
    assert not at.exception, f"Unhandled exception after switching to Product family mode: {at.exception}"

    family_select = next((sb for sb in at.selectbox if sb.key == "trend_family_select"), None)
    assert family_select is not None, "Product family selectbox did not appear after switching modes"
    assert family_select.label == "PU Material Family"

    all_text = "\n".join(w.value for w in list(at.caption) + list(at.markdown))
    assert "foam family" not in all_text.lower(), f"'foam family' leaked once in family mode: {all_text!r}"
    assert f"Pooling {len(family_fixture['grade_ids'])} grade(s) in PU Material Family" in all_text, (
        "Pooling caption missing or not using 'PU Material Family' wording"
    )
    # Business logic unchanged: both seeded grades are still offered/pooled
    # by name in the caption, proving CR-18 only touched the wording, not
    # which grades a family mode actually pools.
    for name in family_fixture["grade_names"]:
        assert name in all_text, f"Pooled grade {name!r} missing from the pooling caption text"


def test_analysis_unit_picker_internal_dict_identifiers_unchanged():
    """The customer-facing wording changed; the internal contract every
    caller (analytics.py, reports.py, helpers.py's own Expert Notes link
    plumbing) relies on must not have - mode stays "family", link_type
    stays "pu_material_family", per CR-18's Internal Compatibility Boundary."""
    import inspect
    src = inspect.getsource(helpers.analysis_unit_picker)
    assert '"mode": "family"' in src
    assert '"link_type": "pu_material_family"' in src
    assert '"mode": "grade"' in src
    assert '"link_type": "foam_grade"' in src


# ---------------------------------------------------------------------------
# 5. Generated Word report text - subject_desc (reports.py, 3 occurrences)
#    and the Expert Notes aggregate report's link-type label.
# ---------------------------------------------------------------------------

def test_trend_analysis_report_docx_subject_desc_says_pu_material_family():
    fake_unit = {
        "mode": "family", "label": "CR18 Report Family",
        "member_grade_names": ["Grade A", "Grade B"], "grade_ids": [1, 2],
    }
    import pandas as pd
    data = reports.build_trend_analysis_report_data(
        session=None, unit=fake_unit, property_name="Density",
        series=pd.DataFrame({"tested_at": []}), pooling_grades=True,
        chart_result=None, capability=None, cusum=None, trend=None,
        change_rows=[], include_trials=False,
    )
    assert data["subject_desc"] == "PU Material Family CR18 Report Family (pooling grades: Grade A, Grade B)"
    docx_bytes = reports.render_trend_analysis_report_docx(data)
    text = _extract_docx_text(docx_bytes)
    assert "PU Material Family" in text
    assert "foam family" not in text.lower()


def test_expert_notes_aggregate_report_link_type_label_says_pu_material_family():
    from types import SimpleNamespace
    notes = [
        SimpleNamespace(
            confidence_level="High", source="Manual", linked_entity_type="pu_material_family",
            vector_store_file_id=None,
        ),
    ]
    data = reports.build_expert_notes_report_data(session=None, notes=notes, scope_label="All companies")
    link_labels = [row["Linked to"] for row in data["link_type_rows"]]
    assert "PU Material Family" in link_labels
    assert "Foam Family" not in link_labels
    docx_bytes = reports.render_expert_notes_report_docx(data)
    text = _extract_docx_text(docx_bytes)
    assert "PU Material Family" in text
    assert "Foam Family" not in text


if __name__ == "__main__":
    sys.exit(pytest.main([__file__, "-v"]))
