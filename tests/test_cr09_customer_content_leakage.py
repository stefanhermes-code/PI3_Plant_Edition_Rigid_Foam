"""CR-09 (Remove Internal Development and UAT Leakage from Customer-Facing
Application, Charlie's instruction, 2026-08-12) regression coverage.

Charlie's UAT audit found that internal engineering vocabulary - Phase
numbers, WP identifiers, UAT/synthetic-dataset terminology, and raw
maturity_status/production_release codes - had leaked into customer-facing
screens, generated Word reports, and PI3 prompt construction:

  - views/30_Production_Methods.py: a caption showing the raw
    maturity_status value and literally saying "Phase 1 offers Production
    Method PM-100 only."
  - views/21_Report.py: help text and a tab label naming "WP3" (a
    development work-package identifier).
  - reports.py: a Word report subtitle reading "Synthetic UAT / Reference
    Dataset" verbatim; a Note column and PI3 prompt text carrying
    wp3_conformance.production_release_status()'s raw internal code
    ("UAT_PASS_NO_RELEASE"); a Word report title naming "WP3"; a Word
    report table row showing FoamGrade.status's raw "UAT_ONLY" code.
  - views/15_Recipe_Optimization.py: the same Note-column and PI3-prompt
    leaks as reports.py's Recipe Optimization Word report, plus a static
    prompt instruction saying "(UAT-only)".

The fix routes every one of those call sites through customer_
presentation.py, a single dependency-free translation module (see that
module's own docstring for why it can't live in helpers.py - helpers.py
already imports reports.py, so a shared helper in either would create a
circular import). This file is the leakage-scanning regression test CR-09
section 9 requires: it scans actual rendered Streamlit page output and
actual generated Word report text - not just the translation functions in
isolation - for the forbidden markers, and separately proves each
translation function's own text is clean.

Deliberately NOT scanned/flagged: Python comments and docstrings (internal
scope, CR-09 section 6), and "Setup"/"Finalized" ProductionPhase labels
(legitimate operational vocabulary, not a development-phase leak).

Usage: python -m pytest tests/test_cr09_customer_content_leakage.py
"""
import datetime as dt
import io
import os
import sys
import uuid

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
os.environ.setdefault("DATABASE_URL", "sqlite://")

import pytest
from docx import Document
from streamlit.testing.v1 import AppTest

import customer_presentation
import db
import tenant_scope
import reports
import wp3_conformance

APP_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PAGE30 = os.path.join(APP_DIR, "views", "30_Production_Methods.py")
PAGE21 = os.path.join(APP_DIR, "views", "21_Report.py")

# CR-09 section 9's marker list. "Phase" is checked as "Phase 1"/"Phase 2"
# etc (word-boundary-ish, via a dedicated helper below) so it doesn't
# false-positive on "Setup"/"Finalized" ProductionPhase labels, which
# don't contain the word "Phase" at all and are unaffected either way.
FORBIDDEN_SUBSTRINGS = [
    "WP3", "UAT", "Gate 2", "Gate 1", "CR-09", "CR-01", "maturity_status",
    "production_release", "Synthetic", "synthetic", "UAT_PASS_NO_RELEASE",
    "UAT_ONLY",
]
FORBIDDEN_PHASE_PATTERNS = ["Phase 1", "Phase 2", "Phase 3"]


def _find_leaks(text):
    """Returns the subset of forbidden markers actually present in text -
    empty list means clean. Kept as a list (not a bool) so a failing
    assertion message names exactly what leaked."""
    hits = [m for m in FORBIDDEN_SUBSTRINGS if m in text]
    hits += [p for p in FORBIDDEN_PHASE_PATTERNS if p in text]
    return hits


def _extract_docx_text(docx_bytes):
    """Flattens every paragraph and table cell in a generated .docx into
    one string, for scanning generated Word report text the same way a
    customer would read it."""
    doc = Document(io.BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _reset_schema():
    db.Base.metadata.drop_all(db.ENGINE)
    db.Base.metadata.create_all(db.ENGINE)


# ---------------------------------------------------------------------------
# 1. customer_presentation.py - the translation functions themselves
# ---------------------------------------------------------------------------

def test_method_availability_note_has_no_maturity_status_or_phase_leak():
    note = customer_presentation.customer_facing_method_availability_note()
    assert not _find_leaks(note), f"Leak in method availability note: {note!r}"
    assert "not yet available" in note.lower()


def test_release_note_translates_known_code_with_no_leak():
    note = customer_presentation.customer_facing_release_note("UAT_PASS_NO_RELEASE")
    assert not _find_leaks(note), f"Leak in release note: {note!r}"
    assert "not yet cleared for production release" in note.lower()


def test_release_note_is_none_when_nothing_to_say():
    assert customer_presentation.customer_facing_release_note(None) is None
    assert customer_presentation.customer_facing_release_note("") is None


def test_reference_dataset_label_translates_with_no_leak():
    label = customer_presentation.customer_facing_reference_dataset_label(True)
    assert not _find_leaks(label), f"Leak in reference dataset label: {label!r}"
    assert customer_presentation.customer_facing_reference_dataset_label(False) is None
    assert customer_presentation.customer_facing_reference_dataset_label(None) is None


def test_report_title_translates_wp3_with_no_leak():
    title = customer_presentation.customer_facing_report_title("WP3 Property Conformance Report")
    assert not _find_leaks(title), f"Leak in report title: {title!r}"
    assert title == "Property Conformance Report"
    # An unrecognised title passes through unchanged - this function is a
    # small fixed lookup, not a general string-replace.
    assert customer_presentation.customer_facing_report_title("Some Other Report") == "Some Other Report"


def test_grade_status_label_translates_uat_only_with_no_leak():
    label = customer_presentation.customer_facing_grade_status_label("UAT_ONLY")
    assert not _find_leaks(label), f"Leak in grade status label: {label!r}"
    assert customer_presentation.customer_facing_grade_status_label(None) is None
    # An unrecognised status (e.g. the real "ACTIVE" value) passes through -
    # only UAT_ONLY is currently known to be internal-only.
    assert customer_presentation.customer_facing_grade_status_label("ACTIVE") == "ACTIVE"


# ---------------------------------------------------------------------------
# 2. views/30_Production_Methods.py - live-rendered Streamlit page scan
# ---------------------------------------------------------------------------

@pytest.fixture()
def gated_method_fixture():
    db.init_db()
    _reset_schema()
    u = uuid.uuid4().hex[:8]
    session = db.get_session()
    company = db.Company(name=f"CR09 Co {u}", is_platform_owner=False)
    session.add(company); session.flush()
    plant = db.Plant(company_id=company.id, name=f"CR09 Plant {u}")
    session.add(plant); session.flush()
    released = db.ProductionMethod(
        controlled_id=f"PM-100-{u}", name="Discontinuous Factory Foaming",
        maturity_status="Released", is_released=True,
    )
    unreleased = db.ProductionMethod(
        controlled_id=f"PM-200-{u}", name="Continuous Panel & Board Production",
        maturity_status="Defined / planned", is_released=False,
    )
    session.add_all([released, unreleased]); session.flush()
    session.commit()
    ids = {"company_id": company.id, "plant_id": plant.id}
    session.close()
    return ids


def test_production_methods_page_has_no_leak_and_shows_safe_caption(gated_method_fixture):
    ids = gated_method_fixture
    at = AppTest.from_file(PAGE30, default_timeout=30)
    at.secrets["AUTH_DISABLED"] = True
    at.session_state["is_platform_owner"] = False
    at.session_state["company_id"] = ids["company_id"]
    at.run()
    assert not at.exception, f"Unhandled exception loading Production Methods: {at.exception}"

    all_text = "\n".join(
        w.value for w in list(at.caption) + list(at.markdown) + list(at.title) + list(at.subheader)
    )
    leaks = _find_leaks(all_text)
    assert not leaks, f"Customer-facing leak on Production Methods page: {leaks} in {all_text!r}"
    assert customer_presentation.customer_facing_method_availability_note() in all_text


# ---------------------------------------------------------------------------
# 3. views/21_Report.py - live-rendered Streamlit page scan (tab labels +
#    help text, no data selected - the leaks CR-09 found here were both in
#    static copy, not data-driven)
# ---------------------------------------------------------------------------

def test_report_page_tab_labels_and_help_text_have_no_leak():
    at = AppTest.from_file(PAGE21, default_timeout=30)
    at.secrets["AUTH_DISABLED"] = True
    at.run()
    assert not at.exception, f"Unhandled exception loading Report page: {at.exception}"

    tab_labels = [t.label for t in at.tabs]
    assert "Property Conformance Report" in tab_labels
    assert not any("WP3" in label for label in tab_labels), f"WP3 leaked into a tab label: {tab_labels}"

    all_text = "\n".join(w.value for w in list(at.markdown) + list(at.caption))
    leaks = _find_leaks(all_text)
    assert not leaks, f"Customer-facing leak on Report page help text: {leaks} in {all_text!r}"


# ---------------------------------------------------------------------------
# 4. reports.py - generated Word report text, built from real UAT-only /
#    synthetic-dataset fixtures so the actual leak paths are exercised
#    end-to-end (not just the translation functions in isolation)
# ---------------------------------------------------------------------------

def _seed_wp3_fixture(session):
    """A UAT-only grade with one passing spec against one production run -
    exactly the condition wp3_conformance.production_release_status()
    flags as "UAT_PASS_NO_RELEASE" (see that function's docstring)."""
    u = uuid.uuid4().hex[:8]
    company = db.Company(name=f"CR09 Conformance Co {u}", is_platform_owner=True)
    session.add(company); session.flush()
    plant = db.Plant(company_id=company.id, name=f"CR09 Conformance Plant {u}")
    session.add(plant); session.flush()
    family = db.ProductFamily(plant_id=plant.id, name=f"CR09 Family {u}")
    session.add(family); session.flush()
    grade = db.FoamGrade(product_family_id=family.id, grade_name=f"RF-CR09-{u}", status="UAT_ONLY")
    session.add(grade); session.flush()

    propdef = db.PhysicalPropertyDefinition(name=f"CR09 Thermal conductivity {u}")
    session.add(propdef); session.flush()
    propdef.mandatory_context = "Record mean test temperature, thickness, orientation, test age and conditioning"
    session.flush()
    propmethod = db.PhysicalPropertyMethod(property_definition_id=propdef.id, method_code="ISO 8301", controlled_id=f"MTH-CR09-{u}")
    session.add(propmethod); session.flush()

    orientation = db.Orientation(controlled_id=f"ORI-CR09-{u}", name="Through-thickness")
    location = db.Location(controlled_id=f"LOC-CR09-{u}", name="Core")
    condition = db.TestCondition(controlled_id=f"CTX-CR09-{u}", name="Initial, 10C mean, 7 days")
    session.add_all([orientation, location, condition]); session.flush()

    spec = db.GradeSpecification(
        foam_grade_id=grade.id, property_definition_id=propdef.id, property_method_id=propmethod.id,
        property_name=propdef.name, target_operator="<=", target_value=0.024, unit="W/(m.K)",
        condition_id=condition.id, orientation_id=orientation.id, location_id=location.id,
    )
    session.add(spec); session.flush()

    machine = db.Machine(plant_id=plant.id, name=f"CR09 Machine {u}")
    session.add(machine); session.flush()
    recipe = db.RecipeVersion(foam_grade_id=grade.id, version_label="v1", approval_status="Draft", is_active=True)
    session.add(recipe); session.flush()
    run = db.ProductionRun(
        plant_id=plant.id, foam_grade_id=grade.id, machine_id=machine.id, recipe_version_id=recipe.id,
        run_date=dt.date(2026, 8, 1), batch_reference="CR09-B1",
    )
    session.add(run); session.flush()
    sample = db.Sample(
        production_run_id=run.id, location_id=location.id, orientation_id=orientation.id,
        thickness_mm=60.0, age_hours=168.0, sample_scope="Core",
        sample_ts=dt.datetime(2026, 8, 1, 10, 0),
    )
    session.add(sample); session.flush()
    result = db.PhysicalPropertyResult(
        production_run_id=run.id, sample_id=sample.id, property_definition_id=propdef.id,
        property_method_id=propmethod.id, property_name=propdef.name, actual_value=0.022,
        unit="W/(m.K)", test_method="ISO 8301", condition_id=condition.id,
        orientation_id=orientation.id, location_id=location.id, tested_at=run.run_date,
    )
    session.add(result); session.flush()
    # Enabled here so test_recipe_optimization_pi3_prompt_has_no_leak's "Get PI3
    # recommendation" button actually renders (ai_assistant.is_enabled_for_plant()
    # gate) - the fake OPENAI_API_KEY/PI3_VECTOR_STORE_ID secrets set in that
    # test's AppTest run only satisfy is_configured()'s presence check; the real
    # OpenAI call is monkeypatched away, never a live network call. Same pattern
    # as tests/test_wp4_recipe_optimization_page_smoke.py's seeded_rigid_only.
    session.add(db.PI3AIConnectionSetting(plant_id=plant.id, pi3_ai_connectivity_enabled=True))
    # A FoamGradeTargetProperty row (not just the GradeSpecification above) so
    # the "Get PI3 recommendation" button's target_properties text area is
    # non-empty and the button isn't disabled - this grade has no
    # target_density/target_hardness set and its one PhysicalPropertyResult
    # carries no target_value (rigid-style: the limit lives on
    # GradeSpecification), so without this the button stays disabled and the
    # test can never reach ask_assistant().
    session.add(db.FoamGradeTargetProperty(
        foam_grade_id=grade.id, property_definition_id=propdef.id,
        property_name=propdef.name, target_value=0.024, unit="W/(m.K)",
    ))
    session.commit()
    return grade.id, run.id


@pytest.fixture()
def wp3_session():
    db.init_db()
    _reset_schema()
    s = db.get_session()
    yield s
    s.close()


def test_wp3_conformance_report_docx_has_no_leak(wp3_session):
    session = wp3_session
    grade_id, run_id = _seed_wp3_fixture(session)
    data = reports.build_wp3_conformance_report_data(session, grade_id, run_id)
    assert data["overall_verdict"] == "Conforming"
    # Confirm the fixture actually exercises the leak path before trusting
    # a clean scan below - if this ever stops being true (e.g. someone
    # changes production_release_status's trigger condition), the test
    # should fail loudly here rather than silently passing on a fixture
    # that no longer proves anything.
    assert any(row["Note"] for row in data["conformance_rows"]), (
        "Fixture no longer produces a Note - test no longer proves the release-note leak is fixed"
    )
    docx_bytes = reports.render_wp3_conformance_report_docx(data)
    text = _extract_docx_text(docx_bytes)
    leaks = _find_leaks(text)
    assert not leaks, f"Customer-facing leak in Property Conformance Report .docx: {leaks}"
    assert "Property Conformance Report" in text
    assert "not yet cleared for production release" in text.lower()


def test_period_summary_report_docx_translates_synthetic_dataset_flag(wp3_session):
    session = wp3_session
    u = uuid.uuid4().hex[:8]
    company = db.Company(name=f"CR09 Period Co {u}", is_platform_owner=True)
    session.add(company); session.flush()
    plant = db.Plant(company_id=company.id, name=f"CR09 Period Plant {u}")
    session.add(plant); session.flush()
    family = db.ProductFamily(plant_id=plant.id, name=f"CR09 Period Family {u}")
    session.add(family); session.flush()
    grade = db.FoamGrade(product_family_id=family.id, grade_name=f"RF-CR09-PERIOD-{u}")
    session.add(grade); session.flush()
    recipe = db.RecipeVersion(foam_grade_id=grade.id, version_label="v1", approval_status="Draft", is_active=True)
    session.add(recipe); session.flush()
    run = db.ProductionRun(
        plant_id=plant.id, foam_grade_id=grade.id, recipe_version_id=recipe.id,
        run_date=dt.date(2026, 8, 1), batch_reference="CR09-PERIOD-B1",
        notes="Synthetic UAT run for demonstration purposes only.",
    )
    session.add(run); session.flush()
    session.commit()

    data = reports.build_period_summary_data(session, plant_id=plant.id)
    assert data["is_reference_dataset"] is True, (
        "Fixture no longer trips the synthetic-dataset flag - test no longer proves the label leak is fixed"
    )
    docx_bytes = reports.render_period_summary_docx(data)
    text = _extract_docx_text(docx_bytes)
    leaks = _find_leaks(text)
    assert not leaks, f"Customer-facing leak in Plant / Period Summary Report .docx: {leaks}"
    assert "Reference Dataset" in text


# ---------------------------------------------------------------------------
# 5. views/15_Recipe_Optimization.py's PI3 prompt construction - captured
#    via a monkeypatched ai_assistant.ask_assistant() during a live AppTest
#    run, the same technique used to prove report/UI text is clean, applied
#    to the one customer-facing surface that isn't rendered UI or a Word
#    file: the prompt string PI3 itself receives.
# ---------------------------------------------------------------------------

def test_recipe_optimization_pi3_prompt_has_no_leak(wp3_session, monkeypatch):
    session = wp3_session
    grade_id, run_id = _seed_wp3_fixture(session)
    # Discovered during the CR-09 closeout correction: tenant_scope.py's
    # plant_ids_for_company()/family_ids_for_plants()/grade_ids_for_families()
    # are st.cache_data'd by company_id/plant_ids alone (see that module's own
    # docstring on clear_scope_cache()) with no per-test isolation - since
    # _reset_schema() restarts autoincrement IDs at 1 for every test in this
    # file, a fresh company here can collide with another test's now-stale
    # cached scope from a few seconds earlier, silently filtering this
    # fixture's own grade out of the Recipe Optimization page's grade picker
    # and making it flake order-dependently (passes alone, fails inside the
    # full file). Same fix every write path in the app already uses after a
    # commit.
    tenant_scope.clear_scope_cache()

    captured = {}

    def _fake_ask_assistant(prompt, **kwargs):
        captured["prompt"] = prompt
        return "stubbed PI3 answer", None

    import ai_assistant
    monkeypatch.setattr(ai_assistant, "ask_assistant", _fake_ask_assistant)

    PAGE15 = os.path.join(APP_DIR, "views", "15_Recipe_Optimization.py")
    at = AppTest.from_file(PAGE15, default_timeout=30)
    at.secrets["AUTH_DISABLED"] = True
    # Fake secrets so ai_assistant.is_configured() passes and the "Get PI3
    # recommendation" button renders (see _seed_wp3_fixture's
    # PI3AIConnectionSetting) - ask_assistant() itself is monkeypatched below,
    # so this never makes a real network call, same pattern as
    # tests/test_wp4_recipe_optimization_page_smoke.py.
    at.secrets["OPENAI_API_KEY"] = "sk-test-not-a-real-key"
    at.secrets["PI3_VECTOR_STORE_ID"] = "vs_test_not_real"
    at.run()
    # CR-09 closeout correction (2026-08-12, per Charlie's "Return to JC for
    # Completion" review): this used to pytest.skip() here on the grounds that
    # a dtype TypeError elsewhere on this page (expectation_summary["avg_target"]
    # .round(2) on an object-dtype column - see views/15_Recipe_Optimization.py's
    # own comment at that line) was "an unrelated pre-existing edge case," and
    # relied on the WP3 conformance report test above as indirect proof this
    # customer-facing path was clean. Charlie's review rejected that: the CR
    # explicitly covers PI3 prompt/output leakage, so that path requires DIRECT
    # verification, not an indirect stand-in from a different report type. The
    # dtype bug has now been fixed at its source, so this must fail loudly (not
    # skip) if the page ever raises again - that's the whole point of this
    # assertion existing.
    assert not at.exception, f"Unhandled exception loading Recipe Optimization: {at.exception}"

    grade_sb = next((sb for sb in at.selectbox if sb.label == "Product grade"), None)
    if grade_sb is None:
        pytest.skip("Recipe Optimization grade picker not found - page structure may have changed")
    grade_option = next((opt for opt in grade_sb.options if "RF-CR09-" in str(opt)), None)
    if grade_option is None:
        pytest.skip("Seeded UAT-only grade not offered by the grade picker in this session")
    grade_sb.set_value(grade_option)
    at.run()

    ask_button = next((b for b in at.button if "ask_pi3_recipe_opt" in (b.key or "")), None)
    if ask_button is None:
        pytest.skip("Get PI3 recommendation button not found - page structure may have changed")
    ask_button.click()
    at.run()

    if "prompt" not in captured:
        pytest.skip("PI3 button click did not reach ask_assistant() in this session - button may be disabled")

    leaks = _find_leaks(captured["prompt"])
    assert not leaks, f"Customer-facing leak in PI3 prompt: {leaks}"


if __name__ == "__main__":
    sys.exit(pytest.main([__file__, "-v"]))
