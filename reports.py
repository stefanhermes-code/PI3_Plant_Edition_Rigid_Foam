"""Report generation: data assembly + Word (.docx) rendering.

"Report" is one of PI3 Plant Edition's own standard, always-included
capabilities (see pages/10_PI3_AI_Connectivity.py's docstring: "Standard
version (always included): Search, Compare, Retrieve, Structure, Report,
Review and Approval.") - this module is what had been missing. It is not
gated behind PI3 connectivity; every logged-in user can generate reports.

Two report types that predate the 2026-08-04 Reports redesign, each with a
data-assembly function (plain dict, no Streamlit import, easy to unit
test) and a Word (.docx) renderer (PDF/Excel renderers were removed
2026-08-04; every report now exports to Word only):

- build_period_summary_data() / render_period_summary_docx()
  One plant/product family/date range: KPIs, pass rate, recurring issues,
  the run list, and a breakdown by product grade.
- build_trial_report_data() / render_trial_report_docx()
  One closed Customer Trial or Optimization Trial (see db.CustomerTrial /
  db.OptimizationTrial - the two independent lab-trial flows, added
  2026-08-03): objective/hypothesis, what changed, outcome/conclusion,
  reviewer sign-off - a formal closeout writeup. Rebuilt 2026-08-04 to
  cover these two (the trials people actually use) after the old
  TrialRecord concept (a formal-experiment flag on a production run) was
  removed - zero real rows across 244 production runs, fully superseded
  by these two independent, self-contained closeout flows.

pages/21_Report.py wires these to selectors, an in-app preview, and
st.download_button for both file formats.

Two purpose-built reports for the Recipes page (pages/3_Recipe_Version_
Record.py), added 2026-08-04 as the first output of the app-wide Reports
redesign (see PI3_Gaps note: reports must be aggregated/purpose-built
answers to a specific question, never a raw-data dump or a PI3-narrative
document - the CSV export on every table already covers raw-data needs,
and PI3's own Word download on relevant pages already covers narrative):

- build_recipe_formulation_record_data() / render_recipe_formulation_record_pdf()
  / render_recipe_formulation_record_excel()
  One recipe version: the formulation itself (materials/php/supplier/
  role), its quality specs vs. actual results aggregated over a chosen
  date range (across every production run/customer trial/optimization
  trial built on this recipe version), and cost per kg - "is this
  formulation meeting spec, and what does it cost" in one document for
  internal use/approval (not customer-facing - a customer-facing version
  would need to omit the formulation itself).
- build_where_used_report_data() / render_where_used_report_pdf()
  / render_where_used_report_excel()
  One raw material: every recipe version (active and retired) that uses
  it with its php/role, the target properties of every product grade
  affected, and any Customer/Optimization Trial precedent tied to a
  recipe version containing it - "if I replace this material, what's
  affected and what trials already exist to lean on."

A third purpose-built report, added 2026-08-04 for one production run -
REPLACING the earlier build_run_report_data() / render_run_report_pdf() /
render_run_report_excel() (removed the same day: a flat header plus four
raw tables - recipe components, process settings, quality results,
quality issues - with no synthesis, exactly the "factual but not adding
value" pattern the whole Reports redesign exists to fix). Lives on the
Report page (pages/21_Report.py), not the Production Run page itself: per
user direction, a report whose subject is a single simple choice (pick
one run from a dropdown) belongs on the Report page alongside the other
selector-driven reports; a report whose subject needs a comprehensive
multi-field selection first (date range, product grade, etc. - see the
Quality Test Result report below) belongs on its own page instead, next
to where that selection naturally happens.

- build_batch_release_record_data() / render_batch_release_record_pdf()
  / render_batch_release_record_excel()
  One production run: the recipe used (in full - not just a reference),
  and a rolled-up quality conformance verdict (Pass/Fail per tested
  property plus one overall Conforming/Non-conforming/Incomplete verdict)
  and any quality issues recorded. If - and only if - a flag is raised
  (a failed result or a recorded quality issue), the report widens to
  pull supporting context from every other tab on the Production Run
  page: the Finalized phase's actual component stream readings (with
  any non-"Valid" calibration status called out), and any Production
  Events logged during the run - "does this batch look wrong, and if so
  what else was going on at the time" in one document, not five
  separate tab exports. A clean run stays a short document; a flagged
  one pulls in exactly what's relevant, not everything that exists.
  (Setup-vs-Finalized fall-plate position deviations were part of this
  widened context before WP7 Phase 0, 2026-08-13, removed the active
  fall-plate section-position sub-workflow; this report has not claimed
  that context since - see the "Batch Release / Conformance Record"
  section comment below for the corresponding dead-code cleanup.)

A fifth report, added 2026-08-04, lives on its own page rather than the
Report page - per the same placement principle stated above, this report's
subject is a comprehensive multi-field selection (Pass/Fail, Property, and
Foam scope), not a single dropdown choice, so it lives on
pages/5_Physical_Property_Result.py, right below the filter controls and
the existing on-page Pareto chart it shares its scope with:

- build_quality_test_report_data() / render_quality_test_report_pdf()
  / render_quality_test_report_excel()
  Takes the exact set of PhysicalPropertyResult rows the page has already
  scoped (tenant) and filtered (Pass/Fail, Property, Foam scope) - the
  report never re-derives its own selection, so it always matches what's
  on screen. Aggregates that set into a pass-rate summary, a failures-by-
  property breakdown (bar chart), a pass/fail-by-foam-grade breakdown
  (bar chart, only shown when the selection spans more than one grade),
  and a curated table of just the failing results (target/actual/
  deviation) - not the full underlying row set, which the page's own CSV
  export already covers.

A sixth report, added 2026-08-04, follows the same placement logic as the
Quality Test Result report - its subject is a comprehensive multi-field
selection (Severity, Foam scope, and the breakdown's group-by choice), not
a single dropdown choice, so it lives on pages/6_Quality_Observation.py
itself, right below that page's own filters and breakdown-by-issue chart:

- build_quality_issue_report_data() / render_quality_issue_report_pdf()
  / render_quality_issue_report_excel()
  Takes the exact set of QualityObservation rows the page has already
  scoped and filtered (Severity, Foam scope) - never re-derives its own
  selection. Aggregates that set into a severity/recurring-vs-one-off
  summary, a confidence-level breakdown, an issues-by-type-or-category
  breakdown (bar chart, matching whichever grouping the page's own
  breakdown chart was using), and a curated table of just the priority
  issues (High severity and/or Recurring) - not every row, which the
  page's own CSV export already covers.

An eighth report, added 2026-08-04, follows the same placement logic as
Batch Release Record and Trial Closeout Report - picking one sample is a
single simple choice, so it lives on the Report page rather than pages
9/11/12 (which already have the aggregate, multi-field-selection Sample
Report added the same day):

- build_sample_certificate_data() / render_sample_certificate_pdf()
  / render_sample_certificate_excel()
  One sample, across any of the three sources (Production Run / Customer
  Trial / Optimization Trial): where it came from (header context), the
  sample itself (zone, sampled time, notes), the recipe used in full (not
  customer-facing as-is, same caveat as Recipe Formulation Record and
  Batch Release Record), that sample's own quality test results, and a
  rolled-up Conforming/Non-conforming/Incomplete verdict - a per-sample
  certificate of analysis, distinct from the aggregate Sample Report
  (which covers many samples at once via charts, no per-sample detail).

A ninth, narrower report type lives here too:

- build_pi3_qa_report_data() / render_pi3_qa_report_docx()
  A single "Ask PI3" question-and-answer exchange (see
  helpers.render_ask_pi3_section) - the question, PI3's answer, and an
  appendix of the exact data PI3 checked to produce it (SQL + rows
  returned, or the verified-analysis arguments and result). Like every
  other report in this module, this is Word (.docx) only, and it is always
  built from this same code path, so every export has identical
  formatting regardless of who generates it or what was asked - a
  hand-maintained Word template would drift over time; this can't.

A tenth through fifteenth report, added 2026-08-04 for the Industrial
Intelligence section (per user direction: pages with a PI3 recommendation/
interpretation get a Context/Analysis/Conclusions report ON the page
itself, built from that page's own deterministic analysis - never the PI3
answer, which already has its own separate Word download via
render_pi3_docx_download). Each build_*_report_data() takes the exact
already-computed analysis object(s) the calling page has on screen at that
moment - ranked DataFrame, SPC results, diff list, etc. - and never
re-derives them, so the report always matches what's on screen:

- build_recipe_optimization_report_data() / render_recipe_optimization_report_pdf()
  / render_recipe_optimization_report_excel() (pages/15_Recipe_Optimization.py)
  Current formulation cost, whether the current recipe meets target per
  property, and the top ingredient-dosage correlation for the selected
  property.
- build_trend_analysis_report_data() / render_trend_analysis_report_pdf()
  / render_trend_analysis_report_excel() (pages/16_Trend_Analysis.py)
  Control chart, process capability, CUSUM, and trend-test results, plus
  the machine-change/quality-issue timeline - the full SPC toolkit for one
  property.
- build_correlation_report_data() / render_correlation_report_pdf()
  / render_correlation_report_excel() (pages/17_Process_Property_Correlation.py)
  The ranked machine-setting-vs-property correlation table.
- build_root_cause_report_data() / render_root_cause_report_pdf()
  / render_root_cause_report_excel() (pages/18_Root_Cause_Assistant.py)
  The deterministic run-vs-prior-run diff (recipe/machine/process-setting
  shifts) behind a flagged quality issue.
- build_machine_settings_report_data() / render_machine_settings_report_pdf()
  / render_machine_settings_report_excel() (pages/19_Machine_Settings_Optimization.py)
  The ranked setting-optimization table (which range of each setting
  separates good outcomes from bad).
- build_expert_notes_report_data() / render_expert_notes_report_pdf()
  / render_expert_notes_report_excel() (pages/20_Expert_Notes.py)
  Doesn't fit the pattern above (no PI3 recommendation of its own) - an
  always-visible aggregate breakdown of the notes already shown on the
  page, by confidence level, source (Manual vs. PI3), and linked-entity
  type, distinct from the existing conditional per-note Word re-download
  (kept as-is - see that button's own comment on pages/20_Expert_Notes.py).
"""

import datetime as dt
import io
import math
import os
import re

import pandas as pd
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.graphics.charts.barcharts import VerticalBarChart
from reportlab.graphics.charts.linecharts import HorizontalLineChart
from reportlab.graphics.shapes import Drawing
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

import customer_presentation
from analytics import (
    production_run_output_summary,
    production_run_process_parameters,
    recipe_version_cost,
)
from db import (
    CONFIDENCE_LEVELS,
    SEVERITIES,
    ComponentStreamReading,
    CustomerTrial,
    FoamGrade,
    FoamGradeTargetProperty,
    GradeSpecification,
    Machine,
    OptimizationTrial,
    Plant,
    PhysicalPropertyMethod,
    PhysicalPropertyResult,
    ProductFamily,
    ProductionEvent,
    ProductionMethod,
    ProductionPhase,
    ProductionRun,
    QualityObservation,
    RawMaterial,
    RecipeComponent,
    RecipeVersion,
    Sample,
)
import quality_issue_taxonomy
from quality_standards import compute_pass_fail
import wp3_conformance

# CR-22 / F22-04 (AF22-01): mirrors helpers.BLOCK_REFERENCE_METHOD_
# CONTROLLED_ID / helpers.block_reference_applicable() - duplicated here
# (not imported) because helpers.py already does `import reports` at
# module scope, and reports.py importing helpers back would be a
# circular import (helpers.py's own module-level statements run before
# its function definitions exist, so a `from helpers import ...` reached
# via that cycle would fail). Both constants must be changed together if
# this ever changes.
_BLOCK_REFERENCE_METHOD_CONTROLLED_ID = "PM-500"


def _block_reference_applicable(production_method):
    """True only for PM-500 Rigid Block Production - see the module-level
    comment above and helpers.block_reference_applicable()."""
    return bool(production_method) and production_method.controlled_id == _BLOCK_REFERENCE_METHOD_CONTROLLED_ID

STYLES = getSampleStyleSheet()
STYLES.add(ParagraphStyle(name="Small", parent=STYLES["Normal"], fontSize=8, leading=10))
STYLES.add(ParagraphStyle(
    name="SmallBold", parent=STYLES["Normal"], fontSize=8, leading=10, fontName="Helvetica-Bold",
))


# ---------------------------------------------------------------------------
# Shared rendering helpers
# ---------------------------------------------------------------------------

def _pdf_bytes(build_story):
    """build_story(story: list) appends reportlab flowables to story."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4, topMargin=18 * mm, bottomMargin=18 * mm,
        leftMargin=16 * mm, rightMargin=16 * mm,
    )
    story = []
    build_story(story)
    doc.build(story)
    return buf.getvalue()


def _build_bar_chart_drawing(categories, values, width=460, height=170,
                              bar_color=colors.HexColor("#4A7A9D"), zero_floor=True):
    """Pure Drawing-builder for a vertical bar chart, shared by the PDF
    flowable (_bar_chart) and the Word rasterized-PNG embed (_docx_bar_chart)
    below - kept as one function so both formats always draw the identical
    chart, never two hand-maintained near-duplicates."""
    drawing = Drawing(width, height)
    chart = VerticalBarChart()
    chart.x = 45
    chart.y = 30
    chart.height = height - 60
    chart.width = width - 65
    chart.data = [values]
    chart.categoryAxis.categoryNames = [str(c)[:16] for c in categories]
    chart.categoryAxis.labels.angle = 30
    chart.categoryAxis.labels.dy = -12
    chart.categoryAxis.labels.dx = -6
    chart.categoryAxis.labels.fontSize = 7
    if zero_floor:
        chart.valueAxis.valueMin = 0
    chart.bars[0].fillColor = bar_color
    drawing.add(chart)
    return drawing


def _bar_chart(story, title, categories, values, note=None, width=460, height=170,
                bar_color=colors.HexColor("#4A7A9D"), zero_floor=True):
    """A simple vertical bar chart flowable for the PDF - categories/values
    are same-length parallel lists. Used wherever a report should show a
    breakdown at a glance rather than force the reader to scan a table for
    it (e.g. failures by property).

    zero_floor=True (the default, matching every existing caller) pins the
    value axis to start at 0 - correct for counts, which are never
    negative. Callers charting a signed metric (e.g. % deviation from
    target, a correlation coefficient) must pass zero_floor=False so
    reportlab auto-scales the axis instead - pinning min to 0 would clip
    negative bars off the chart entirely."""
    story.append(Spacer(1, 8))
    story.append(Paragraph(title, STYLES["Heading3"]))
    if note:
        story.append(_p(note))
    if not categories or not any(v not in (None, 0) for v in values):
        story.append(_p("No data recorded."))
        return
    story.append(_build_bar_chart_drawing(categories, values, width, height, bar_color, zero_floor))


_LINE_COLOR_HEX = ["#4A7A9D", "#C0392B", "#7F8C8D", "#27AE60"]
_LINE_COLORS = [colors.HexColor(h) for h in _LINE_COLOR_HEX]


def _build_line_chart_drawing(categories, series, width=460, height=180):
    """Pure Drawing-builder for a multi-line chart, shared by the PDF
    flowable (_line_chart) and the Word rasterized-PNG embed (_docx_line_chart)
    below - see _build_bar_chart_drawing's docstring for why this is split
    out rather than duplicated."""
    drawing = Drawing(width, height)
    chart = HorizontalLineChart()
    chart.x = 45
    chart.y = 35
    chart.height = height - 65
    chart.width = width - 65
    chart.data = [vals for _, vals in series]
    chart.categoryAxis.categoryNames = [str(c)[:10] for c in categories]
    chart.categoryAxis.labels.angle = 30
    chart.categoryAxis.labels.dy = -12
    chart.categoryAxis.labels.dx = -6
    chart.categoryAxis.labels.fontSize = 6
    for i, (_, vals) in enumerate(series):
        color = _LINE_COLORS[i % len(_LINE_COLORS)]
        chart.lines[i].strokeColor = color
        chart.lines[i].strokeWidth = 1.5
    drawing.add(chart)
    return drawing


def _line_chart(story, title, categories, series, note=None, width=460, height=180):
    """A simple multi-line chart flowable for the PDF - categories is the
    shared x-axis labels (e.g. tested_at dates), series is an ordered dict/
    list of (label, values) pairs, each a same-length list of y-values
    (None gaps are not supported by reportlab's HorizontalLineChart, so
    callers should pre-filter to rows where every series has a value).
    Used for genuine time-series (Trend Analysis's control chart / CUSUM)
    where a bar chart would flatten the thing that actually matters - the
    shape of the line over time."""
    story.append(Spacer(1, 8))
    story.append(Paragraph(title, STYLES["Heading3"]))
    if note:
        story.append(_p(note))
    if not categories or not series or not any(any(v is not None for v in vals) for _, vals in series):
        story.append(_p("No data recorded."))
        return
    drawing = _build_line_chart_drawing(categories, series, width, height)
    story.append(drawing)
    legend_bits = []
    for i, (label, _) in enumerate(series):
        legend_bits.append(f'<font color="{_LINE_COLORS[i % len(_LINE_COLORS)].hexval()}">■</font> {label}')
    story.append(_p("  ".join(legend_bits)))


def _p(text, style="Normal"):
    """Paragraph with basic XML-escaping, since free-text fields (notes,
    hypotheses, ...) may contain '&', '<', '>' which reportlab's
    Paragraph would otherwise try to interpret as markup."""
    if text is None:
        text = "—"
    escaped = (
        str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    )
    return Paragraph(escaped, STYLES[style])


def _key_value_table(pairs, col_widths=(35 * mm, 55 * mm, 35 * mm, 55 * mm)):
    """pairs: list of (label, value) - rendered two-per-row as a compact
    header block (label/value/label/value)."""
    rows = []
    for i in range(0, len(pairs), 2):
        row = []
        for label, value in pairs[i:i + 2]:
            row.extend([label, "—" if value in (None, "") else str(value)])
        while len(row) < 4:
            row.append("")
        rows.append(row)
    t = Table(rows, colWidths=list(col_widths))
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME", (2, 0), (2, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    return t


def _section(story, title, rows, col_widths=None):
    """A heading followed by a table built from a list of dicts (all
    sharing the same keys), or a plain "no data" note if rows is empty."""
    story.append(Spacer(1, 8))
    story.append(Paragraph(title, STYLES["Heading3"]))
    if not rows:
        story.append(_p("No data recorded."))
        return
    headers = list(rows[0].keys())
    table_rows = [headers]
    for row in rows:
        table_rows.append(["—" if row.get(h) in (None, "") else str(row.get(h)) for h in headers])
    kwargs = {"colWidths": col_widths} if col_widths else {}
    t = Table(table_rows, repeatRows=1, **kwargs)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DCE6EC")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B0BEC5")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.append(t)


def _wrapped_section(story, title, rows, columns, col_widths):
    """Like _section, but every cell is a Paragraph so a long free-text
    column (e.g. "Suspected cause") wraps within its column instead of
    overflowing the page - _section's plain-string cells don't wrap,
    which is fine for short values but breaks down once a table mixes a
    long descriptive column with several rows. columns is the explicit,
    ordered list of dict keys to display - lets the caller show a
    narrower slice of a wider dict (dropping a couple of less-essential
    columns) to keep the table print-width, all without touching the
    underlying data itself. col_widths must be supplied (no auto-sizing)
    and should sum to the usable page width."""
    story.append(Spacer(1, 8))
    story.append(Paragraph(title, STYLES["Heading3"]))
    if not rows:
        story.append(_p("No data recorded."))
        return
    table_rows = [[_p(c, style="SmallBold") for c in columns]]
    for row in rows:
        table_rows.append([_p(row.get(c), style="Small") for c in columns])
    t = Table(table_rows, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DCE6EC")),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B0BEC5")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.append(t)


def _title_block(story, title, subtitle=None):
    story.append(Paragraph(title, STYLES["Title"]))
    if subtitle:
        story.append(Paragraph(subtitle, STYLES["Normal"]))
    story.append(Spacer(1, 6))


def plant_label(session, plant_id):
    if plant_id is None:
        return "All plants"
    p = session.get(Plant, plant_id)
    return p.name if p else "—"


def product_family_label(session, product_family_id):
    if product_family_id is None:
        return "All product families"
    f = session.get(ProductFamily, product_family_id)
    return f.name if f else "—"


# ---------------------------------------------------------------------------
# 1. Plant / Period Summary Report
# ---------------------------------------------------------------------------

def build_period_summary_data(
    session, plant_id=None, product_family_id=None, date_from=None, date_to=None,
    allowed_plant_ids=None, production_method_id=None,
):
    """allowed_plant_ids is the tenant-scope guardrail (see tenant_scope.py):
    None = unfiltered (platform owner viewing "All companies"), otherwise the
    list of plant ids the calling company is allowed to see - applied
    unconditionally, on top of whatever single-plant choice plant_id
    represents. Without this, a non-owner user who leaves the on-screen
    "Plant" selector at its default "All plants" would get a report across
    every plant in the database, not just their own company's.

    production_method_id (added 2026-08-10, per Charlie's flat-PM technical
    completion instruction): an optional isolation filter against each run's
    own immutable Production Method snapshot (ProductionRun.production_
    method_id - see db.py) - this report pools every run in a plant/date
    range, which can span more than one Production Method once a plant
    activates a second one, so pass-rate/quality-issue totals could
    otherwise silently blend two methods' history together. None = every
    method pooled (the report still surfaces the split via
    method_breakdown below, so the pooling is visible either way)."""
    runs_q = session.query(ProductionRun)
    if allowed_plant_ids is not None:
        runs_q = runs_q.filter(ProductionRun.plant_id.in_(allowed_plant_ids))
    if plant_id:
        runs_q = runs_q.filter(ProductionRun.plant_id == plant_id)
    if product_family_id:
        runs_q = runs_q.join(FoamGrade, ProductionRun.foam_grade_id == FoamGrade.id).filter(
            FoamGrade.product_family_id == product_family_id
        )
    if date_from:
        runs_q = runs_q.filter(ProductionRun.run_date >= date_from)
    if date_to:
        runs_q = runs_q.filter(ProductionRun.run_date <= date_to)
    if production_method_id:
        runs_q = runs_q.filter(ProductionRun.production_method_id == production_method_id)
    runs = runs_q.order_by(ProductionRun.run_date).all()
    run_ids = [r.id for r in runs]

    results = (
        session.query(PhysicalPropertyResult)
        .filter(PhysicalPropertyResult.production_run_id.in_(run_ids)).all()
        if run_ids else []
    )
    # WP6-S09 fix (2026-08-09, UAT-012 per Charlie's review): a rigid-foam
    # grade's Pass/Fail is governed by its own GradeSpecification rows
    # (method/condition/orientation-aware limits - see wp3_conformance.py),
    # not the flexible app's flat compute_pass_fail(property_name,
    # target_value, actual_value). Every real rigid PhysicalPropertyResult
    # has target_value = NULL (the ceiling/floor lives in the spec's own
    # upper_limit/lower_limit instead), so the flat path silently scored
    # every real result as unresolved and this report's pass rate was
    # always blank against real data. Uses the same _is_rigid_grade /
    # compute_conformance_report pattern as build_batch_release_record_data
    # (UAT-011), aggregated per (grade, run) the same way
    # wp3_conformance.compute_grade_conformance_summary already does
    # app-wide; a grade with no chemistry_id (flexible-style, none exist in
    # this app today) keeps the original flat-property scoring as a
    # fallback.
    pass_count = 0
    fail_count = 0
    unresolved_count = 0  # EXCLUDED_CONTEXT / INVALID / NO_RESULT rows - see note below
    total_checks = 0
    runs_by_grade = {}
    for r in runs:
        runs_by_grade.setdefault(r.foam_grade_id, []).append(r)

    for grade_id, grade_runs in runs_by_grade.items():
        grade = session.get(FoamGrade, grade_id)
        if _is_rigid_grade(grade):
            for run in grade_runs:
                for row in wp3_conformance.compute_conformance_report(session, grade_id, production_run_id=run.id):
                    total_checks += 1
                    if row["verdict"] == "Pass":
                        pass_count += 1
                    elif row["verdict"] == "Fail":
                        fail_count += 1
                    else:
                        unresolved_count += 1
        else:
            grade_run_ids = {r.id for r in grade_runs}
            legacy_results = [r for r in results if r.production_run_id in grade_run_ids]
            legacy_verdicts = [compute_pass_fail(r.property_name, r.target_value, r.actual_value) for r in legacy_results]
            total_checks += len(legacy_verdicts)
            pass_count += legacy_verdicts.count("Pass")
            fail_count += legacy_verdicts.count("Fail")
            unresolved_count += sum(1 for v in legacy_verdicts if v not in ("Pass", "Fail"))

    total_scored = pass_count + fail_count
    pass_rate = round(100 * pass_count / total_scored) if total_scored else None
    # WP6-S09 (UAT-012, per Charlie's review): pass_rate on its own is
    # dangerously misleading when most grade-specification checks couldn't
    # be evaluated at all - e.g. 1 evaluable check that happens to Pass
    # reads as "100%" even when 48 other checks in the same period are
    # EXCLUDED_CONTEXT/INVALID/NO_RESULT. total_checks/unresolved_count let
    # the report say plainly how much of the picture that percentage
    # actually covers.
    coverage_pct = round(100 * total_scored / total_checks) if total_checks else None

    # WP6-S09 (UAT-012): flag the report clearly when the date range
    # includes synthetic/demonstration runs, rather than presenting a
    # pass-rate/quality-issue summary that reads as real plant history.
    # Every real seeded production run's notes field states this
    # explicitly ("Synthetic UAT run", "Synthetic end-to-end data set") -
    # there is no dedicated is_synthetic column on ProductionRun (see
    # db.py), so that note text is the authoritative signal.
    #
    # CR-09 (2026-08-12): this used to build the customer-visible label
    # text directly here ("Synthetic UAT / Reference Dataset" - both
    # "Synthetic" and "UAT" are internal engineering terms). It now stores
    # only the internal boolean signal; render_period_summary_report_docx
    # generates the customer-safe label via customer_presentation.
    # customer_facing_reference_dataset_label(), the single shared
    # translation point for this flag.
    is_reference_dataset = any(
        "synthetic" in (r.notes or "").lower() for r in runs
    )

    observations = (
        session.query(QualityObservation)
        .filter(QualityObservation.production_run_id.in_(run_ids)).all()
        if run_ids else []
    )
    recurring = [o for o in observations if o.frequency == "Recurring"]

    run_rows = [
        {
            "Run ID": r.id, "Date": r.run_date,
            "Product grade": r.foam_grade.grade_name if r.foam_grade else "—",
            "Recipe version": r.recipe_version.version_label if r.recipe_version else "—",
            "Production Unit or Cell": r.machine.name if r.machine else "—",
            "Production Method": r.production_method.name if r.production_method else "—",
            "Batch reference": r.batch_reference or "—",
        }
        for r in runs
    ]

    issue_rows = [
        {
            "Observed": o.observed_at, "Run": o.production_run_id, "Issue type": o.observation_type,
            "Severity": o.severity or "—", "Frequency": o.frequency or "—",
            "Confidence": o.confidence_level or "—",
        }
        for o in observations
    ]

    grade_counts = {}
    for r in runs:
        gname = r.foam_grade.grade_name if r.foam_grade else "—"
        grade_counts[gname] = grade_counts.get(gname, 0) + 1
    grade_breakdown = [{"Product grade": k, "Production runs": v} for k, v in sorted(grade_counts.items())]

    # Production Method breakdown (added 2026-08-10): shown regardless of
    # whether production_method_id isolates the report to one method - when
    # unfiltered, this is what makes an "All methods pooled" report legible
    # rather than silently blending PM-100 and PM-200 runs into one number.
    method_counts = {}
    for r in runs:
        mname = r.production_method.name if r.production_method else "—"
        method_counts[mname] = method_counts.get(mname, 0) + 1
    method_breakdown = [
        {"Production Method": k, "Production runs": v} for k, v in sorted(method_counts.items())
    ]
    production_method_label_text = "All methods"
    if production_method_id:
        pm = session.get(ProductionMethod, production_method_id)
        production_method_label_text = pm.name if pm else "All methods"

    return {
        "plant": plant_label(session, plant_id),
        "product_family": product_family_label(session, product_family_id),
        "production_method": production_method_label_text,
        "date_from": date_from,
        "date_to": date_to,
        "is_reference_dataset": is_reference_dataset,
        "total_runs": len(runs),
        "pass_rate": pass_rate,
        "total_results_scored": total_scored,
        "total_checks_attempted": total_checks,
        "unresolved_checks": unresolved_count,
        "coverage_pct": coverage_pct,
        "total_quality_issues": len(observations),
        # WP6-S09 (UAT-012): the 12 controlled UAT failure cases (WP3 Gate 2)
        # live as computed Fail verdicts from GradeSpecification/
        # PhysicalPropertyResult (see pass_rate above), not as rows in this
        # QualityObservation-backed count - labeled explicitly so this
        # number is never mistaken for the full set of quality failures in
        # the period.
        "quality_issues_label": "Recorded production quality issues",
        "recurring_issues": len(recurring),
        "runs": run_rows,
        "quality_issues": issue_rows,
        "grade_breakdown": grade_breakdown,
        "method_breakdown": method_breakdown,
    }


def render_period_summary_pdf(data):
    def build(story):
        _title_block(
            story, "Plant / Period Summary Report",
            f"{data['plant']} · {data['product_family']} · {data['date_from'] or 'earliest'} to {data['date_to'] or 'latest'}",
        )
        story.append(_key_value_table([
            ("Production runs", data["total_runs"]),
            ("Quality test pass rate", f"{data['pass_rate']}%" if data["pass_rate"] is not None else "—"),
            ("Quality issues", data["total_quality_issues"]),
            ("Recurring quality issues", data["recurring_issues"]),
        ], col_widths=(45 * mm, 40 * mm, 45 * mm, 40 * mm)))
        _section(story, "Production runs in range", data["runs"])
        _section(story, "Quality issues in range", data["quality_issues"])
        _section(story, "Breakdown by product grade", data["grade_breakdown"])
    return _pdf_bytes(build)


def render_period_summary_docx(data):
    doc = Document()
    subtitle = f"{data['plant']} · {data['product_family']} · {data['date_from'] or 'earliest'} to {data['date_to'] or 'latest'}"
    if data.get("production_method") and data["production_method"] != "All methods":
        subtitle = f"{subtitle} · Production Method: {data['production_method']}"
    reference_dataset_label = customer_presentation.customer_facing_reference_dataset_label(
        data.get("is_reference_dataset")
    )
    if reference_dataset_label:
        subtitle = f"{subtitle} · {reference_dataset_label}"
    _docx_report_header(doc, "Plant / Period Summary Report", subtitle)
    issue_label = data.get("quality_issues_label", "Quality issues")
    # WP6-S09 fix (2026-08-09, per Charlie's WP6 sequence item 3, "correct
    # the UAT-012 headline metrics"): the headline block itself used to lead
    # with a bare "Quality test pass rate: 100%" and relegate coverage to a
    # paragraph below it - a reader who only looks at the headline table
    # (the whole point of a headline) never sees that the 100% came from 1
    # of 49 attempted checks. The headline now carries the coverage figures
    # directly, and the pass-rate row's own label says what it's a rate of,
    # rather than reading as an unqualified plant-wide number.
    headline_rows = [("Production runs", data["total_runs"])]
    if data.get("total_checks_attempted"):
        headline_rows.extend([
            ("Grade-specification checks attempted", data["total_checks_attempted"]),
            ("Checks evaluated to Pass/Fail", data["total_results_scored"]),
            ("Coverage (evaluated / attempted)", f"{data.get('coverage_pct')}%" if data.get("coverage_pct") is not None else "—"),
            (
                "Quality test pass rate (of checks evaluated)",
                f"{data['pass_rate']}%" if data["pass_rate"] is not None else "—",
            ),
        ])
    else:
        headline_rows.append(
            ("Quality test pass rate", f"{data['pass_rate']}%" if data["pass_rate"] is not None else "—")
        )
    headline_rows.extend([
        (issue_label, data["total_quality_issues"]),
        ("Recurring quality issues", data["recurring_issues"]),
    ])
    _docx_kv_table(doc, headline_rows)
    if data.get("total_checks_attempted"):
        doc.add_paragraph(
            f"Pass rate is calculated over {data['total_results_scored']} of "
            f"{data['total_checks_attempted']} grade-specification checks attempted "
            f"({data.get('coverage_pct')}% coverage); the remaining "
            f"{data['unresolved_checks']} could not be evaluated to Pass/Fail "
            "(excluded context, invalid/incomplete result context, or no matching "
            "result) and are not counted in either direction. A pass rate near a "
            "small evaluable count is not a statement about the untested majority."
        )
    _docx_section(doc, "Production runs in range", data["runs"])
    _docx_section(doc, f"{issue_label} in range", data["quality_issues"])
    _docx_section(doc, "Breakdown by product grade", data["grade_breakdown"])
    _docx_section(doc, "Breakdown by Production Method", data.get("method_breakdown") or [])
    return _docx_bytes(doc)


# ---------------------------------------------------------------------------
# 2. Trial Closeout Report
#
# Covers the two independent lab-trial flows (see db.CustomerTrial /
# db.OptimizationTrial, added 2026-08-03) - NOT the old TrialRecord concept
# (a formal-experiment flag on a production run), which was removed
# 2026-08-04 after confirming zero real rows across 244 production runs.
# The two models have different closeout fields (a sales-driven Customer
# Trial has customer_name/outcome/customer_feedback; an internally-driven
# Optimization Trial has hypothesis/conclusion/reuse_recommendation), so
# build_trial_report_data() normalizes both into one common "narrative
# fields" list of (label, value) pairs the renderer can walk without
# needing to know which trial type produced them.
# ---------------------------------------------------------------------------

def build_trial_report_data(session, source_type, trial_id):
    """source_type is "Customer Trial" or "Optimization Trial" (see
    db.SAMPLE_SOURCE_TYPES) - the same source-type string used throughout
    the app to disambiguate the three mutually-exclusive parents a sample/
    quality result can belong to."""
    if source_type == "Customer Trial":
        trial = session.get(CustomerTrial, trial_id)
    elif source_type == "Optimization Trial":
        trial = session.get(OptimizationTrial, trial_id)
    else:
        return None
    if trial is None:
        return None
    grade = trial.foam_grade
    plant = trial.plant

    quality_issues = [
        {
            "Issue type": o.observation_type, "Severity": o.severity or "—",
            "Frequency": o.frequency or "—", "Confidence": o.confidence_level or "—",
        }
        for o in session.query(QualityObservation).filter(
            (QualityObservation.customer_trial_id == trial_id)
            if source_type == "Customer Trial"
            else (QualityObservation.optimization_trial_id == trial_id)
        ).all()
    ]

    if source_type == "Customer Trial":
        narrative_fields = [
            ("Customer", trial.customer_name),
            ("Sales opportunity reference", trial.sales_opportunity_reference or "—"),
            ("Requested by", trial.requested_by or "—"),
            ("Trial objective", trial.trial_objective or "—"),
            ("Outcome", trial.outcome or "—"),
            ("Customer feedback", trial.customer_feedback or "—"),
            ("Follow-up action", trial.follow_up_action or "—"),
        ]
    else:
        narrative_fields = [
            ("Improvement initiative reference", trial.improvement_initiative_reference or "—"),
            ("Hypothesis", trial.hypothesis or "—"),
            ("What changed", trial.what_changed or "—"),
            ("Result against target", trial.result_against_target or "—"),
            ("Conclusion", trial.conclusion or "—"),
            ("Reuse recommendation", trial.reuse_recommendation or "—"),
        ]

    return {
        "source_type": source_type,
        "trial_id": trial.id,
        "foam_grade": grade.grade_name if grade else "—",
        "plant": plant.name if plant else "—",
        "status": trial.status,
        "responsible_person": trial.responsible_person or "—",
        "trial_date": trial.trial_date,
        "batch_reference": trial.batch_reference or "—",
        "notes": trial.notes or "",
        "narrative_fields": narrative_fields,
        "reviewed_by": trial.reviewed_by or "—",
        # Only OptimizationTrial has a separate approved_by (CustomerTrial's
        # closeout is reviewed_by only - see db.py's REQUIRED_CLOSEOUT_FIELDS
        # on each model).
        "approved_by": getattr(trial, "approved_by", None) or "—",
        "date_closed": trial.date_closed,
        "quality_issues": quality_issues,
    }


def render_trial_report_pdf(data):
    def build(story):
        _title_block(
            story, f"Trial Closeout Report — {data['source_type']} #{data['trial_id']}",
            f"{data['foam_grade']} · {data['plant']} · {data['status']}",
        )
        story.append(_key_value_table([
            ("Status", data["status"]), ("Responsible", data["responsible_person"]),
            ("Product grade", data["foam_grade"]), ("Plant", data["plant"]),
            ("Trial date", data["trial_date"]), ("Batch reference", data["batch_reference"]),
            ("Reviewed by", data["reviewed_by"]), ("Approved by", data["approved_by"]),
            ("Date closed", data["date_closed"]), ("", ""),
        ]))
        if data["notes"]:
            story.append(Spacer(1, 6))
            story.append(_p(f"Notes: {data['notes']}"))
        story.append(Spacer(1, 8))
        for label, value in data["narrative_fields"]:
            story.append(Paragraph(label, STYLES["Heading3"]))
            story.append(_p(value))
            story.append(Spacer(1, 6))
        _section(story, "Quality issues observed", data["quality_issues"])
    return _pdf_bytes(build)


def render_trial_report_docx(data):
    doc = Document()
    _docx_report_header(
        doc, f"Trial Closeout Report — {data['source_type']} #{data['trial_id']}",
        f"{data['foam_grade']} · {data['plant']} · {data['status']}",
    )
    _docx_kv_table(doc, [
        ("Status", data["status"]), ("Responsible", data["responsible_person"]),
        ("Product grade", data["foam_grade"]), ("Plant", data["plant"]),
        ("Trial date", data["trial_date"]), ("Batch reference", data["batch_reference"]),
        ("Reviewed by", data["reviewed_by"]), ("Approved by", data["approved_by"]),
        ("Date closed", data["date_closed"]),
    ])
    if data["notes"]:
        doc.add_paragraph(f"Notes: {data['notes']}")
    for label, value in data["narrative_fields"]:
        _docx_heading(doc, label, size=11.5, color=_HTC_GREY, space_before=10)
        doc.add_paragraph(value or "—")
    _docx_section(doc, "Quality issues observed", data["quality_issues"])
    return _docx_bytes(doc)


# ---------------------------------------------------------------------------
# 3. Recipe / Formulation Record Report
#
# Internal-use record for one recipe version: the formulation, its quality
# specs vs. actual results aggregated over a chosen date range, and cost
# per kg. NOT customer-facing as-is (see reports.py module docstring) -
# the recipe/formulation section is the whole reason it can't be handed
# to a customer.
#
# "Quality specs & results" pulls actual PhysicalPropertyResult rows from
# every production run, customer trial, and optimization trial built on
# this recipe version (the three mutually-exclusive parents - see
# db.SAMPLE_SOURCE_TYPES), filtered to the caller's date range, then
# aggregated per property (average actual, pass rate, sample count)
# against that product grade's target - never a flat row-by-row dump, per
# the Reports redesign ruling that raw data belongs in each page's own
# CSV export, not in a report.
# ---------------------------------------------------------------------------

def _recipe_version_target_properties(grade):
    """FoamGrade's target specs as a flat list of {property_name, target_value,
    unit} dicts - density/hardness (fixed columns on FoamGrade) plus any
    additional targets recorded in FoamGradeTargetProperty."""
    if grade is None:
        return []
    targets = []
    if grade.target_density is not None:
        targets.append({"property_name": "Density", "target_value": grade.target_density, "unit": "kg/m³"})
    if grade.target_hardness is not None:
        # "40% IFD / hardness" is the canonical property_name used app-wide
        # (see quality_standards.INDUSTRY_TOLERANCES and pages/15_Recipe_
        # Optimization.py's own target_by_name dict) - matching it here,
        # rather than a differently-worded label, is what lets this target
        # line up with actual PhysicalPropertyResult rows recorded against
        # that same property name instead of showing as two separate rows.
        targets.append({"property_name": "40% IFD / hardness", "target_value": grade.target_hardness, "unit": "N"})
    for tp in grade.target_properties:
        targets.append({"property_name": tp.property_name, "target_value": tp.target_value, "unit": tp.unit or ""})
    return targets


def _property_results_for_recipe_version(session, recipe_version_id, date_from=None, date_to=None):
    """Every PhysicalPropertyResult tied (via production run, customer trial,
    or optimization trial - see db.SAMPLE_SOURCE_TYPES) to this recipe
    version, filtered to [date_from, date_to] on that parent's own date
    field (run_date / trial_date). Three separate joins rather than one
    query, since which date field applies depends on which of the three
    parent types the result belongs to."""
    results = []

    run_q = (
        session.query(PhysicalPropertyResult)
        .join(ProductionRun, PhysicalPropertyResult.production_run_id == ProductionRun.id)
        .filter(ProductionRun.recipe_version_id == recipe_version_id)
    )
    if date_from:
        run_q = run_q.filter(ProductionRun.run_date >= date_from)
    if date_to:
        run_q = run_q.filter(ProductionRun.run_date <= date_to)
    results.extend(run_q.all())

    ct_q = (
        session.query(PhysicalPropertyResult)
        .join(CustomerTrial, PhysicalPropertyResult.customer_trial_id == CustomerTrial.id)
        .filter(CustomerTrial.recipe_version_id == recipe_version_id)
    )
    if date_from:
        ct_q = ct_q.filter(CustomerTrial.trial_date >= date_from)
    if date_to:
        ct_q = ct_q.filter(CustomerTrial.trial_date <= date_to)
    results.extend(ct_q.all())

    ot_q = (
        session.query(PhysicalPropertyResult)
        .join(OptimizationTrial, PhysicalPropertyResult.optimization_trial_id == OptimizationTrial.id)
        .filter(OptimizationTrial.recipe_version_id == recipe_version_id)
    )
    if date_from:
        ot_q = ot_q.filter(OptimizationTrial.trial_date >= date_from)
    if date_to:
        ot_q = ot_q.filter(OptimizationTrial.trial_date <= date_to)
    results.extend(ot_q.all())

    return results


def build_recipe_formulation_record_data(session, recipe_version_id, date_from=None, date_to=None):
    rv = session.get(RecipeVersion, recipe_version_id)
    if rv is None:
        return None
    grade = rv.foam_grade
    family = grade.product_family if grade else None

    ordered_components = sorted(
        rv.components,
        key=lambda c: (c.role_in_formulation or "", c.raw_material_name or ""),
    )
    components = [
        {
            "Material": c.raw_material_name,
            "Supplier": c.supplier or "—",
            "PHP": c.php,
            "Role": c.role_in_formulation or "—",
            "Notes": c.notes or "—",
        }
        for c in ordered_components
    ]

    targets_by_name = {t["property_name"]: t for t in _recipe_version_target_properties(grade)}
    results = _property_results_for_recipe_version(session, rv.id, date_from, date_to)
    by_property = {}
    for r in results:
        by_property.setdefault(r.property_name, []).append(r)

    quality_rows = []
    # Show every declared target even if zero results fell in range (an
    # honest "no data yet" beats silently omitting the row), plus any
    # measured property that isn't a formally declared target.
    for prop_name in sorted(set(targets_by_name) | set(by_property)):
        rs = by_property.get(prop_name, [])
        verdicts = [compute_pass_fail(r.property_name, r.target_value, r.actual_value) for r in rs]
        pass_ct, fail_ct = verdicts.count("Pass"), verdicts.count("Fail")
        scored = pass_ct + fail_ct
        actuals = [r.actual_value for r in rs if r.actual_value is not None]
        target = targets_by_name.get(prop_name)
        target_value = target["target_value"] if target else next(
            (r.target_value for r in rs if r.target_value is not None), None
        )
        unit = (target["unit"] if target else None) or next((r.unit for r in rs if r.unit), "")
        quality_rows.append({
            "Property": prop_name,
            "Target": target_value,
            "Avg. actual": round(sum(actuals) / len(actuals), 2) if actuals else None,
            "Unit": unit,
            "Results in range": len(rs),
            "Pass rate": f"{round(100 * pass_ct / scored)}%" if scored else "—",
        })

    cost = recipe_version_cost(session, rv)
    # Same php-parts-as-kg convention as pages/15_Recipe_Optimization.py's
    # _cost_per_kg() - a formulation's php total already IS its cost basis
    # (see analytics.recipe_version_cost's own docstring), so cost per kg
    # is simply total cost / total php, once any component is priced.
    cost_per_kg = (
        round(cost["total_cost"] / cost["total_php"], 2)
        if cost["total_cost"] is not None and cost["total_php"]
        else None
    )

    return {
        "recipe_version_id": rv.id,
        "version_label": rv.version_label,
        "foam_grade": grade.grade_name if grade else "—",
        "product_family": family.name if family else "—",
        "approval_status": rv.approval_status,
        "is_active": rv.is_active,
        "effective_date": rv.effective_date,
        "created_by": rv.created_by or "—",
        "change_note": rv.change_note or "",
        "ratio_index": rv.ratio_index,
        "components": components,
        "date_from": date_from,
        "date_to": date_to,
        "quality_rows": quality_rows,
        "cost_per_kg": cost_per_kg,
        "cost_priced_php": cost["priced_php"],
        "cost_total_php": cost["total_php"],
        "cost_missing_materials": cost["missing"],
    }


def render_recipe_formulation_record_pdf(data):
    def build(story):
        _title_block(
            story, f"Recipe / Formulation Record — {data['version_label']}",
            f"{data['foam_grade']} · {data['product_family']} · "
            f"{'Active recipe' if data['is_active'] else 'Retired version'}",
        )
        story.append(_key_value_table([
            ("Approval status", data["approval_status"]), ("Active", "Yes" if data["is_active"] else "No"),
            ("Effective date", data["effective_date"]), ("Created by", data["created_by"]),
            ("Ratio / index", f"{data['ratio_index']:.3f}" if data["ratio_index"] is not None else "—"),
            ("", ""),
        ]))
        if data["change_note"]:
            story.append(Spacer(1, 6))
            story.append(_p(f"Change note: {data['change_note']}"))
        _section(story, "Formulation (recipe components)", data["components"])

        date_range = f"{data['date_from'] or 'earliest'} to {data['date_to'] or 'latest'}"
        _section(story, f"Quality specs vs. results ({date_range})", data["quality_rows"])

        story.append(Spacer(1, 8))
        story.append(Paragraph("Cost", STYLES["Heading3"]))
        if data["cost_per_kg"] is None:
            story.append(_p("No priced components - cost per kg cannot be calculated."))
        else:
            coverage = f"{data['cost_priced_php']:.2f} / {data['cost_total_php']:.2f} php priced"
            story.append(_key_value_table([
                ("Cost per kg", data["cost_per_kg"]), ("Cost coverage", coverage),
            ]))
            if data["cost_missing_materials"]:
                story.append(_p("Unpriced materials (excluded from total): " + ", ".join(data["cost_missing_materials"])))
    return _pdf_bytes(build)


def render_recipe_formulation_record_docx(data):
    doc = Document()
    _docx_report_header(
        doc, f"Recipe / Formulation Record — {data['version_label']}",
        f"{data['foam_grade']} · {data['product_family']} · "
        f"{'Active recipe' if data['is_active'] else 'Retired version'}",
    )
    _docx_kv_table(doc, [
        ("Approval status", data["approval_status"]), ("Active", "Yes" if data["is_active"] else "No"),
        ("Effective date", data["effective_date"]), ("Created by", data["created_by"]),
        ("Ratio / index", f"{data['ratio_index']:.3f}" if data["ratio_index"] is not None else "—"),
    ])
    if data["change_note"]:
        doc.add_paragraph(f"Change note: {data['change_note']}")
    _docx_section(doc, "Formulation (recipe components)", data["components"])

    date_range = f"{data['date_from'] or 'earliest'} to {data['date_to'] or 'latest'}"
    _docx_section(doc, f"Quality specs vs. results ({date_range})", data["quality_rows"])

    _docx_heading(doc, "Cost", size=12, color=_HTC_GREY, space_before=10)
    if data["cost_per_kg"] is None:
        doc.add_paragraph("No priced components - cost per kg cannot be calculated.")
    else:
        coverage = f"{data['cost_priced_php']:.2f} / {data['cost_total_php']:.2f} php priced"
        _docx_kv_table(doc, [("Cost per kg", data["cost_per_kg"]), ("Cost coverage", coverage)])
        if data["cost_missing_materials"]:
            doc.add_paragraph(
                "Unpriced materials (excluded from total): " + ", ".join(data["cost_missing_materials"])
            )
    return _docx_bytes(doc)


# ---------------------------------------------------------------------------
# 4. Where Used Report
#
# Given a raw material, answers "which recipes use this, and what depends
# on it" - the reverse lookup a Plant Manager needs before considering a
# material substitution. Scoped inherently by the tenant boundary: it
# joins on RecipeComponent.raw_material_id, a real FK to one specific
# (already company-scoped) raw_materials row, not a name match, so it
# can't cross into another company's recipes.
# ---------------------------------------------------------------------------

def build_where_used_report_data(session, raw_material_id):
    rm = session.get(RawMaterial, raw_material_id)
    if rm is None:
        return None

    components = (
        session.query(RecipeComponent)
        .filter(RecipeComponent.raw_material_id == rm.id)
        .all()
    )
    recipe_version_ids = {c.recipe_version_id for c in components}
    versions = (
        session.query(RecipeVersion).filter(RecipeVersion.id.in_(recipe_version_ids)).all()
        if recipe_version_ids else []
    )
    version_by_id = {v.id: v for v in versions}

    def _sort_key(c):
        v = version_by_id.get(c.recipe_version_id)
        grade = v.foam_grade if v else None
        return (grade.grade_name if grade else "", v.version_label if v else "")

    usage_rows = []
    grade_ids, family_names = set(), set()
    for c in sorted(components, key=_sort_key):
        v = version_by_id.get(c.recipe_version_id)
        grade = v.foam_grade if v else None
        family = grade.product_family if grade else None
        if grade:
            grade_ids.add(grade.id)
        if family:
            family_names.add(family.name)
        usage_rows.append({
            "Product grade": grade.grade_name if grade else "—",
            "Product family": family.name if family else "—",
            "Recipe version": v.version_label if v else "—",
            "Status": "Active" if v and v.is_active else "Retired",
            "PHP": c.php,
            "Role": c.role_in_formulation or "—",
            "Approval status": v.approval_status if v else "—",
        })

    grades = session.query(FoamGrade).filter(FoamGrade.id.in_(grade_ids)).all() if grade_ids else []
    target_rows = []
    for g in sorted(grades, key=lambda g: g.grade_name):
        for t in _recipe_version_target_properties(g):
            target_rows.append({
                "Product grade": g.grade_name, "Property": t["property_name"],
                "Target": t["target_value"], "Unit": t["unit"],
            })

    trial_rows = []
    if recipe_version_ids:
        customer_trials = (
            session.query(CustomerTrial)
            .filter(CustomerTrial.recipe_version_id.in_(recipe_version_ids))
            .order_by(CustomerTrial.trial_date.desc())
            .all()
        )
        for t in customer_trials:
            trial_rows.append({
                "Trial type": "Customer Trial", "Trial ID": t.id,
                "Product grade": t.foam_grade.grade_name if t.foam_grade else "—",
                "Status": t.status, "Trial date": t.trial_date, "Outcome": t.outcome or "—",
            })
        optimization_trials = (
            session.query(OptimizationTrial)
            .filter(OptimizationTrial.recipe_version_id.in_(recipe_version_ids))
            .order_by(OptimizationTrial.trial_date.desc())
            .all()
        )
        for t in optimization_trials:
            trial_rows.append({
                "Trial type": "Optimization Trial", "Trial ID": t.id,
                "Product grade": t.foam_grade.grade_name if t.foam_grade else "—",
                "Status": t.status, "Trial date": t.trial_date, "Outcome": t.conclusion or "—",
            })

    return {
        "raw_material_id": rm.id,
        "raw_material_name": rm.name,
        "category": rm.category or "—",
        "default_supplier": rm.default_supplier or "—",
        "active": rm.active,
        "recipe_version_count": len(recipe_version_ids),
        "foam_grade_count": len(grade_ids),
        "product_family_count": len(family_names),
        "usage_rows": usage_rows,
        "target_rows": target_rows,
        "trial_rows": trial_rows,
    }


def render_where_used_report_pdf(data):
    def build(story):
        _title_block(
            story, f"Where Used Report — {data['raw_material_name']}",
            f"{data['category']} · Default supplier: {data['default_supplier']} · "
            f"{'Active' if data['active'] else 'Inactive'} material",
        )
        story.append(_key_value_table([
            ("Recipe versions using this material", data["recipe_version_count"]),
            ("Product grades affected", data["foam_grade_count"]),
            ("Product families affected", data["product_family_count"]),
            ("", ""),
        ]))
        _section(story, "Recipes using this material", data["usage_rows"])
        _section(story, "Target properties of affected product grades", data["target_rows"])
        _section(story, "Trial precedent (Customer / Optimization Trials on these recipes)", data["trial_rows"])
    return _pdf_bytes(build)


def render_where_used_report_docx(data):
    doc = Document()
    _docx_report_header(
        doc, f"Where Used Report — {data['raw_material_name']}",
        f"{data['category']} · Default supplier: {data['default_supplier']} · "
        f"{'Active' if data['active'] else 'Inactive'} material",
    )
    _docx_kv_table(doc, [
        ("Recipe versions using this material", data["recipe_version_count"]),
        ("Product grades affected", data["foam_grade_count"]),
        ("Product families affected", data["product_family_count"]),
    ])
    _docx_section(doc, "Recipes using this material", data["usage_rows"])
    _docx_section(doc, "Target properties of affected product grades", data["target_rows"])
    _docx_section(doc, "Trial precedent (Customer / Optimization Trials on these recipes)", data["trial_rows"])
    return _docx_bytes(doc)


# ---------------------------------------------------------------------------
# 5. Batch Release / Conformance Record (Production Run)
#
# Purpose (per user direction 2026-08-04): "did this batch meet spec, and
# is there anything on record I should know" - not a transcription of
# every stored field on the run. A clean run gets a short document,
# header + recipe + verdict. A FLAGGED run (a failed quality result, or
# any recorded quality issue) widens to pull relevant context from every
# other tab on the Production Run page - actual component stream
# readings and Production Events - because a flag's explanation often
# isn't on the same tab as the flag itself. (Setup-vs-Finalized
# process-setting/fall-plate-position deviations were part of this
# widened context before WP7 Phase 0, 2026-08-13, removed the active
# fall-plate section-position sub-workflow and WP7 Phase 5, 2026-08-15,
# retired ProductionPhase's active machine-setting authority; the two
# comparison-tolerance constants that backed that removed section
# (_SETTING_DEVIATION_EPSILON, _FALLPLATE_POSITION_DEVIATION_MM) were
# removed alongside it as dead code, per Decision Ledger D5-02/D5-03.)
# ---------------------------------------------------------------------------

_PARAMETER_CATEGORIES_FOR_REPORT = ("Process Setting", "Environment", "Outcome")


def _effective_limit(min_base, max_base, min_override, max_override):
    """Winning acceptance limit for one process-parameter report row - the
    applicability-level override (min_value_override/max_value_override)
    takes precedence over the definition's own default min_value/max_value
    when populated, per Charlie's WP7 Phase 4 Closeout Review Return to JC
    (Material Completion Item 1.2): "use min_value_override /
    max_value_override when populated, otherwise the definition min_value
    / max_value". Returns (effective_min, effective_max) - either side may
    be None (no bound on that side); both None means no approved limit at
    all, i.e. this parameter stays informational rather than Pass/Fail."""
    eff_min = min_override if min_override is not None else min_base
    eff_max = max_override if max_override is not None else max_base
    return eff_min, eff_max


def _limit_text(eff_min, eff_max):
    if eff_min is not None and eff_max is not None:
        return f"{eff_min}–{eff_max}"
    if eff_min is not None:
        return f">= {eff_min}"
    if eff_max is not None:
        return f"<= {eff_max}"
    return "—"


def _conformance_text(eff_min, eff_max, actual_value, data_type):
    """Per Charlie's Item 1.2: "produce a conformance result only when an
    approved limit exists" - a row with no approved limit remains
    informational, never silently shown as Pass. Only Float/Integer
    Actual values are evaluated numerically against the limit; a limit
    recorded against a non-numeric definition (not expected under today's
    controlled vocabulary, but the schema doesn't forbid it) also stays
    informational since there is nothing numeric to compare."""
    if eff_min is None and eff_max is None:
        return "Informational (no approved limit)"
    if data_type not in ("Float", "Integer"):
        return "Informational (no approved limit)"
    if actual_value is None:
        return "No Actual value recorded"
    if eff_min is not None and actual_value < eff_min:
        return "Fail"
    if eff_max is not None and actual_value > eff_max:
        return "Fail"
    return "Pass"


def _process_parameter_report_rows(session, run_id):
    """WP7 Phase 4 targeted-completion correction (2026-08-14, per
    Charlie's WP7 Phase 4 Closeout Review Return to JC, Material
    Completion Item 1). Replaces the retired _process_parameter_
    deviations(), which only surfaced rows where Planned and Actual
    actually differed and carried no Category/UOM/limit columns - flagged
    by Charlie's review as not meeting the required report shape: "The
    process-data section must be definition-driven and show Parameter,
    Category, Planned, Actual, numeric Delta where applicable, and
    canonical UOM. Environment and Outcome records must be separated from
    controllable Process Settings" plus "Controlled acceptance limits must
    be applied where they exist... a row with no approved limit remains
    informational."

    Reads analytics.production_run_process_parameters(session, run_id) -
    unchanged shared reader, still Machine > Method > Global precedence,
    still Actual-never-falls-back-to-Planned (Charlie's accepted Phase 4
    shared-reader behavior is not reopened here) - and returns EVERY
    eligible definition's row (definition-driven, not a deviations-only
    filter), bucketed by parameter_category into three separate lists so
    Environment/Outcome observations are structurally kept out of the
    controllable Process Setting table rather than merely sorted next to
    it. A category with no eligible definitions for this run's Production
    Method/Machine contributes an empty list (rendered as "No data
    recorded" by _section/_docx_section) - the honest empty state, same
    "never a legacy fallback" rule as the shared reader itself.

    Returns {"Process Setting": [...], "Environment": [...], "Outcome":
    [...]}. Each row: Parameter, Category, Planned, Actual, Delta (the
    shared reader's own Float/Integer-only delta, None otherwise -
    rendered as "—" by _section/_docx_section), UOM, Limit (display
    text for the winning min/max, see _limit_text), Conformance (Pass/
    Fail only when an approved limit exists and Actual is recorded and
    numeric; otherwise explicitly informational, see _conformance_text -
    never a silent Pass)."""
    buckets = {cat: [] for cat in _PARAMETER_CATEGORIES_FOR_REPORT}
    for row in production_run_process_parameters(session, run_id):
        category = row["parameter_category"]
        if category not in buckets:
            continue  # controlled vocabulary; no other category is expected today
        eff_min, eff_max = _effective_limit(
            row.get("min_value"), row.get("max_value"),
            row.get("min_value_override"), row.get("max_value_override"),
        )
        buckets[category].append({
            "Parameter": row["name"] or row["controlled_id"],
            "Category": category,
            "Planned": row["planned_value"],
            "Actual": row["actual_value"],
            "Delta": row["delta"],
            "UOM": row["unit_symbol"] or "—",
            "Limit": _limit_text(eff_min, eff_max),
            "Conformance": _conformance_text(eff_min, eff_max, row["actual_value"], row["data_type"]),
        })
    return buckets


# _fallplate_deviations (fall-plate section-position changes between Setup
# and Finalized) was removed under WP7 Phase 0, 2026-08-13, along with the
# rest of the active "Tool Geometry and Fill Configuration" sub-workflow -
# see pages/4_Production_Run_Trial_Record.py's module docstring.
# FallplateSectionPosition rows already recorded remain in the database and
# readable directly off that model; this report section simply no longer
# runs. The Batch Release report's "story"/"_docx_section" call sites and
# the "fallplate_deviations" data-dict key were removed alongside it.


def _is_rigid_grade(grade):
    """A grade is 'rigid' (has a controlled chemistry, hence real
    grade_specifications rows to evaluate against) vs. a legacy
    flexible-style grade with no chemistry - same is_rigid convention
    already used by pages/15_Recipe_Optimization.py's WP4 branch."""
    return bool(grade and grade.chemistry_id is not None)


def _conformance_verdict(conformance_rows):
    """Overall verdict text from wp3_conformance.compute_conformance_report()'s
    raw rows. "No specification on file" when the grade has zero
    grade_specifications rows at all (nothing to report against - distinct
    from "Incomplete testing", which means specs exist but couldn't all be
    evaluated)."""
    if not conformance_rows:
        return "No specification on file"
    verdicts = [r["verdict"] for r in conformance_rows]
    if "Fail" in verdicts:
        return "Non-conforming"
    if verdicts and all(v == "Pass" for v in verdicts):
        return "Conforming"
    return "Incomplete testing"


def _conformance_rows_for_display(session, results_lookup, conformance_rows):
    """Reshapes wp3_conformance.compute_conformance_report() rows into the
    flat Property/Specification/Actual/Unit/Pass-Fail/Test method/Condition/
    Spec reference/Tested columns the Batch Release Record and Sample
    Certificate tables render - joining back to the matched
    PhysicalPropertyResult (results_lookup, keyed by id, already loaded by
    the caller) for the test-method/condition/tested-at context
    compute_conformance_report's own return rows don't carry.

    WP6-S09 (Charlie's UAT-011/UAT-014 technical review, 2026-08-09): this
    replaces resolving Target/Pass-Fail from a target_value stored directly
    on the result row (quality_standards.compute_pass_fail, still used
    below for legacy/flexible-style grades) with resolution against the
    grade's own controlled grade_specifications - the traceable conformance
    basis Charlie's review requires. "Spec reference" cites the
    grade_specifications row's own id (GS-<id>) - there is no separate
    revision field on this table; the row id plus its Source Register link
    (via GradeSpecification.source), where present, is the traceable
    reference this schema actually carries.
    """
    out = []
    for row in conformance_rows:
        result = results_lookup.get(row["result_id"]) if row.get("result_id") else None
        spec = session.get(GradeSpecification, row["spec_id"]) if row.get("spec_id") else None

        if spec is None:
            target_text = "—"
        else:
            op = (spec.target_operator or "<=").strip()
            if op == "between":
                target_text = f"{spec.lower_limit}–{spec.upper_limit} {spec.unit or ''}".strip()
            else:
                limit = spec.target_value
                if limit is None:
                    limit = spec.upper_limit if op == "<=" else (spec.lower_limit if op == ">=" else None)
                target_text = f"{op} {limit} {spec.unit or ''}".strip() if limit is not None else "—"

        status = row.get("status")
        if status in ("Pass", "Fail"):
            pass_fail_text = status
        else:
            pass_fail_text = {
                "EXCLUDED_CONTEXT": f"Excluded ({row.get('excluded_reason') or 'context mismatch'})",
                "INVALID": f"Invalid ({row.get('excluded_reason') or 'incomplete context'})",
                "NO_RESULT": "No result recorded",
            }.get(status, status or "—")

        out.append({
            "Property": row.get("property_name"),
            "Specification": target_text,
            "Actual": row.get("actual_value"),
            "Unit": (spec.unit if spec else None) or (result.unit if result else "") or "",
            "Pass/Fail": pass_fail_text,
            "Test method": (result.test_method if result else None) or "—",
            "Condition": (
                result.condition.name if (result is not None and result.condition)
                else (spec.condition.name if (spec is not None and spec.condition) else "—")
            ),
            "Spec reference": f"GS-{spec.id}" if spec is not None else "—",
            "Tested": result.tested_at if result else None,
        })
    return out


def build_batch_release_record_data(session, run_id):
    run = session.get(ProductionRun, run_id)
    if run is None:
        return None
    grade = run.foam_grade
    family = grade.product_family if grade else None
    recipe = run.recipe_version

    ordered_components = (
        sorted(recipe.components, key=lambda c: (c.role_in_formulation or "", c.raw_material_name or ""))
        if recipe else []
    )
    recipe_components = [
        {
            "Material": c.raw_material_name, "Supplier": c.supplier or "—",
            "PHP": c.php, "Role": c.role_in_formulation or "—", "Notes": c.notes or "—",
        }
        for c in ordered_components
    ]

    results = (
        session.query(PhysicalPropertyResult)
        .filter(PhysicalPropertyResult.production_run_id == run_id).all()
    )

    if _is_rigid_grade(grade):
        conformance_rows = wp3_conformance.compute_conformance_report(
            session, grade.id, production_run_id=run_id
        )
        results_by_id = {r.id: r for r in results}
        quality_results = _conformance_rows_for_display(session, results_by_id, conformance_rows)
        quality_verdict = _conformance_verdict(conformance_rows)
        verdicts = [r["verdict"] for r in conformance_rows]  # None entries for Excluded/Invalid/No-result rows
    else:
        quality_results = [
            {
                "Property": r.property_name, "Target": r.target_value, "Actual": r.actual_value,
                "Unit": r.unit or "",
                "Pass/Fail": compute_pass_fail(r.property_name, r.target_value, r.actual_value) or "—",
                "Tested": r.tested_at,
            }
            for r in results
        ]
        verdicts = [compute_pass_fail(r.property_name, r.target_value, r.actual_value) for r in results]
        if not results:
            quality_verdict = "No testing recorded"
        elif "Fail" in verdicts:
            quality_verdict = "Non-conforming"
        elif verdicts and all(v == "Pass" for v in verdicts):
            quality_verdict = "Conforming"
        else:
            quality_verdict = "Incomplete testing"

    quality_issues = [
        {
            "Issue type": o.observation_type, "Severity": o.severity or "—",
            "Frequency": o.frequency or "—", "Confidence": o.confidence_level or "—",
            "Suspected cause": o.suspected_cause or "—",
        }
        for o in session.query(QualityObservation).filter(QualityObservation.production_run_id == run_id).all()
    ]

    has_flags = ("Fail" in verdicts) or bool(quality_issues)
    flag_reasons = []
    fail_count = verdicts.count("Fail")
    if fail_count:
        flag_reasons.append(f"{fail_count} failed quality result(s)")
    if quality_issues:
        flag_reasons.append(f"{len(quality_issues)} quality issue(s) recorded")

    # WP7 Phase 4 cutover (2026-08-14): ProductionOutputSummary is the
    # active output fact for this report - see Charlie's Downstream
    # Reader Cutover Execution Instruction section 6. Shown unconditionally
    # (not gated by has_flags), since disposition is a core release
    # decision this "Conformance Record" exists to surface, not
    # supplementary flag context. None (no recorded row) stays None -
    # never inferred from the retired compute_runtime_output() geometry
    # formula.
    output_summary = production_run_output_summary(session, run_id)

    process_setting_rows, environment_rows, outcome_rows = [], [], []
    stream_readings, stream_calibration_flags, production_events = [], [], []
    if has_flags:
        rows_by_category = _process_parameter_report_rows(session, run_id)
        process_setting_rows = rows_by_category["Process Setting"]
        environment_rows = rows_by_category["Environment"]
        outcome_rows = rows_by_category["Outcome"]

        # WP7 Phase 4 targeted-completion correction (2026-08-14, Charlie's
        # Closeout Review Return to JC, Material Completion Item 1.3):
        # Material Metering now reads exclusively via production_run_id -
        # the active Phase 2 run anchor - never via a located Finalized
        # ProductionPhase. ComponentStreamReading.production_phase_id is
        # nullable and Phase-5-scoped for eventual retirement; a run-linked
        # reading with production_phase_id left unset still surfaces here.
        # This removes Batch Release's last direct ProductionPhase read.
        readings = (
            session.query(ComponentStreamReading)
            .filter(ComponentStreamReading.production_run_id == run_id).all()
        )
        stream_readings = [
            {
                "Stream": rd.stream_name, "Flow": rd.flow, "Unit": rd.flow_unit or "",
                "Pump speed": rd.pump_speed, "Total delivered": rd.flow_total_qty,
                "Temperature (°C)": rd.temperature_c, "Pressure (bar)": rd.pressure_bar,
                "Calibration": rd.calibration_status or "—",
            }
            for rd in readings
        ]
        stream_calibration_flags = [
            rd.stream_name for rd in readings if rd.calibration_status and rd.calibration_status != "Valid"
        ]

        production_events = [
            {
                "Time": e.event_ts, "Type": e.event_type, "Severity": e.severity or "—",
                "Description": e.description or "—",
            }
            for e in session.query(ProductionEvent)
            .filter(ProductionEvent.production_run_id == run_id)
            .order_by(ProductionEvent.event_ts).all()
        ]

    return {
        "run_id": run.id,
        "plant": run.plant.name if run.plant else "—",
        "product_family": family.name if family else "—",
        "foam_grade": grade.grade_name if grade else "—",
        "machine": run.machine.name if run.machine else "—",
        # Immutable snapshot taken at run creation (db.py ProductionRun.
        # production_method_id) - per Charlie's flat-PM technical completion
        # instruction ("run-specific reports use the stored Production
        # Method snapshot"), never the machine's or grade's CURRENT method,
        # so a re-tagged machine never rewrites history for a past run.
        "production_method": run.production_method.name if run.production_method else "—",
        "run_date": run.run_date,
        "batch_reference": run.batch_reference or "—",
        "block_reference": run.block_reference or "—",
        # CR-22 / F22-04 (AF22-01): Block reference is customer-facing
        # only for PM-500 Rigid Block Production - the renderers below
        # omit the "Block reference" row entirely for every other method.
        "block_reference_applicable": _block_reference_applicable(run.production_method),
        "operator": run.operator_or_team_reference or "—",
        "notes": run.notes or "",
        "recipe_version_label": recipe.version_label if recipe else "—",
        "recipe_approval_status": recipe.approval_status if recipe else "—",
        "recipe_effective_date": recipe.effective_date if recipe else None,
        "recipe_ratio_index": recipe.ratio_index if recipe else None,
        "recipe_components": recipe_components,
        "quality_results": quality_results,
        "quality_verdict": quality_verdict,
        "quality_issues": quality_issues,
        "output_summary": output_summary,
        "has_flags": has_flags,
        "flag_reasons": flag_reasons,
        "process_setting_rows": process_setting_rows,
        "environment_rows": environment_rows,
        "outcome_rows": outcome_rows,
        "stream_readings": stream_readings,
        "stream_calibration_flags": stream_calibration_flags,
        "production_events": production_events,
    }


def render_batch_release_record_pdf(data):
    def build(story):
        _title_block(
            story, f"Batch Release Record — Run #{data['run_id']}",
            f"{data['plant']} · {data['foam_grade']} · {data['run_date'] or '—'} · "
            f"Verdict: {data['quality_verdict']}",
        )
        # CR-22 / F22-04 (AF22-01): "Block reference" row omitted entirely
        # for every method except PM-500 Rigid Block Production.
        batch_release_kv = [
            ("Plant", data["plant"]), ("Product family", data["product_family"]),
            ("Product grade", data["foam_grade"]), ("Production Unit or Cell", data["machine"]),
            ("Run date", data["run_date"]), ("Batch reference", data["batch_reference"]),
        ]
        if data["block_reference_applicable"]:
            batch_release_kv.append(("Block reference", data["block_reference"]))
        batch_release_kv += [("Operator/team", data["operator"]), ("Quality verdict", data["quality_verdict"]), ("", "")]
        story.append(_key_value_table(batch_release_kv))
        if data["notes"]:
            story.append(Spacer(1, 6))
            story.append(_p(f"Notes: {data['notes']}"))

        # WP7 Phase 4 cutover (2026-08-14): ProductionOutputSummary is the
        # active output fact for this report - see Charlie's Downstream
        # Reader Cutover Execution Instruction section 6. None means no row
        # recorded yet, shown honestly rather than inferred.
        story.append(Spacer(1, 8))
        story.append(Paragraph("Production output", STYLES["Heading3"]))
        out = data["output_summary"]
        if out is None:
            story.append(_p("No Production Output has been recorded yet for this run."))
        else:
            unit = out["unit_symbol"] or ""
            story.append(_key_value_table([
                ("Planned quantity", f"{out['planned_quantity']} {unit}".strip() if out["planned_quantity"] is not None else "—"),
                ("Actual quantity", f"{out['actual_quantity']} {unit}".strip() if out["actual_quantity"] is not None else "—"),
                ("Disposition", out["disposition"] or "—"),
                ("Disposition notes", out["disposition_notes"] or "—"),
            ]))

        story.append(Spacer(1, 8))
        story.append(Paragraph("Recipe used", STYLES["Heading3"]))
        story.append(_key_value_table([
            ("Recipe version", data["recipe_version_label"]), ("Approval status", data["recipe_approval_status"]),
            ("Effective date", data["recipe_effective_date"]),
            ("Ratio / index", f"{data['recipe_ratio_index']:.3f}" if data["recipe_ratio_index"] is not None else "—"),
        ]))
        _section(story, "Formulation", data["recipe_components"])

        _section(story, "Quality test results", data["quality_results"])
        _section(story, "Quality issues", data["quality_issues"])

        if data["has_flags"]:
            story.append(Spacer(1, 10))
            story.append(Paragraph("Flagged — supporting context from other tabs", STYLES["Heading2"]))
            story.append(_p("Flagged because: " + "; ".join(data["flag_reasons"])))
            _section(story, "Process settings — Planned vs Actual (method-aware, with controlled limits)", data["process_setting_rows"])
            _section(story, "Environment — recorded observations", data["environment_rows"])
            _section(story, "Outcome — recorded observations", data["outcome_rows"])
            _section(story, "Material metering and actual usage (Actual Run and Cycle Data phase)", data["stream_readings"])
            if data["stream_calibration_flags"]:
                story.append(_p(
                    "⚠ Non-valid calibration status recorded for: " + ", ".join(data["stream_calibration_flags"])
                ))
            _section(story, "Production events during this run", data["production_events"])
    return _pdf_bytes(build)


def render_batch_release_record_docx(data):
    doc = Document()
    _docx_report_header(
        doc, f"Batch Release Record — Run #{data['run_id']}",
        f"{data['plant']} · {data['foam_grade']} · {data['run_date'] or '—'} · "
        f"Verdict: {data['quality_verdict']}",
    )
    # CR-22 / F22-04 (AF22-01): "Block reference" row omitted entirely for
    # every method except PM-500 Rigid Block Production.
    batch_release_docx_kv = [
        ("Plant", data["plant"]), ("Product family", data["product_family"]),
        ("Product grade", data["foam_grade"]), ("Production Unit or Cell", data["machine"]),
        ("Production method", data["production_method"]),
        ("Run date", data["run_date"]), ("Batch reference", data["batch_reference"]),
    ]
    if data["block_reference_applicable"]:
        batch_release_docx_kv.append(("Block reference", data["block_reference"]))
    batch_release_docx_kv += [("Operator/team", data["operator"]), ("Quality verdict", data["quality_verdict"])]
    _docx_kv_table(doc, batch_release_docx_kv)
    if data["notes"]:
        doc.add_paragraph(f"Notes: {data['notes']}")

    # WP7 Phase 4 cutover (2026-08-14): ProductionOutputSummary is the
    # active output fact for this report - see Charlie's Downstream Reader
    # Cutover Execution Instruction section 6. None means no row recorded
    # yet, shown honestly rather than inferred.
    _docx_heading(doc, "Production output", size=12, color=_HTC_GREY, space_before=10)
    out = data["output_summary"]
    if out is None:
        doc.add_paragraph("No Production Output has been recorded yet for this run.")
    else:
        unit = out["unit_symbol"] or ""
        _docx_kv_table(doc, [
            ("Planned quantity", f"{out['planned_quantity']} {unit}".strip() if out["planned_quantity"] is not None else "—"),
            ("Actual quantity", f"{out['actual_quantity']} {unit}".strip() if out["actual_quantity"] is not None else "—"),
            ("Disposition", out["disposition"] or "—"),
            ("Disposition notes", out["disposition_notes"] or "—"),
        ])

    _docx_heading(doc, "Recipe used", size=12, color=_HTC_GREY, space_before=10)
    _docx_kv_table(doc, [
        ("Recipe version", data["recipe_version_label"]), ("Approval status", data["recipe_approval_status"]),
        ("Effective date", data["recipe_effective_date"]),
        ("Ratio / index", f"{data['recipe_ratio_index']:.3f}" if data["recipe_ratio_index"] is not None else "—"),
    ])
    _docx_section(doc, "Formulation", data["recipe_components"])
    _docx_section(doc, "Quality test results", data["quality_results"])
    _docx_section(doc, "Quality issues", data["quality_issues"])

    if data["has_flags"]:
        _docx_heading(doc, "Flagged — supporting context from other tabs", size=15, space_before=14)
        doc.add_paragraph("Flagged because: " + "; ".join(data["flag_reasons"]))
        _docx_section(doc, "Process settings — Planned vs Actual (method-aware, with controlled limits)", data["process_setting_rows"])
        _docx_section(doc, "Environment — recorded observations", data["environment_rows"])
        _docx_section(doc, "Outcome — recorded observations", data["outcome_rows"])
        _docx_section(doc, "Material metering and actual usage (Actual Run and Cycle Data phase)", data["stream_readings"])
        if data["stream_calibration_flags"]:
            doc.add_paragraph(
                "⚠ Non-valid calibration status recorded for: " + ", ".join(data["stream_calibration_flags"])
            )
        _docx_section(doc, "Production events during this run", data["production_events"])
    return _docx_bytes(doc)


# ---------------------------------------------------------------------------
# Sample Certificate of Analysis
#
# Placement/mechanism (per user direction 2026-08-04): picking one sample
# is a single simple choice - same as Batch Release Record and Trial
# Closeout Report - so this lives on the Report page (pages/21_Report.py)
# rather than on pages 9/11/12 themselves (which already have the
# aggregate, multi-field-selection Sample Report).
#
# "What is this sample, what recipe made it, what did testing find, and
# did it pass" in one document - works across all three sample sources
# (Production Run / Customer Trial / Optimization Trial - see
# db.SAMPLE_SOURCE_TYPES). Includes the full recipe formulation (materials/
# php/supplier/role), same as Batch Release Record and Recipe Formulation
# Record - per that same precedent, this makes the report NOT customer-
# facing as-is; a customer-facing version would need to omit section C
# (the formulation table).
# ---------------------------------------------------------------------------

def _sample_source(session, sample):
    """(source_type, source object) for a Sample, resolving whichever of
    the three mutually exclusive parent FKs is set - see
    db.SAMPLE_SOURCE_TYPES / db.sample_source_fk_field()."""
    if sample.production_run_id is not None:
        return "Production Run", session.get(ProductionRun, sample.production_run_id)
    if sample.customer_trial_id is not None:
        return "Customer Trial", session.get(CustomerTrial, sample.customer_trial_id)
    if sample.optimization_trial_id is not None:
        return "Optimization Trial", session.get(OptimizationTrial, sample.optimization_trial_id)
    return "—", None


def build_sample_certificate_data(session, sample_id):
    sample = session.get(Sample, sample_id)
    if sample is None:
        return None
    source_type, source = _sample_source(session, sample)
    grade = source.foam_grade if source else None
    plant = source.plant if source else None
    recipe = source.recipe_version if source else None

    # Production Method (added 2026-08-10, per Charlie's flat-PM technical
    # completion instruction): only the Production Run source has one -
    # inherited from that run's own immutable snapshot, never the machine's
    # CURRENT method - the two lab-trial sources are explicitly "N/A (lab
    # trial)", same convention as pages 5/6/9's production_method_label().
    if source_type == "Production Run":
        header_fields = [
            ("Source", f"Production Run #{source.id}"),
            ("Run date", source.run_date), ("Batch reference", source.batch_reference or "—"),
        ]
        # CR-22 / F22-04 (AF22-01): Block reference row omitted entirely
        # for every method except PM-500 Rigid Block Production.
        if _block_reference_applicable(source.production_method):
            header_fields.append(("Block reference", source.block_reference or "—"))
        header_fields += [
            ("Production Unit or Cell", source.machine.name if source.machine else "—"),
            ("Production Method", source.production_method.name if source.production_method else "—"),
            ("Operator/team", source.operator_or_team_reference or "—"),
        ]
    elif source_type == "Customer Trial":
        header_fields = [
            ("Source", f"Customer Trial #{source.id}"),
            ("Customer", source.customer_name), ("Trial date", source.trial_date),
            ("Status", source.status), ("Responsible", source.responsible_person or "—"),
            ("Batch reference", source.batch_reference or "—"),
            ("Production Method", "N/A (lab trial)"),
        ]
    elif source_type == "Optimization Trial":
        header_fields = [
            ("Source", f"Optimization Trial #{source.id}"),
            ("Improvement initiative reference", source.improvement_initiative_reference or "—"),
            ("Trial date", source.trial_date), ("Status", source.status),
            ("Responsible", source.responsible_person or "—"),
            ("Batch reference", source.batch_reference or "—"),
            ("Production Method", "N/A (lab trial)"),
        ]
    else:
        header_fields = [("Source", "—")]

    ordered_components = (
        sorted(recipe.components, key=lambda c: (c.role_in_formulation or "", c.raw_material_name or ""))
        if recipe else []
    )
    recipe_components = [
        {
            "Material": c.raw_material_name, "Supplier": c.supplier or "—",
            "PHP": c.php, "Role": c.role_in_formulation or "—", "Notes": c.notes or "—",
        }
        for c in ordered_components
    ]

    results = (
        session.query(PhysicalPropertyResult)
        .filter(PhysicalPropertyResult.sample_id == sample.id).all()
    )

    # WP6-S09 fix (2026-08-09, UAT-014 per Charlie's review): same
    # grade-specification resolution as UAT-011/UAT-012 (see
    # _is_rigid_grade / _conformance_rows_for_display / _conformance_verdict
    # above build_batch_release_record_data) rather than the flat
    # compute_pass_fail(target_value) path, which evaluates to "Not
    # computed" for every real rigid result (target_value is always NULL -
    # see wp3_conformance.evaluate_specification's docstring). sample_id is
    # passed to compute_conformance_report so conformance is scoped to just
    # this sample's own results, not every sample under the same run/trial.
    if _is_rigid_grade(grade):
        source_kwargs = {}
        if source_type == "Production Run":
            source_kwargs["production_run_id"] = source.id
        elif source_type == "Customer Trial":
            source_kwargs["customer_trial_id"] = source.id
        elif source_type == "Optimization Trial":
            source_kwargs["optimization_trial_id"] = source.id
        conformance_rows = (
            wp3_conformance.compute_conformance_report(session, grade.id, sample_id=sample.id, **source_kwargs)
            if source_kwargs else []
        )
        results_by_id = {r.id: r for r in results}
        quality_results = _conformance_rows_for_display(session, results_by_id, conformance_rows)
        overall_verdict = _conformance_verdict(conformance_rows)
        verdicts = [r["verdict"] for r in conformance_rows]  # None entries for Excluded/Invalid/No-result rows
        pass_count = verdicts.count("Pass")
        fail_count = verdicts.count("Fail")
    else:
        quality_results = [
            {
                "Property": r.property_name, "Target": r.target_value, "Actual": r.actual_value,
                "Unit": r.unit or "",
                "Pass/Fail": compute_pass_fail(r.property_name, r.target_value, r.actual_value) or "Not computed",
                "Method": r.test_method or "—", "Rev.": r.method_revision or "—",
                "Replicate": r.replicate_no, "Tested": r.tested_at,
            }
            for r in sorted(results, key=lambda r: r.property_name)
        ]
        verdicts = [compute_pass_fail(r.property_name, r.target_value, r.actual_value) for r in results]
        pass_count = verdicts.count("Pass")
        fail_count = verdicts.count("Fail")
        if not results:
            overall_verdict = "No testing recorded"
        elif fail_count:
            overall_verdict = "Non-conforming"
        elif verdicts and all(v == "Pass" for v in verdicts):
            overall_verdict = "Conforming"
        else:
            overall_verdict = "Incomplete testing"

    return {
        "sample_id": sample.id,
        "source_type": source_type,
        "source_id": source.id if source else None,
        "header_fields": header_fields,
        "foam_grade": grade.grade_name if grade else "—",
        "plant": plant.name if plant else "—",
        "zone_label": sample.zone_label or "—",
        "sample_ts": sample.sample_ts,
        "sample_notes": sample.notes or "",
        "recipe_version_label": recipe.version_label if recipe else "—",
        "recipe_approval_status": recipe.approval_status if recipe else "—",
        "recipe_effective_date": recipe.effective_date if recipe else None,
        "recipe_ratio_index": recipe.ratio_index if recipe else None,
        "recipe_components": recipe_components,
        "quality_results": quality_results,
        "pass_count": pass_count,
        "fail_count": fail_count,
        "overall_verdict": overall_verdict,
    }


def render_sample_certificate_pdf(data):
    def build(story):
        _title_block(
            story, f"Sample Certificate of Analysis — Sample #{data['sample_id']}",
            f"{data['plant']} · {data['foam_grade']} · Sample Location Reference: {data['zone_label']} · "
            f"Verdict: {data['overall_verdict']}",
        )
        story.append(Paragraph("Sample source", STYLES["Heading3"]))
        story.append(_key_value_table(data["header_fields"] + [("Product grade", data["foam_grade"]), ("Plant", data["plant"])]))

        story.append(Spacer(1, 8))
        story.append(Paragraph("Sample", STYLES["Heading3"]))
        story.append(_key_value_table([
            ("Sample ID", data["sample_id"]), ("Sample Location Reference", data["zone_label"]),
            ("Sampled", data["sample_ts"]), ("", ""),
        ]))
        if data["sample_notes"]:
            story.append(_p(f"Notes: {data['sample_notes']}"))

        story.append(Spacer(1, 8))
        story.append(Paragraph("Recipe used", STYLES["Heading3"]))
        story.append(_key_value_table([
            ("Recipe version", data["recipe_version_label"]), ("Approval status", data["recipe_approval_status"]),
            ("Effective date", data["recipe_effective_date"]),
            ("Ratio / index", f"{data['recipe_ratio_index']:.3f}" if data["recipe_ratio_index"] is not None else "—"),
        ]))
        _section(story, "Formulation", data["recipe_components"])

        story.append(Spacer(1, 8))
        story.append(Paragraph("Quality test results", STYLES["Heading3"]))
        story.append(_key_value_table([
            ("Overall verdict", data["overall_verdict"]),
            ("Pass", data["pass_count"]), ("Fail", data["fail_count"]), ("", ""),
        ]))
        _section(story, "Results (target vs. actual)", data["quality_results"])
    return _pdf_bytes(build)


def render_sample_certificate_docx(data):
    doc = Document()
    _docx_report_header(
        doc, f"Sample Certificate of Analysis — Sample #{data['sample_id']}",
        f"{data['plant']} · {data['foam_grade']} · Sample Location Reference: {data['zone_label']} · "
        f"Verdict: {data['overall_verdict']}",
    )
    _docx_heading(doc, "Sample source", size=12, color=_HTC_GREY, space_before=6)
    _docx_kv_table(doc, data["header_fields"] + [("Product grade", data["foam_grade"]), ("Plant", data["plant"])])

    _docx_heading(doc, "Sample", size=12, color=_HTC_GREY, space_before=10)
    _docx_kv_table(doc, [
        ("Sample ID", data["sample_id"]), ("Sample Location Reference", data["zone_label"]), ("Sampled", data["sample_ts"]),
    ])
    if data["sample_notes"]:
        doc.add_paragraph(f"Notes: {data['sample_notes']}")

    _docx_heading(doc, "Recipe used", size=12, color=_HTC_GREY, space_before=10)
    _docx_kv_table(doc, [
        ("Recipe version", data["recipe_version_label"]), ("Approval status", data["recipe_approval_status"]),
        ("Effective date", data["recipe_effective_date"]),
        ("Ratio / index", f"{data['recipe_ratio_index']:.3f}" if data["recipe_ratio_index"] is not None else "—"),
    ])
    _docx_section(doc, "Formulation", data["recipe_components"])

    _docx_heading(doc, "Quality test results", size=12, color=_HTC_GREY, space_before=10)
    _docx_kv_table(doc, [
        ("Overall verdict", data["overall_verdict"]),
        ("Pass", data["pass_count"]), ("Fail", data["fail_count"]),
    ])
    _docx_section(doc, "Results (target vs. actual)", data["quality_results"])
    return _docx_bytes(doc)


# ---------------------------------------------------------------------------
# 6. Test Results Report (Physical Property Result page) (CR-01, 2026-08-10:
# renamed from "Quality Test Result Report" - label-only rename)
#
# Placement (per user direction 2026-08-04): this report's subject is a
# comprehensive multi-field selection (Pass/Fail, Property, Foam scope) the
# reader has to build up first, not a single dropdown choice - so it lives
# on pages/5_Physical_Property_Result.py itself, right below the same
# filter controls and Pareto chart it shares its scope with, rather than
# on the Report page. build_quality_test_report_data() never re-derives
# tenant scope or filters on its own - it purely aggregates the exact
# PhysicalPropertyResult id set the page has already scoped and filtered,
# so the report always matches what's on screen at the moment it's
# generated.
# ---------------------------------------------------------------------------

def _qtr_source_and_grade(result):
    """(source label, human-readable parent description, product grade name)
    for a PhysicalPropertyResult, resolving whichever of the three
    mutually exclusive parents (production run / customer trial /
    optimization trial - see db.SAMPLE_SOURCE_TYPES) it belongs to.
    Mirrors pages/5_Physical_Property_Result.py's own
    _result_source_desc()/_result_foam_grade_id() - kept as a local copy
    here since reports.py doesn't import from page modules."""
    if result.production_run_id is not None:
        run = result.production_run
        grade = run.foam_grade if run else None
        desc = f"Run #{run.id} — {grade.grade_name if grade else '—'} · {run.run_date}" if run else f"Run #{result.production_run_id}"
        return "Production Run", desc, grade.grade_name if grade else "—"
    if result.customer_trial_id is not None:
        t = result.customer_trial
        grade = t.foam_grade if t else None
        desc = f"Trial #{t.id} — {t.customer_name}" if t else f"Trial #{result.customer_trial_id}"
        return "Customer Trial", desc, grade.grade_name if grade else "—"
    if result.optimization_trial_id is not None:
        t = result.optimization_trial
        grade = t.foam_grade if t else None
        ref = (t.improvement_initiative_reference or "(no reference)") if t else ""
        desc = f"Trial #{t.id} — {ref}" if t else f"Trial #{result.optimization_trial_id}"
        return "Optimization Trial", desc, grade.grade_name if grade else "—"
    return "—", "—", "—"


def build_quality_test_report_data(session, result_ids, scope):
    """result_ids: PhysicalPropertyResult ids already scoped (tenant) and
    filtered (Pass/Fail, Property, Foam scope) by the caller - see the
    module-level note above. scope: dict of already-formatted display
    strings describing what was selected - pass_fail_label,
    property_label, foam_scope_label - shown in the report header so the
    reader knows exactly what this report does and doesn't cover."""
    results = (
        session.query(PhysicalPropertyResult)
        .filter(PhysicalPropertyResult.id.in_(result_ids)).all()
        if result_ids else []
    )

    detail = []
    for r in results:
        verdict = compute_pass_fail(r.property_name, r.target_value, r.actual_value) or "Not computed"
        source_label, source_desc, grade_name = _qtr_source_and_grade(r)
        detail.append({
            "result": r, "verdict": verdict,
            "source_desc": source_desc, "grade_name": grade_name,
        })

    total = len(detail)
    pass_count = sum(1 for d in detail if d["verdict"] == "Pass")
    fail_count = sum(1 for d in detail if d["verdict"] == "Fail")
    not_computed_count = total - pass_count - fail_count
    total_scored = pass_count + fail_count
    pass_rate = round(100 * pass_count / total_scored) if total_scored else None

    # Failures by property - the same grouping as the on-page Pareto chart
    # (helpers.render_pareto_chart), recomputed here since reports.py
    # can't import a Streamlit-rendering helper. Fail-only, not
    # pass+fail+not-computed per property, since "which properties are
    # behind the failures" is the actionable question a Pareto answers.
    property_fail_counts = {}
    for d in detail:
        if d["verdict"] == "Fail":
            prop = d["result"].property_name
            property_fail_counts[prop] = property_fail_counts.get(prop, 0) + 1
    property_breakdown = [
        {"Property": k, "Fail count": v}
        for k, v in sorted(property_fail_counts.items(), key=lambda kv: -kv[1])
    ]

    grade_counts = {}
    for d in detail:
        bucket = grade_counts.setdefault(d["grade_name"], {"Pass": 0, "Fail": 0, "Not computed": 0})
        bucket[d["verdict"]] += 1
    grade_breakdown = [
        {"Product grade": g, "Pass count": c["Pass"], "Fail count": c["Fail"]}
        for g, c in sorted(grade_counts.items())
    ]
    # Only meaningful as a chart when the selection actually spans more
    # than one grade - a single-grade selection would just re-draw the
    # header metrics as a one-bar "chart".
    show_grade_breakdown = len(grade_counts) > 1

    failing_results = [
        {
            "Source": d["source_desc"], "Property": d["result"].property_name,
            "Target": d["result"].target_value, "Actual": d["result"].actual_value,
            "Unit": d["result"].unit or "",
            "Deviation": (
                round(d["result"].actual_value - d["result"].target_value, 2)
                if d["result"].actual_value is not None and d["result"].target_value is not None
                else None
            ),
            "Product grade": d["grade_name"], "Tested": d["result"].tested_at,
        }
        for d in sorted(detail, key=lambda d: (d["result"].property_name, d["source_desc"]))
        if d["verdict"] == "Fail"
    ]

    return {
        "scope": scope,
        "total_results": total,
        "pass_count": pass_count,
        "fail_count": fail_count,
        "not_computed_count": not_computed_count,
        "pass_rate": pass_rate,
        "property_breakdown": property_breakdown,
        "grade_breakdown": grade_breakdown,
        "show_grade_breakdown": show_grade_breakdown,
        "failing_results": failing_results,
    }


def render_quality_test_report_pdf(data):
    scope = data["scope"]

    def build(story):
        _title_block(
            story, "Test Results Report",
            f"Pass/Fail: {scope['pass_fail_label']} · Property: {scope['property_label']} · "
            f"Product scope: {scope['foam_scope_label']}",
        )
        story.append(_key_value_table([
            ("Results", data["total_results"]),
            ("Pass rate", f"{data['pass_rate']}%" if data["pass_rate"] is not None else "—"),
            ("Pass", data["pass_count"]), ("Fail", data["fail_count"]),
            ("Not computed", data["not_computed_count"]), ("", ""),
        ]))

        _bar_chart(
            story, "Pass / Fail breakdown",
            ["Pass", "Fail", "Not computed"],
            [data["pass_count"], data["fail_count"], data["not_computed_count"]],
        )

        prop_rows = data["property_breakdown"]
        _bar_chart(
            story, "Failures by property",
            [row["Property"] for row in prop_rows], [row["Fail count"] for row in prop_rows],
            note="Which tested properties are behind the failures in this selection." if prop_rows else None,
        )

        if data["show_grade_breakdown"]:
            grade_rows = data["grade_breakdown"]
            _bar_chart(
                story, "Failures by product grade",
                [row["Product grade"] for row in grade_rows], [row["Fail count"] for row in grade_rows],
                note="Shown because this selection spans more than one product grade.",
            )

        _section(story, "Failing results (target vs. actual)", data["failing_results"])
    return _pdf_bytes(build)


def render_quality_test_report_docx(data):
    scope = data["scope"]
    doc = Document()
    _docx_report_header(
        doc, "Test Results Report",
        f"Pass/Fail: {scope['pass_fail_label']} · Property: {scope['property_label']} · "
        f"Product scope: {scope['foam_scope_label']}",
    )
    _docx_kv_table(doc, [
        ("Results", data["total_results"]),
        ("Pass rate", f"{data['pass_rate']}%" if data["pass_rate"] is not None else "—"),
        ("Pass", data["pass_count"]), ("Fail", data["fail_count"]),
        ("Not computed", data["not_computed_count"]),
    ])

    _docx_bar_chart(
        doc, "Pass / Fail breakdown",
        ["Pass", "Fail", "Not computed"],
        [data["pass_count"], data["fail_count"], data["not_computed_count"]],
    )

    prop_rows = data["property_breakdown"]
    _docx_bar_chart(
        doc, "Failures by property",
        [row["Property"] for row in prop_rows], [row["Fail count"] for row in prop_rows],
        note="Which tested properties are behind the failures in this selection." if prop_rows else None,
    )

    if data["show_grade_breakdown"]:
        grade_rows = data["grade_breakdown"]
        _docx_bar_chart(
            doc, "Failures by product grade",
            [row["Product grade"] for row in grade_rows], [row["Fail count"] for row in grade_rows],
            note="Shown because this selection spans more than one product grade.",
        )

    _docx_section(doc, "Failing results (target vs. actual)", data["failing_results"])
    return _docx_bytes(doc)


# ---------------------------------------------------------------------------
# 6. Quality Issues Report (Quality Observation page) (CR-01, 2026-08-10:
# pluralized from "Quality Issue Report" - label-only rename)
#
# Placement/mechanism (per user direction 2026-08-04): same logic as the
# Test Results report - this report's subject is a comprehensive
# multi-field selection (Severity, Foam scope, and the breakdown's
# group-by choice) built up on pages/6_Quality_Observation.py itself, not
# a single dropdown choice, so it lives there rather than on the Report
# page, and it is generated from exactly whatever the page's own filters
# currently have selected.
# ---------------------------------------------------------------------------

def _qi_source_and_grade(obs):
    """(source label, human-readable parent description, product grade name)
    for a QualityObservation, resolving whichever of the three mutually
    exclusive parents it belongs to. Mirrors pages/6_Quality_Observation.
    py's own _obs_source_desc()/_obs_foam_grade_id() - kept as a local
    copy here since reports.py doesn't import from page modules."""
    if obs.production_run_id is not None:
        run = obs.production_run
        grade = run.foam_grade if run else None
        desc = f"Run #{run.id} — {grade.grade_name if grade else '—'} · {run.run_date}" if run else f"Run #{obs.production_run_id}"
        return "Production Run", desc, grade.grade_name if grade else "—"
    if obs.customer_trial_id is not None:
        t = obs.customer_trial
        grade = t.foam_grade if t else None
        desc = f"Trial #{t.id} — {t.customer_name}" if t else f"Trial #{obs.customer_trial_id}"
        return "Customer Trial", desc, grade.grade_name if grade else "—"
    if obs.optimization_trial_id is not None:
        t = obs.optimization_trial
        grade = t.foam_grade if t else None
        ref = (t.improvement_initiative_reference or "(no reference)") if t else ""
        desc = f"Trial #{t.id} — {ref}" if t else f"Trial #{obs.optimization_trial_id}"
        return "Optimization Trial", desc, grade.grade_name if grade else "—"
    return "—", "—", "—"


def build_quality_issue_report_data(session, observation_ids, scope):
    """observation_ids: QualityObservation ids already scoped (tenant) and
    filtered (Severity, Foam scope) by the caller - this function purely
    aggregates that exact set, it never re-derives scope or filters
    itself. scope: dict of already-formatted display strings -
    severity_label, foam_scope_label, group_by_label ("Issue type" or
    "Issue category", matching the page's own breakdown-chart toggle)."""
    observations = (
        session.query(QualityObservation)
        .filter(QualityObservation.id.in_(observation_ids)).all()
        if observation_ids else []
    )

    total = len(observations)
    severity_counts = {s: 0 for s in SEVERITIES}
    for o in observations:
        if o.severity in severity_counts:
            severity_counts[o.severity] += 1
    severity_breakdown = [{"Severity": s, "Count": c} for s, c in severity_counts.items()]

    recurring_count = sum(1 for o in observations if o.frequency == "Recurring")
    one_off_count = total - recurring_count

    confidence_counts = {c: 0 for c in CONFIDENCE_LEVELS}
    for o in observations:
        if o.confidence_level in confidence_counts:
            confidence_counts[o.confidence_level] += 1
    confidence_breakdown = [{"Confidence level": k, "Count": v} for k, v in confidence_counts.items()]

    # Issues by type or category - the same grouping as the page's own
    # breakdown-by-issue Pareto chart, recomputed here since reports.py
    # can't import a Streamlit-rendering helper.
    group_by_col = scope.get("group_by_label") or "Issue type"
    issue_group_counts = {}
    for o in observations:
        if group_by_col == "Issue category":
            label = (quality_issue_taxonomy.lookup(o.observation_type) or {}).get("category") or "Other / not yet classified"
        else:
            label = o.observation_type
        issue_group_counts[label] = issue_group_counts.get(label, 0) + 1
    issue_breakdown = [
        {group_by_col: k, "Count": v}
        for k, v in sorted(issue_group_counts.items(), key=lambda kv: -kv[1])
    ]

    # Priority issues - High severity and/or Recurring, the two signals
    # that mark an issue as needing attention rather than a one-off, minor
    # note. Not every row in the selection, which the page's own CSV
    # export already covers.
    priority_issues = []
    for o in observations:
        if o.severity == "High" or o.frequency == "Recurring":
            _source_label, source_desc, grade_name = _qi_source_and_grade(o)
            priority_issues.append({
                "Source": source_desc, "Issue": o.observation_type,
                "Severity": o.severity or "—", "Frequency": o.frequency or "—",
                "Confidence": o.confidence_level or "—", "Product grade": grade_name,
                "Suspected cause": o.suspected_cause or "—", "Observed": o.observed_at,
            })
    priority_issues.sort(key=lambda r: (0 if r["Severity"] == "High" else 1, 0 if r["Frequency"] == "Recurring" else 1))

    return {
        "scope": scope,
        "total_issues": total,
        "severity_breakdown": severity_breakdown,
        "recurring_count": recurring_count,
        "one_off_count": one_off_count,
        "confidence_breakdown": confidence_breakdown,
        "issue_breakdown": issue_breakdown,
        "group_by_col": group_by_col,
        "priority_issues": priority_issues,
    }


def render_quality_issue_report_pdf(data):
    scope = data["scope"]
    group_col = data["group_by_col"]

    def build(story):
        _title_block(
            story, "Quality Issues Report",
            f"Severity: {scope['severity_label']} · Product scope: {scope['foam_scope_label']} · "
            f"Grouped by: {scope['group_by_label']}",
        )
        high_count = next((r["Count"] for r in data["severity_breakdown"] if r["Severity"] == "High"), 0)
        story.append(_key_value_table([
            ("Issues", data["total_issues"]), ("High severity", high_count),
            ("Recurring", data["recurring_count"]), ("One-off", data["one_off_count"]),
        ]))

        sev_rows = data["severity_breakdown"]
        _bar_chart(
            story, "Severity breakdown",
            [r["Severity"] for r in sev_rows], [r["Count"] for r in sev_rows],
        )

        issue_rows = data["issue_breakdown"]
        _bar_chart(
            story, f"Issues by {group_col.lower()}",
            [r[group_col] for r in issue_rows], [r["Count"] for r in issue_rows],
            note=f"Which {group_col.lower()}(s) occur most often in this selection." if issue_rows else None,
        )

        conf_rows = data["confidence_breakdown"]
        _bar_chart(
            story, "Confidence level breakdown",
            [r["Confidence level"] for r in conf_rows], [r["Count"] for r in conf_rows],
        )

        # _wrapped_section, not _section: this table mixes a long free-
        # text column (Suspected cause) with several rows, which
        # overflowed the page under _section's no-wrap plain-string
        # cells. Confidence/Product grade/Observed are dropped here to keep
        # the remaining columns legible at print width - Source already
        # carries the product grade and date.
        _wrapped_section(
            story, "Priority issues (High severity and/or Recurring)",
            data["priority_issues"], ["Source", "Issue", "Severity", "Frequency", "Suspected cause"],
            [45 * mm, 35 * mm, 18 * mm, 22 * mm, 58 * mm],
        )
    return _pdf_bytes(build)


def render_quality_issue_report_docx(data):
    scope = data["scope"]
    group_col = data["group_by_col"]
    doc = Document()
    _docx_report_header(
        doc, "Quality Issues Report",
        f"Severity: {scope['severity_label']} · Product scope: {scope['foam_scope_label']} · "
        f"Grouped by: {scope['group_by_label']}",
    )
    high_count = next((r["Count"] for r in data["severity_breakdown"] if r["Severity"] == "High"), 0)
    _docx_kv_table(doc, [
        ("Issues", data["total_issues"]), ("High severity", high_count),
        ("Recurring", data["recurring_count"]), ("One-off", data["one_off_count"]),
    ])

    sev_rows = data["severity_breakdown"]
    _docx_bar_chart(doc, "Severity breakdown", [r["Severity"] for r in sev_rows], [r["Count"] for r in sev_rows])

    issue_rows = data["issue_breakdown"]
    _docx_bar_chart(
        doc, f"Issues by {group_col.lower()}",
        [r[group_col] for r in issue_rows], [r["Count"] for r in issue_rows],
        note=f"Which {group_col.lower()}(s) occur most often in this selection." if issue_rows else None,
    )

    conf_rows = data["confidence_breakdown"]
    _docx_bar_chart(
        doc, "Confidence level breakdown",
        [r["Confidence level"] for r in conf_rows], [r["Count"] for r in conf_rows],
    )

    _docx_section(
        doc, "Priority issues (High severity and/or Recurring)", data["priority_issues"],
        columns=["Source", "Issue", "Severity", "Frequency", "Suspected cause"],
    )
    return _docx_bytes(doc)


# ---------------------------------------------------------------------------
# Sample Report (Production Samples / Customer Trials & Samples /
# Optimization Trials & Samples pages)
#
# Placement/mechanism (per user direction 2026-08-04): same pattern as the
# Quality Test Result and Quality Issue reports - a comprehensive,
# multi-field selection built up on the page itself (which run/trial(s),
# which company/plant scope is active), so it lives on each of pages 9/11/
# 12 rather than as a single dropdown choice on the Report page. This
# function is source_type-aware (Production Run / Customer Trial /
# Optimization Trial - see db.SAMPLE_SOURCE_TYPES) but otherwise identical
# across all three pages, so it's one shared function rather than three
# near-duplicates.
#
# This report purely aggregates the exact Sample id set the calling page
# has already scoped (tenant) and filtered - it does not apply any scope
# or filtering of its own. Answers two questions a raw sample list can't:
# how complete is this selection's traceability (coverage - the % of
# samples that actually have a quality test result attached), and what did
# testing find once it happened (pass/fail rate of just the results linked
# to these samples) - both as charts, per the Reports redesign ruling that
# no report in this app dumps a raw row-by-row table.
# ---------------------------------------------------------------------------

def build_sample_report_data(session, source_type, sample_ids, scope):
    """sample_ids: Sample ids already scoped (tenant) and filtered by the
    calling page - see the module note above. scope: dict of already-
    formatted display strings for the report header (selection_label
    describing which run/trial(s)/date range was selected)."""
    samples = (
        session.query(Sample).filter(Sample.id.in_(sample_ids)).all()
        if sample_ids else []
    )
    total_samples = len(samples)

    zone_counts = {}
    for s in samples:
        zone = s.zone_label or "Unspecified"
        zone_counts[zone] = zone_counts.get(zone, 0) + 1
    zone_breakdown = [
        {"Sample Location Reference": k, "Sample count": v}
        for k, v in sorted(zone_counts.items(), key=lambda kv: -kv[1])
    ]

    sample_id_set = {s.id for s in samples}
    linked_results = (
        session.query(PhysicalPropertyResult)
        .filter(PhysicalPropertyResult.sample_id.in_(sample_id_set)).all()
        if sample_id_set else []
    )
    samples_with_results = len({r.sample_id for r in linked_results if r.sample_id is not None})
    coverage_pct = round(100 * samples_with_results / total_samples) if total_samples else None

    pass_count = fail_count = not_computed_count = 0
    for r in linked_results:
        verdict = compute_pass_fail(r.property_name, r.target_value, r.actual_value) or "Not computed"
        if verdict == "Pass":
            pass_count += 1
        elif verdict == "Fail":
            fail_count += 1
        else:
            not_computed_count += 1
    total_scored = pass_count + fail_count
    pass_rate = round(100 * pass_count / total_scored) if total_scored else None

    return {
        "source_type": source_type,
        "scope": scope,
        "total_samples": total_samples,
        "zone_breakdown": zone_breakdown,
        "samples_with_results": samples_with_results,
        "samples_without_results": total_samples - samples_with_results,
        "coverage_pct": coverage_pct,
        "pass_count": pass_count,
        "fail_count": fail_count,
        "not_computed_count": not_computed_count,
        "pass_rate": pass_rate,
    }


def render_sample_report_pdf(data):
    scope = data["scope"]

    def build(story):
        _title_block(story, f"Sample Report — {data['source_type']}", scope.get("selection_label", ""))
        story.append(_key_value_table([
            ("Samples in selection", data["total_samples"]),
            ("With a quality result", data["samples_with_results"]),
            ("Coverage", f"{data['coverage_pct']}%" if data["coverage_pct"] is not None else "—"),
            ("Pass rate (linked results)", f"{data['pass_rate']}%" if data["pass_rate"] is not None else "—"),
            ("Pass", data["pass_count"]), ("Fail", data["fail_count"]),
        ]))

        zone_rows = data["zone_breakdown"]
        _bar_chart(
            story, "Samples by Sample Location Reference",
            [row["Sample Location Reference"] for row in zone_rows], [row["Sample count"] for row in zone_rows],
        )

        _bar_chart(
            story, "Linked quality result outcomes",
            ["Pass", "Fail", "Not computed"],
            [data["pass_count"], data["fail_count"], data["not_computed_count"]],
            note="Pass/Fail of the quality test results recorded against these samples - not every "
                 "sample has one yet, see Coverage above.",
        )
    return _pdf_bytes(build)


def render_sample_report_docx(data):
    scope = data["scope"]
    doc = Document()
    _docx_report_header(doc, f"Sample Report — {data['source_type']}", scope.get("selection_label", ""))
    _docx_kv_table(doc, [
        ("Samples in selection", data["total_samples"]),
        ("With a quality result", data["samples_with_results"]),
        ("Coverage", f"{data['coverage_pct']}%" if data["coverage_pct"] is not None else "—"),
        ("Pass rate (linked results)", f"{data['pass_rate']}%" if data["pass_rate"] is not None else "—"),
        ("Pass", data["pass_count"]), ("Fail", data["fail_count"]),
    ])

    zone_rows = data["zone_breakdown"]
    _docx_bar_chart(doc, "Samples by Sample Location Reference", [row["Sample Location Reference"] for row in zone_rows], [row["Sample count"] for row in zone_rows])

    _docx_bar_chart(
        doc, "Linked quality result outcomes",
        ["Pass", "Fail", "Not computed"],
        [data["pass_count"], data["fail_count"], data["not_computed_count"]],
        note="Pass/Fail of the quality test results recorded against these samples - not every "
             "sample has one yet, see Coverage above.",
    )
    return _docx_bytes(doc)


# ---------------------------------------------------------------------------
# 7. PI3 Q&A Report (DOCX only)
# ---------------------------------------------------------------------------

_HTC_LOGO_PATH = "assets/htc_global_logo_blue_steel.png"
_HTC_BLUE = RGBColor(0x1B, 0x6F, 0xA8)  # matches .streamlit/config.toml primaryColor
_HTC_GREY = RGBColor(0x5A, 0x6B, 0x74)


def build_pi3_qa_report_data(
    question, answer, tool_log, page_context="", plant_name=None,
    foam_grade_name=None, asked_by=None, asked_at=None,
):
    """Plain-dict data assembly for one 'Ask PI3' question/answer exchange -
    no Streamlit or python-docx import, so this half is easy to unit test
    on its own. `tool_log` is exactly what ai_assistant.ask_plant_question()
    returns: a list of dicts, each either
    {"tool": "query_plant_data", "sql", "rows_returned", "rows", ["error"]}
    or {"tool": "get_verified_analysis", "args", "result"}."""
    return {
        "question": (question or "").strip(),
        "answer": (answer or "").strip(),
        "tool_log": tool_log or [],
        "page_context": (page_context or "").strip(),
        "plant_name": plant_name,
        "foam_grade_name": foam_grade_name,
        "asked_by": asked_by,
        "asked_at": asked_at or dt.datetime.utcnow(),
    }


def _docx_heading(doc, text, size=13, color=_HTC_BLUE, space_before=12):
    """A styled heading paragraph, built on a real Word "Heading N" style
    (rather than a bold Normal paragraph) so it behaves like an actual
    heading: it shows up in Word's Navigation Pane/outline view and in any
    auto-generated table of contents, and - the practical reason this
    matters here - `keep_with_next` below stops Word from ever stranding a
    heading alone at the bottom of a page with its content pushed to the
    next one.

    Appearance is still fully controlled here via explicit run-level
    formatting (size/color/bold), same as before, so the look stays
    identical regardless of what Word template the opening machine has -
    using a named style doesn't reintroduce that risk, it just makes the
    style semantically real on top of the same explicit formatting."""
    level = 2 if size >= 15 else (3 if size >= 12 else 4)
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


# Recognizes the fixed structure PI3's system prompt always produces (see
# ai_assistant.py's SYSTEM_PROMPT, section "9) Default Response Structure"):
# short, punctuation-free numbered top-level section titles ("1. Direct
# Answer"), optional single-letter-lettered sub-sections ("A. Reduced
# silicone performance"), and "- " prefixed list items. Matching on these
# turns PI3's plain text into real headings/bullets instead of one flat
# Normal paragraph per line - deliberately conservative (short line, starts
# with a capital letter, no internal period, no trailing period) so an
# ordinary sentence that happens to start with a number or single letter
# doesn't get misread as a heading.
_TOP_HEADING_RE = re.compile(r"^\d{1,2}\.\s+[A-Z][^.\n]{2,78}$")
_SUB_HEADING_RE = re.compile(r"^[A-J]\.\s+[A-Z][^.\n]{2,98}$")
_BULLET_RE = re.compile(r"^[-•*]\s+(\S.*)$")

# PI3's free-form question-answering prompt (ai_assistant.PLANT_QUERY_SYSTEM_PROMPT,
# used by the Ask PI3 box) doesn't forbid markdown the way the fixed-prompt
# one does, so its answers commonly contain **bold** / *italic* / `code`
# inline. Without this, those markers came through as literal asterisks/
# backticks in the Word doc instead of real formatting.
_INLINE_MD_RE = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`")

# PI3 also frequently answers with a GFM-style markdown table (a header
# row, a "|---|---:|" alignment/separator row, then data rows) when
# comparing several components or properties - confirmed in production
# output (e.g. a Recipe Optimization report's "Recommended formulation
# direction" and "Target-property focus" tables). Before this existed,
# every line of a markdown table fell through to the plain-paragraph case
# below and rendered as literal "| Water | 3.00 php | ... |" / "|---|---:|"
# text - unreadable, and the single biggest formatting complaint on these
# reports. _PIPE_ROW_RE spots a candidate row; _SEP_CELL_RE recognizes the
# separator row's cells (dashes, optionally colon-flanked for alignment)
# so a real header+data table can be told apart from an ordinary line that
# merely happens to contain a "|" character.
_PIPE_ROW_RE = re.compile(r"^\|.*\|$")
_SEP_CELL_RE = re.compile(r"^:?-{3,}:?$")


def _split_md_table_row(line):
    """Split a markdown table row ('| a | b |') into ['a', 'b'], honoring a
    backslash-escaped pipe inside a cell and dropping the empty strings
    produced by the row's own leading/trailing pipes."""
    protected = line.replace("\\|", "\x00")
    cells = [c.strip().replace("\x00", "|") for c in protected.split("|")]
    if cells and cells[0] == "":
        cells = cells[1:]
    if cells and cells[-1] == "":
        cells = cells[:-1]
    return cells


def _is_md_table_separator(line):
    if not _PIPE_ROW_RE.match(line):
        return False
    cells = _split_md_table_row(line)
    return bool(cells) and all(_SEP_CELL_RE.match(c) for c in cells)


def _strip_inline_markdown(text):
    """Plain-text version of a line with markdown markers removed but their
    content kept - used for heading lines, which are already bold/colored
    by their own style, so there's no run-formatting reason to parse
    **bold**/*italic* there, just to avoid showing the literal markers."""
    return _INLINE_MD_RE.sub(lambda m: next(g for g in m.groups() if g is not None), text)


def _add_runs_with_inline_markdown(paragraph, text, size=None):
    """Append `text` to `paragraph` as one or more runs, converting
    **bold**, *italic*, and `code` markdown spans into real run formatting
    instead of leaving the literal markers in the output."""
    pos = 0
    for m in _INLINE_MD_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            if size:
                run.font.size = size
        if m.group(1) is not None:
            run = paragraph.add_run(m.group(1))
            run.bold = True
        elif m.group(2) is not None:
            run = paragraph.add_run(m.group(2))
            run.italic = True
        else:
            run = paragraph.add_run(m.group(3))
            run.font.name = "Consolas"
        if size:
            run.font.size = size
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        if size:
            run.font.size = size


def _docx_markdown_table(doc, header_cells, data_rows):
    """Renders a parsed markdown table (a header cell list plus a list of
    data-row cell lists, both already produced by _split_md_table_row) as a
    real bordered Word table - same "Light Grid Accent 1" style used for
    PI3's own SQL-result appendix table (_docx_data_table), so a markdown
    table and a data table look consistent in the same report. Ragged rows
    (a data row with a different cell count than the header) are padded or
    truncated to the header's column count rather than raising - a
    model-written table is exactly the kind of input that occasionally
    comes out uneven."""
    ncols = len(header_cells)
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Light Grid Accent 1"
    for cell, header_text in zip(table.rows[0].cells, header_cells):
        run = cell.paragraphs[0].add_run(_strip_inline_markdown(header_text))
        run.bold = True
        run.font.size = Pt(9)
    for row_cells in data_rows:
        padded = (row_cells + [""] * ncols)[:ncols]
        cells = table.add_row().cells
        for cell, cell_text in zip(cells, padded):
            _add_runs_with_inline_markdown(cell.paragraphs[0], cell_text or "—", size=Pt(9))
    return table


def _render_ai_answer_body(doc, text):
    """Render a PI3 answer's plain text into real Word structure: numbered
    top-level sections and lettered sub-sections become real headings,
    "- " list items become real bulleted paragraphs, a markdown pipe table
    (header row + "|---|" separator row + data rows) becomes a real Word
    table, inline **bold**/*italic*/`code` markdown becomes real run
    formatting, and everything else stays a normal paragraph. Replaces the
    previous behavior of dumping one flat Normal paragraph per non-blank
    line verbatim, which produced an unreadable wall of text with no
    headings, bullets, or tables, and left literal markdown markers
    (including whole "| a | b |" / "|---|---:|" table rows) in place."""
    lines = [raw_line.strip() for raw_line in (text or "").split("\n")]
    n = len(lines)
    i = 0
    while i < n:
        line = lines[i]
        if not line:
            i += 1
            continue

        if _PIPE_ROW_RE.match(line) and i + 1 < n and _is_md_table_separator(lines[i + 1]):
            header_cells = _split_md_table_row(line)
            j = i + 2
            data_rows = []
            while j < n and _PIPE_ROW_RE.match(lines[j]):
                data_rows.append(_split_md_table_row(lines[j]))
                j += 1
            _docx_markdown_table(doc, header_cells, data_rows)
            i = j
            continue

        bullet_match = _BULLET_RE.match(line)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_inline_markdown(p, bullet_match.group(1), size=Pt(10.5))
            i += 1
            continue
        if _TOP_HEADING_RE.match(line):
            _docx_heading(doc, _strip_inline_markdown(line), size=13, space_before=14)
            i += 1
            continue
        if _SUB_HEADING_RE.match(line):
            _docx_heading(doc, _strip_inline_markdown(line), size=11.5, color=_HTC_GREY, space_before=10)
            i += 1
            continue
        p = doc.add_paragraph()
        _add_runs_with_inline_markdown(p, line)
        i += 1


def _docx_kv_table(doc, pairs):
    table = doc.add_table(rows=0, cols=2)
    table.autofit = True
    for label, value in pairs:
        row = table.add_row().cells
        label_run = row[0].paragraphs[0].add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(9.5)
        row[1].paragraphs[0].add_run("—" if value in (None, "") else str(value)).font.size = Pt(9.5)
    return table


def _docx_data_table(doc, rows, max_rows=200):
    """Renders a list-of-dicts as a bordered table, capped at max_rows so a
    very large query result doesn't produce an unusable multi-hundred-page
    appendix - the SQL that produced it is always shown alongside, so the
    full result set is still reproducible."""
    if not rows:
        doc.add_paragraph("No rows returned.").runs[0].font.size = Pt(9)
        return
    shown = rows[:max_rows]
    headers = list(shown[0].keys())
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for cell, header in zip(table.rows[0].cells, headers):
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(8.5)
    for row_data in shown:
        cells = table.add_row().cells
        for cell, header in zip(cells, headers):
            value = row_data.get(header)
            cell.paragraphs[0].add_run("—" if value in (None, "") else str(value)).font.size = Pt(8.5)
    if len(rows) > max_rows:
        note = doc.add_paragraph(f"... {len(rows) - max_rows} further row(s) not shown.")
        note.runs[0].italic = True
        note.runs[0].font.size = Pt(8.5)


def render_pi3_qa_report_docx(data):
    """Renders one PI3 Q&A exchange as DOCX bytes: HTC-branded header,
    a metadata block, the question, PI3's answer, an advisory-boundary
    disclaimer, and an appendix showing exactly what PI3 checked to
    produce the answer. Same layout, same styling, every single time -
    that consistency is the whole point of generating this from code
    rather than starting from a hand-edited Word file each time."""
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # --- Header: logo + title -------------------------------------------
    header = doc.add_table(rows=1, cols=2)
    header.autofit = False
    header.columns[0].width = Cm(3.6)
    logo_cell, title_cell = header.rows[0].cells
    if os.path.exists(_HTC_LOGO_PATH):
        run = logo_cell.paragraphs[0].add_run()
        run.add_picture(_HTC_LOGO_PATH, width=Cm(3.0))
    title_p = title_cell.paragraphs[0]
    title_run = title_p.add_run("PI3 Q&A Report")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = _HTC_BLUE
    subtitle_p = title_cell.add_paragraph()
    subtitle_run = subtitle_p.add_run("Rigid foam expert system | HTC Global Co. Ltd")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = _HTC_GREY

    doc.add_paragraph()

    # --- Metadata ----------------------------------------------------------
    _docx_kv_table(doc, [
        ("Generated", data["asked_at"].strftime("%Y-%m-%d %H:%M UTC")),
        ("Plant", data.get("plant_name") or "—"),
        ("Product grade", data.get("foam_grade_name") or "—"),
        ("Asked by", data.get("asked_by") or "—"),
        ("Page context", data.get("page_context") or "—"),
    ])

    # --- Question / answer ---------------------------------------------
    # "Question asked"/"PI3's answer"/"Appendix" are the report's top-level
    # sections, so all three sit at Heading 2 - that leaves Heading 3 free
    # for PI3's own numbered sections ("1. Direct Answer") to nest properly
    # underneath "PI3's answer" instead of sitting as its siblings.
    _docx_heading(doc, "Question asked", size=15)
    doc.add_paragraph(data["question"] or "—")

    _docx_heading(doc, "PI3's answer", size=15)
    _render_ai_answer_body(doc, data["answer"] or "—")

    disclaimer = doc.add_paragraph()
    disclaimer.paragraph_format.space_before = Pt(10)
    disc_run = disclaimer.add_run(
        "This is historical reference for the reviewer's own investigation, not an "
        "instruction. Confirm through your own investigation before acting on it."
    )
    disc_run.italic = True
    disc_run.font.size = Pt(9)
    disc_run.font.color.rgb = _HTC_GREY

    # --- Appendix: exactly what PI3 checked ------------------------------
    doc.add_page_break()
    _docx_heading(doc, "Appendix: data PI3 checked", size=15)
    tool_log = data.get("tool_log") or []
    if not tool_log:
        doc.add_paragraph("No tool calls were recorded for this answer.")
    for i, entry in enumerate(tool_log, start=1):
        tool_name = entry.get("tool", "unknown tool")
        _docx_heading(doc, f"{i}. {tool_name}", size=13, color=_HTC_GREY, space_before=14)
        if tool_name == "query_plant_data":
            sql_p = doc.add_paragraph()
            sql_run = sql_p.add_run(entry.get("sql", "—"))
            sql_run.font.name = "Consolas"
            sql_run.font.size = Pt(9)
            if "error" in entry:
                err_p = doc.add_paragraph()
                err_run = err_p.add_run(f"Rejected: {entry['error']}")
                err_run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
                err_run.font.size = Pt(9)
            else:
                count_p = doc.add_paragraph()
                count_run = count_p.add_run(f"{entry.get('rows_returned', 0)} row(s) returned:")
                count_run.font.size = Pt(9)
                _docx_data_table(doc, entry.get("rows") or [])
        elif tool_name == "get_verified_analysis":
            args_p = doc.add_paragraph()
            args_run = args_p.add_run(f"Arguments: {entry.get('args')}")
            args_run.font.size = Pt(9)
            result = entry.get("result")
            if isinstance(result, dict):
                kv_pairs = []
                table_keys = []  # list-of-dicts values, rendered as a sub-table below instead
                for k, v in result.items():
                    if isinstance(v, list) and v and isinstance(v[0], dict):
                        kv_pairs.append((k, f"[{len(v)} row(s) - see table below]"))
                        table_keys.append(k)
                    elif isinstance(v, list):
                        # Plain-value list (e.g. warnings/successes strings) - show the
                        # actual content inline rather than hiding it behind a count.
                        kv_pairs.append((k, "; ".join(str(x) for x in v) if v else "—"))
                    else:
                        kv_pairs.append((k, v))
                _docx_kv_table(doc, kv_pairs)
                for k in table_keys:
                    sub = doc.add_paragraph()
                    sub_run = sub.add_run(k)
                    sub_run.bold = True
                    sub_run.font.size = Pt(9)
                    _docx_data_table(doc, result[k])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Shared Word (.docx) report infrastructure
#
# Added 2026-08-04 per user direction: every "Download PDF" button app-wide
# is being replaced with a "Download Word" button (the user's original,
# earlier instruction was Word format throughout - PDF was a mistake on
# this app's part). Rather than hand-roll python-docx calls in each of the
# 15 render_*_docx functions below, these helpers mirror the PDF-side
# building blocks (_title_block/_key_value_table/_section/_bar_chart/
# _line_chart) one-for-one, so every render_*_docx function reads almost
# identically to its render_*_pdf twin. Charts reuse the exact same
# reportlab Drawing objects the PDF uses (_build_bar_chart_drawing/
# _build_line_chart_drawing above) rasterized to PNG via renderPM, so the
# chart itself is pixel-identical between the two formats, not a
# separately-drawn approximation.
# ---------------------------------------------------------------------------

def _nice_ticks(vmin, vmax, target=5):
    """Round-number axis ticks (the classic 1/2/5-times-a-power-of-ten
    algorithm), so small integer ranges - e.g. a "samples by zone" count of
    0-3 - get labels like 0, 1, 2, 3 instead of the literal vmin+span*i/5
    fractions the first cut of these Pillow charts used (which produced
    duplicate-looking labels like two rows both showing "2" and an odd
    "0.60"). Returns (ticks, nice_min, nice_max)."""
    if vmin == vmax:
        vmin, vmax = vmin - 1, vmax + 1
    raw_step = (vmax - vmin) / target
    exponent = math.floor(math.log10(raw_step)) if raw_step > 0 else 0
    fraction = raw_step / (10 ** exponent)
    nice_fraction = 1 if fraction <= 1 else 2 if fraction <= 2 else 5 if fraction <= 5 else 10
    step = nice_fraction * (10 ** exponent)
    nice_min = math.floor(vmin / step) * step
    nice_max = math.ceil(vmax / step) * step
    n = max(1, round((nice_max - nice_min) / step))
    ticks = [nice_min + i * step for i in range(n + 1)]
    return ticks, nice_min, nice_max, step


def _fmt_tick(v, step):
    if step >= 1:
        return f"{v:,.0f}"
    decimals = max(0, -int(math.floor(math.log10(step))))
    return f"{v:,.{decimals}f}"


def _pil_font(size):
    """ImageFont.load_default() has accepted a `size` argument (giving a
    real scalable bitmap font instead of the tiny fixed 10px default) since
    Pillow 9.2 - this app already requires Pillow 12.x as a transitive
    reportlab dependency, so this is always available. No external .ttf
    file needed, which matters: anything requiring a font file path is one
    more thing that can be missing on a fresh Streamlit Cloud build."""
    return ImageFont.load_default(size=size)


def _pil_bar_chart_png(categories, values, bar_color_hex="#4A7A9D", zero_floor=True,
                        width_px=1400, height_px=620):
    """Pure-Pillow bar chart, rendered directly to PNG bytes for embedding
    in a Word document. Replaces the old reportlab-Drawing +
    renderPM.drawToFile() rasterization path (removed 2026-08-05): reportlab
    5.0's renderPM hard-requires the rlPyCairo backend with no fallback
    (see reportlab.graphics.renderPM._getPMBackend - 'cairo' not in the
    backend name raises immediately), and pycairo publishes no Linux wheels
    at all, so it can never install cleanly on Streamlit Cloud without a
    from-source build needing system cairo dev headers - the same class of
    problem already hit once with scipy/gfortran (see requirements.txt
    history) and not worth repeating. Pillow is a mandatory, always-
    installed reportlab dependency with real manylinux wheels for every
    Python version this app runs on, so drawing the chart with it directly
    sidesteps the cairo dependency entirely rather than trying to install
    it. This crashed every Word report with a chart in production
    (RenderPMError: cannot import desired renderPM backend rlPyCairo) until
    fixed."""
    cats = [str(c)[:16] for c in categories]
    vals = [0.0 if v is None else float(v) for v in values]

    margin_l, margin_r, margin_t, margin_b = 100, 30, 20, 90
    plot_w = width_px - margin_l - margin_r
    plot_h = height_px - margin_t - margin_b

    if zero_floor:
        raw_max = max(vals) if max(vals) > 0 else 1.0
        ticks, _, vmax, step = _nice_ticks(0.0, raw_max)
        vmin = 0.0
        ticks = [t for t in ticks if t >= 0] or [0.0]
    else:
        raw_min, raw_max = min(vals), max(vals)
        if raw_min == raw_max:
            raw_min, raw_max = raw_min - 1, raw_max + 1
        pad = (raw_max - raw_min) * 0.1
        ticks, vmin, vmax, step = _nice_ticks(raw_min - pad, raw_max + pad)
    span = (vmax - vmin) or 1.0

    img = Image.new("RGB", (width_px, height_px), "white")
    draw = ImageDraw.Draw(img)
    font = _pil_font(22)

    def y_for(v):
        return margin_t + plot_h - (v - vmin) / span * plot_h

    for v in ticks:
        y = y_for(v)
        draw.line([(margin_l, y), (width_px - margin_r, y)], fill="#E5E5E5", width=1)
        label = _fmt_tick(v, step)
        bbox = draw.textbbox((0, 0), label, font=font)
        draw.text((margin_l - 12 - (bbox[2] - bbox[0]), y - (bbox[3] - bbox[1]) / 2 - bbox[1]),
                   label, font=font, fill="#555555")

    n = len(cats) or 1
    slot_w = plot_w / n
    bar_w = min(slot_w * 0.55, 100)
    y_zero = y_for(0 if vmin <= 0 <= vmax else vmin)

    for i, (cat, v) in enumerate(zip(cats, vals)):
        cx = margin_l + slot_w * (i + 0.5)
        y_val = y_for(v)
        top, bottom = sorted([y_val, y_zero])
        if bottom - top < 1:
            bottom = top + 1
        draw.rectangle([cx - bar_w / 2, top, cx + bar_w / 2, bottom], fill=bar_color_hex)
        bbox = draw.textbbox((0, 0), cat, font=font)
        draw.text((cx - (bbox[2] - bbox[0]) / 2, margin_t + plot_h + 14), cat, font=font, fill="#333333")

    draw.line([(margin_l, margin_t), (margin_l, margin_t + plot_h)], fill="#999999", width=2)
    draw.line([(margin_l, margin_t + plot_h), (width_px - margin_r, margin_t + plot_h)], fill="#999999", width=2)

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _pil_line_chart_png(categories, series, width_px=1400, height_px=640):
    """Pure-Pillow multi-line chart, rendered directly to PNG bytes for
    embedding in a Word document - see _pil_bar_chart_png's docstring for
    why this replaced the reportlab-Drawing + renderPM path. series is an
    ordered list of (label, values) pairs sharing the categories x-axis;
    None values are simply skipped when drawing each line's polyline
    (a real gap, not interpolated), matching the "no gaps" note the old
    HorizontalLineChart-based version required callers to pre-filter for
    anyway."""
    cats = [str(c)[:10] for c in categories]
    all_vals = [v for _, vals in series for v in vals if v is not None]
    raw_min, raw_max = (min(all_vals), max(all_vals)) if all_vals else (0.0, 1.0)
    if raw_min == raw_max:
        raw_min, raw_max = raw_min - 1, raw_max + 1
    pad = (raw_max - raw_min) * 0.1
    ticks, vmin, vmax, tick_step = _nice_ticks(raw_min - pad, raw_max + pad)
    span = (vmax - vmin) or 1.0

    margin_l, margin_r, margin_t, margin_b = 100, 30, 20, 80
    plot_w = width_px - margin_l - margin_r
    plot_h = height_px - margin_t - margin_b
    n = len(cats)

    img = Image.new("RGB", (width_px, height_px), "white")
    draw = ImageDraw.Draw(img)
    font = _pil_font(20)

    def y_for(v):
        return margin_t + plot_h - (v - vmin) / span * plot_h

    def x_for(i):
        return margin_l + (plot_w * i / (n - 1) if n > 1 else plot_w / 2)

    for v in ticks:
        y = y_for(v)
        draw.line([(margin_l, y), (width_px - margin_r, y)], fill="#E5E5E5", width=1)
        label = _fmt_tick(v, tick_step)
        bbox = draw.textbbox((0, 0), label, font=font)
        draw.text((margin_l - 12 - (bbox[2] - bbox[0]), y - (bbox[3] - bbox[1]) / 2 - bbox[1]),
                   label, font=font, fill="#555555")

    step = max(1, n // 12)
    for i, cat in enumerate(cats):
        if i % step and i != n - 1:
            continue
        x = x_for(i)
        bbox = draw.textbbox((0, 0), cat, font=font)
        draw.text((x - (bbox[2] - bbox[0]) / 2, margin_t + plot_h + 10), cat, font=font, fill="#333333")

    for s_idx, (_, series_vals) in enumerate(series):
        color = _LINE_COLOR_HEX[s_idx % len(_LINE_COLOR_HEX)]
        pts = [(x_for(i), y_for(v)) for i, v in enumerate(series_vals) if v is not None]
        if len(pts) >= 2:
            draw.line(pts, fill=color, width=3, joint="curve")
        for (x, y) in pts:
            draw.ellipse([x - 4, y - 4, x + 4, y + 4], fill=color)

    draw.line([(margin_l, margin_t), (margin_l, margin_t + plot_h)], fill="#999999", width=2)
    draw.line([(margin_l, margin_t + plot_h), (width_px - margin_r, margin_t + plot_h)], fill="#999999", width=2)

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _docx_report_header(doc, title, subtitle=None):
    """HTC-branded header (logo + title + optional subtitle) shared by every
    generated report - same layout as render_pi3_qa_report_docx's header,
    generalized so every report type looks consistent regardless of which
    page generated it."""
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    header = doc.add_table(rows=1, cols=2)
    header.autofit = False
    header.columns[0].width = Cm(3.6)
    logo_cell, title_cell = header.rows[0].cells
    if os.path.exists(_HTC_LOGO_PATH):
        run = logo_cell.paragraphs[0].add_run()
        run.add_picture(_HTC_LOGO_PATH, width=Cm(3.0))
    title_p = title_cell.paragraphs[0]
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = _HTC_BLUE
    if subtitle:
        subtitle_p = title_cell.add_paragraph()
        subtitle_run = subtitle_p.add_run(subtitle)
        subtitle_run.italic = True
        subtitle_run.font.size = Pt(10)
        subtitle_run.font.color.rgb = _HTC_GREY
    doc.add_paragraph()


def _docx_section(doc, title, rows, columns=None, max_rows=200):
    """Word twin of _section (and, since Word tables wrap free text on
    their own, also of _wrapped_section) - a heading followed by a bordered
    table built from a list of dicts, or a plain "No data recorded." line
    if rows is empty. columns, if given, narrows which dict keys are shown
    and in what order - the same purpose _wrapped_section's explicit
    `columns` argument serves on the PDF side."""
    _docx_heading(doc, title, size=12, color=_HTC_GREY, space_before=10)
    if not rows:
        doc.add_paragraph("No data recorded.")
        return
    if columns:
        rows = [{c: r.get(c) for c in columns} for r in rows]
    _docx_data_table(doc, rows, max_rows=max_rows)


def _docx_bar_chart(doc, title, categories, values, note=None, zero_floor=True,
                     bar_color=colors.HexColor("#4A7A9D")):
    """Word twin of _bar_chart - a pure-Pillow PNG embed (see
    _pil_bar_chart_png's docstring) since Word has no native vector-chart
    support via python-docx."""
    _docx_heading(doc, title, size=12, color=_HTC_GREY, space_before=10)
    if note:
        note_p = doc.add_paragraph(note)
        note_p.runs[0].font.size = Pt(9)
    if not categories or not any(v not in (None, 0) for v in values):
        doc.add_paragraph("No data recorded.")
        return
    png_bytes = _pil_bar_chart_png(categories, values, bar_color_hex="#" + bar_color.hexval()[2:],
                                    zero_floor=zero_floor)
    doc.add_picture(io.BytesIO(png_bytes), width=Cm(16))


def _docx_line_chart(doc, title, categories, series, note=None):
    """Word twin of _line_chart - a pure-Pillow PNG embed (see
    _pil_line_chart_png's docstring), with the same color-coded legend line
    underneath."""
    _docx_heading(doc, title, size=12, color=_HTC_GREY, space_before=10)
    if note:
        note_p = doc.add_paragraph(note)
        note_p.runs[0].font.size = Pt(9)
    if not categories or not series or not any(any(v is not None for v in vals) for _, vals in series):
        doc.add_paragraph("No data recorded.")
        return
    png_bytes = _pil_line_chart_png(categories, series)
    doc.add_picture(io.BytesIO(png_bytes), width=Cm(16))
    legend = doc.add_paragraph()
    for i, (label, _) in enumerate(series):
        run = legend.add_run(("   " if i else "") + "■ " + label)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(_LINE_COLOR_HEX[i % len(_LINE_COLOR_HEX)].lstrip("#"))


def _docx_bytes(doc):
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# 10. Recipe Optimization Report (Context / Analysis / Conclusions)
#
# Added 2026-08-04 as part of the Industrial Intelligence reports batch (per
# user direction: "For most of the pages we have a PI3 Word generator on the
# actual page... These reports should have: A. Context B. Analysis
# C. Conclusions. The PI3 button stays on is only for the PI3 analysis or
# answer of a question."). Deliberately NOT the PI3-generated recommendation
# further up the page (that already has its own Word download via
# render_pi3_docx_download) - this is the page's own deterministic analysis:
# does the current recipe meet target, and where it doesn't, which raw
# material's actual dosage is the strongest lead. Lives on
# pages/15_Recipe_Optimization.py itself (a comprehensive multi-field
# selection - product grade, include-trials toggle, correlation property - not
# a single dropdown choice), same placement logic as the Quality Test Result
# Report. build_recipe_optimization_report_data() never re-derives the
# page's own selection - it takes the exact grade, current recipe version,
# cost, expectation summary, and correlation ranking the page has already
# computed, so the report always matches what's on screen.
#
# WP4 (2026-08-07, Converged Joint Implementation Plan section 7.5) added
# build_rigid_recipe_optimization_report_data() alongside this for a
# rigid-product grade's spec-based achievement table (wp3_conformance.
# compute_grade_achievement_summary) instead of the flexible industry-
# tolerance model - see that function's own docstring. It deliberately
# REUSES render_recipe_optimization_report_pdf/docx below rather than
# forking them, since both already build their tables from generic
# lists-of-dicts.
# ---------------------------------------------------------------------------

def build_recipe_optimization_report_data(
    session, grade, current_version, current_cost, expectation_summary,
    corr_property, actual_ranked, include_trials,
):
    """grade: FoamGrade. current_version: RecipeVersion (the active one).
    current_cost: dict from analytics.recipe_version_cost(). expectation_summary:
    DataFrame already computed by the page ("Does the current recipe meet
    target?" - property_name/avg_actual/avg_target/unit/achieved/tolerance/
    n/n_outside). corr_property: the currently selected property for the
    ingredient-correlation drill-down. actual_ranked: DataFrame from
    analytics.rank_component_actual_correlations() for corr_property."""
    plant_name = None
    if grade.product_family and grade.product_family.plant:
        plant_name = grade.product_family.plant.name

    cost_per_kg = None
    if current_cost.get("total_cost") is not None and current_cost.get("total_php"):
        cost_per_kg = round(current_cost["total_cost"] / current_cost["total_php"], 2)
    cost_coverage_pct = (
        round((current_cost["priced_php"] / current_cost["total_php"]) * 100, 0)
        if current_cost.get("total_php") else None
    )

    expectation_rows = []
    achieved_count = 0
    not_achieved_count = 0
    deviation_categories, deviation_values = [], []
    if expectation_summary is not None and not expectation_summary.empty:
        for _, row in expectation_summary.iterrows():
            n_outside = int(row["n_outside"]) if pd.notna(row["n_outside"]) else 0
            n_total = int(row["n"]) if pd.notna(row["n"]) else 0
            expectation_rows.append({
                "Property": row["property_name"],
                "Avg actual": row["avg_actual"],
                "Required (target)": row["avg_target"],
                "UOM": row["unit"] or "—",
                "Tolerance": row.get("tolerance", "—"),
                "Achieved?": row["achieved"],
                "Runs outside tolerance": f"{n_outside} of {n_total}" if n_total else "—",
            })
            if row["achieved"] == "Yes":
                achieved_count += 1
            elif row["achieved"] == "No":
                not_achieved_count += 1
            if pd.notna(row["avg_actual"]) and pd.notna(row["avg_target"]) and row["avg_target"]:
                deviation_categories.append(row["property_name"])
                deviation_values.append(
                    round(100 * (row["avg_actual"] - row["avg_target"]) / row["avg_target"], 1)
                )

    correlation_rows = []
    correlation_categories, correlation_values = [], []
    top_correlation_line = None
    if actual_ranked is not None and not actual_ranked.empty:
        for _, row in actual_ranked.iterrows():
            correlation_rows.append({
                "Raw material": row["raw_material_name"],
                "Runs compared": int(row["n_runs"]),
                "Correlation with outcome": round(row["correlation"], 3),
            })
            correlation_categories.append(row["raw_material_name"])
            correlation_values.append(round(row["correlation"], 3))
        top = actual_ranked.iloc[0]
        top_correlation_line = (
            f"Strongest association for {corr_property}: {top['raw_material_name']} "
            f"(correlation {top['correlation']:+.3f} across {int(top['n_runs'])} production runs' "
            "metered dosage) - a lead to investigate on the floor, not a confirmed cause."
        )

    conclusions = []
    if expectation_rows:
        conclusions.append(
            f"{achieved_count} of {achieved_count + not_achieved_count} tracked properties are "
            f"achieved under the current recipe ({current_version.version_label}); "
            f"{not_achieved_count} are not."
        )
    else:
        conclusions.append(
            f"No quality test results recorded yet under the current recipe "
            f"({current_version.version_label}) to judge against target."
        )
    if top_correlation_line:
        conclusions.append(top_correlation_line)
    else:
        conclusions.append(
            f"Not enough metered stream-reading data paired with {corr_property} results yet to "
            "identify a leading ingredient correlation."
        )
    if cost_per_kg is not None:
        coverage_note = f" ({cost_coverage_pct:.0f}% cost coverage)" if cost_coverage_pct is not None else ""
        conclusions.append(f"Current formulation cost: {cost_per_kg:.2f} USD per kg{coverage_note}.")
    else:
        conclusions.append("No cost data recorded for the current formulation yet.")

    return {
        "grade_name": grade.grade_name,
        "plant_name": plant_name or "—",
        "version_label": current_version.version_label,
        "version_status": current_version.approval_status,
        "include_trials": include_trials,
        "component_count": len(current_version.components or []),
        "cost_per_kg": cost_per_kg,
        "cost_coverage_pct": cost_coverage_pct,
        "expectation_rows": expectation_rows,
        "deviation_categories": deviation_categories,
        "deviation_values": deviation_values,
        "corr_property": corr_property,
        "correlation_rows": correlation_rows,
        "correlation_categories": correlation_categories,
        "correlation_values": correlation_values,
        "conclusions": conclusions,
        "generated_at": dt.datetime.utcnow(),
    }


def build_rigid_recipe_optimization_report_data(
    session, grade, current_version, current_cost, achievement_summary,
    corr_property, actual_ranked, include_trials,
):
    """WP4 (Converged Joint Implementation Plan, section 7.5) rigid-foam
    equivalent of build_recipe_optimization_report_data() above, for the
    same Recipe Optimization Report placement on pages/15_Recipe_
    Optimization.py. Deliberately REUSES render_recipe_optimization_report_
    pdf/docx below rather than duplicating them: both renderers build their
    Analysis table from a plain list-of-dicts via _section/_docx_section, so
    a differently-shaped table (spec context, operator-aware limit text,
    Excluded/Invalid/No result counts - none of which the flexible
    achievement model has) renders through the exact same functions with no
    PDF/DOCX-specific rigid branch needed there. Only the *data assembly* is
    rigid-specific.

    grade: FoamGrade (chemistry_id populated - a rigid-product grade).
    current_version: RecipeVersion (the active one). current_cost: dict
    from analytics.recipe_version_cost(). achievement_summary: the list of
    dicts already computed by the page - wp3_conformance.
    compute_grade_achievement_summary()'s return value, the rigid "Does the
    current recipe meet target?" table. corr_property: the currently
    selected specification's property_name for the ingredient-correlation
    drill-down. actual_ranked: DataFrame from wp3_conformance.
    rank_lot_use_actual_correlations() for that specification - same
    raw_material_name/n_runs/correlation shape as the flexible app's
    rank_component_actual_correlations(), so the correlation section below
    needs no rigid-specific handling either."""
    plant_name = None
    if grade.product_family and grade.product_family.plant:
        plant_name = grade.product_family.plant.name

    cost_per_kg = None
    if current_cost.get("total_cost") is not None and current_cost.get("total_php"):
        cost_per_kg = round(current_cost["total_cost"] / current_cost["total_php"], 2)
    cost_coverage_pct = (
        round((current_cost["priced_php"] / current_cost["total_php"]) * 100, 0)
        if current_cost.get("total_php") else None
    )

    def _limit_text(row):
        op = row["target_operator"] or "<="
        unit = row["unit"] or ""
        if op == "between":
            if row["lower_limit"] is None or row["upper_limit"] is None:
                return "—"
            return f"{row['lower_limit']} – {row['upper_limit']} {unit}".strip()
        if row["target_value"] is None:
            return "—"
        return f"{op} {row['target_value']} {unit}".strip()

    def _context_text(row):
        bits = [b for b in [row["condition"], row["orientation"], row["location"]] if b]
        return " · ".join(bits) if bits else "—"

    expectation_rows = []
    achieved_count = 0
    not_achieved_count = 0
    no_release_specs = []
    deviation_categories, deviation_values = [], []
    for row in achievement_summary or []:
        expectation_rows.append({
            "Property": row["property_name"],
            "Context": _context_text(row),
            "Avg actual (current recipe)": row["avg_actual"] if row["avg_actual"] is not None else "—",
            "Limit / target": _limit_text(row),
            "Achieved?": row["achieved"],
            "Runs Fail": f"{row['n_fail']} of {row['n']}" if row["n"] else "—",
            "Excluded / Invalid / No result": (
                f"{row['n_excluded_context']} / {row['n_invalid']} / {row['n_no_result']}"
            ),
            # CR-09 (2026-08-12): this used to write row["production_release"]
            # (wp3_conformance.production_release_status()'s raw internal
            # code, e.g. "UAT_PASS_NO_RELEASE") straight into a customer
            # report Note column. The backend value and the calculation
            # producing it are unchanged - only how it's displayed here.
            "Note": customer_presentation.customer_facing_release_note(row["production_release"]) or "—",
        })
        if row["achieved"] == "Yes":
            achieved_count += 1
        elif row["achieved"] == "No":
            not_achieved_count += 1
        if row["production_release"]:
            no_release_specs.append(f"{row['property_name']} ({_context_text(row)})")
        # % deviation only has a clean single-target meaning for a plain
        # <=/>=/= operator with a numeric target_value - a "between" spec
        # (lower_limit/upper_limit, no single target_value) is skipped
        # rather than guessed at, same abstain-don't-guess convention as
        # the rest of this WP4 batch.
        if (
            row["avg_actual"] is not None
            and row["target_value"] not in (None, 0)
            and (row["target_operator"] or "<=") != "between"
        ):
            deviation_categories.append(row["property_name"])
            deviation_values.append(
                round(100 * (row["avg_actual"] - row["target_value"]) / row["target_value"], 1)
            )

    correlation_rows = []
    correlation_categories, correlation_values = [], []
    top_correlation_line = None
    if actual_ranked is not None and not actual_ranked.empty:
        for _, row in actual_ranked.iterrows():
            correlation_rows.append({
                "Raw material": row["raw_material_name"],
                "Runs compared": int(row["n_runs"]),
                "Correlation with outcome": round(row["correlation"], 3),
            })
            correlation_categories.append(row["raw_material_name"])
            correlation_values.append(round(row["correlation"], 3))
        top = actual_ranked.iloc[0]
        top_correlation_line = (
            f"Strongest association for {corr_property}: {top['raw_material_name']} "
            f"(correlation {top['correlation']:+.3f} across {int(top['n_runs'])} production runs' "
            "metered lot consumption) - a lead to investigate on the floor, not a confirmed cause."
        )

    conclusions = []
    if expectation_rows:
        conclusions.append(
            f"{achieved_count} of {achieved_count + not_achieved_count} tracked specifications are "
            f"achieved under the current recipe ({current_version.version_label}); "
            f"{not_achieved_count} are not."
        )
    else:
        conclusions.append(
            f"No specifications recorded for this grade, or no quality test results recorded yet "
            f"under the current recipe ({current_version.version_label}) to judge against them."
        )
    if no_release_specs:
        conclusions.append(
            f"{len(no_release_specs)} passing specification(s) are still under internal review and "
            f"not yet cleared for production release: {', '.join(no_release_specs)}."
        )
    if top_correlation_line:
        conclusions.append(top_correlation_line)
    else:
        conclusions.append(
            f"Not enough metered raw-material lot use data paired with {corr_property} results yet "
            "to identify a leading ingredient correlation."
        )
    if cost_per_kg is not None:
        coverage_note = f" ({cost_coverage_pct:.0f}% cost coverage)" if cost_coverage_pct is not None else ""
        conclusions.append(f"Current formulation cost: {cost_per_kg:.2f} USD per kg{coverage_note}.")
    else:
        conclusions.append("No cost data recorded for the current formulation yet.")

    return {
        "grade_name": grade.grade_name,
        "plant_name": plant_name or "—",
        "version_label": current_version.version_label,
        "version_status": current_version.approval_status,
        # Always False: the rigid achievement summary is computed straight
        # from ProductionRun rows only (see wp3_conformance.
        # compute_grade_achievement_summary), never affected by the page's
        # "Include lab trial data" toggle the way the flexible branch's
        # expectation_summary is - so this reflects what actually fed this
        # report, not whatever the toggle happened to be set to.
        "include_trials": False,
        "component_count": len(current_version.components or []),
        "cost_per_kg": cost_per_kg,
        "cost_coverage_pct": cost_coverage_pct,
        "expectation_rows": expectation_rows,
        "deviation_categories": deviation_categories,
        "deviation_values": deviation_values,
        "corr_property": corr_property,
        "correlation_rows": correlation_rows,
        "correlation_categories": correlation_categories,
        "correlation_values": correlation_values,
        "conclusions": conclusions,
        "generated_at": dt.datetime.utcnow(),
    }


def render_recipe_optimization_report_pdf(data):
    def build(story):
        _title_block(
            story, "Recipe Optimization Report",
            f"{data['grade_name']} · Plant: {data['plant_name']} · Recipe "
            f"{data['version_label']} ({data['version_status']})",
        )
        story.append(Paragraph("Context", STYLES["Heading2"]))
        story.append(_key_value_table([
            ("Product grade", data["grade_name"]), ("Plant", data["plant_name"]),
            ("Recipe version", data["version_label"]), ("Status", data["version_status"]),
            ("Ingredients", data["component_count"]),
            ("Cost per kg (USD)", f"{data['cost_per_kg']:.2f}" if data["cost_per_kg"] is not None else "—"),
            (
                "Cost coverage",
                f"{data['cost_coverage_pct']:.0f}%" if data["cost_coverage_pct"] is not None else "—",
            ),
            (
                "Lab trial data included",
                "Yes" if data["include_trials"] else "No (production runs only)",
            ),
        ]))

        story.append(Paragraph("Analysis", STYLES["Heading2"]))
        _section(story, "Does the current recipe meet target?", data["expectation_rows"])
        _bar_chart(
            story, "Deviation from target, by property (%)",
            data["deviation_categories"], data["deviation_values"],
            note="Positive = above target, negative = below target, as a % of target.",
            zero_floor=False,
        )
        _section(
            story, f"Ingredient-dosage correlation with {data['corr_property']}",
            data["correlation_rows"],
        )
        _bar_chart(
            story, f"Correlation with {data['corr_property']}, by raw material",
            data["correlation_categories"], data["correlation_values"],
            zero_floor=False,
        )

        story.append(Paragraph("Conclusions", STYLES["Heading2"]))
        for line in data["conclusions"]:
            story.append(_p(f"• {line}"))
    return _pdf_bytes(build)


def render_recipe_optimization_report_docx(data):
    doc = Document()
    _docx_report_header(
        doc, "Recipe Optimization Report",
        f"{data['grade_name']} · Plant: {data['plant_name']} · Recipe "
        f"{data['version_label']} ({data['version_status']})",
    )
    _docx_heading(doc, "Context", size=15)
    _docx_kv_table(doc, [
        ("Product grade", data["grade_name"]), ("Plant", data["plant_name"]),
        ("Recipe version", data["version_label"]), ("Status", data["version_status"]),
        ("Ingredients", data["component_count"]),
        ("Cost per kg (USD)", f"{data['cost_per_kg']:.2f}" if data["cost_per_kg"] is not None else "—"),
        (
            "Cost coverage",
            f"{data['cost_coverage_pct']:.0f}%" if data["cost_coverage_pct"] is not None else "—",
        ),
        (
            "Lab trial data included",
            "Yes" if data["include_trials"] else "No (production runs only)",
        ),
    ])

    _docx_heading(doc, "Analysis", size=15)
    _docx_section(doc, "Does the current recipe meet target?", data["expectation_rows"])
    _docx_bar_chart(
        doc, "Deviation from target, by property (%)",
        data["deviation_categories"], data["deviation_values"],
        note="Positive = above target, negative = below target, as a % of target.",
        zero_floor=False,
    )
    _docx_section(
        doc, f"Ingredient-dosage correlation with {data['corr_property']}",
        data["correlation_rows"],
    )
    _docx_bar_chart(
        doc, f"Correlation with {data['corr_property']}, by raw material",
        data["correlation_categories"], data["correlation_values"],
        zero_floor=False,
    )

    _docx_heading(doc, "Conclusions", size=15)
    for line in data["conclusions"]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(line)
    return _docx_bytes(doc)


# ---------------------------------------------------------------------------
# 11. Trend Analysis Report (Context / Analysis / Conclusions)
#
# Added 2026-08-04, same batch and placement logic as the Recipe
# Optimization Report above - lives on pages/16_Trend_Analysis.py itself
# (product grade/family, property, recipe/machine filters: a comprehensive
# multi-field selection, not a single dropdown choice). This is the page's
# own deterministic SPC results (control chart, capability, CUSUM, trend
# test - all computed by analytics.py), never the PI3 interpretation
# further down (which has its own separate Word download).
# build_trend_analysis_report_data() takes the exact chart_result/
# capability/cusum/trend/change_rows objects the page has already computed
# - it never re-runs the SPC math itself.
# ---------------------------------------------------------------------------

def build_trend_analysis_report_data(
    session, unit, property_name, series, pooling_grades,
    chart_result, capability, cusum, trend, change_rows, include_trials,
):
    """unit: the dict returned by helpers.analysis_unit_picker() (label/
    mode/member_grade_names/...). series: the DataFrame from
    analytics.property_run_series(), already filtered by the page (recipe
    version / machine). chart_result/capability/cusum/trend: the exact
    deterministic SPC results from analytics.control_chart_analysis()/
    capability_analysis()/cusum_analysis()/trend_test() the page already
    computed. change_rows: the machine-change + quality-issue timeline the
    page already assembled."""
    subject_desc = (
        f"product grade {unit['label']}" if unit["mode"] == "grade"
        else f"product family {unit['label']} (pooling grades: {', '.join(unit['member_grade_names'])})"
    )

    control_categories, control_series = [], []
    control_summary_line = "Not enough results yet for a control chart."
    control_flag_rows = []
    if chart_result and chart_result.get("ready"):
        chart_df = chart_result["chart_df"]
        control_categories = [str(d) for d in chart_df["tested_at"]]
        control_series = [
            ("Actual", chart_df["actual_value"].tolist()),
            ("Center", chart_df["center_line"].tolist()),
            ("Upper limit", chart_df["ucl"].tolist()),
            ("Lower limit", chart_df["lcl"].tolist()),
        ]
        if chart_result["in_control"]:
            control_summary_line = "In control - no unusual patterns found across these runs."
        else:
            control_summary_line = f"{len(chart_result['flags'])} unusual pattern(s) found."
            control_flag_rows = [
                {
                    "Pattern": f["rule"], "First seen": str(f["first_tested_at"]),
                    "Points matching": f["points_matching"],
                }
                for f in chart_result["flags"]
            ]

    capability_line = "Not enough data for a margin-to-spec check."
    capability_kv = []
    if capability is not None:
        cpk = capability["cpk"]
        if cpk >= 1.33:
            read = "comfortable margin to spec"
        elif cpk >= 1.0:
            read = "tight - some results will likely fall outside spec"
        else:
            read = "not enough margin - this process routinely produces results outside spec"
        capability_line = f"Cpk {cpk:.2f} - {read}."
        # WP6-S09 closure (2026-08-09, UAT-016): capability_analysis() now
        # returns a genuinely one-sided result (cpl/lsl or cpu/usl is None,
        # not a fabricated opposite limit) whenever the real spec only has
        # one real limit - see analytics.capability_analysis's docstring.
        # Formatting capability['cpl']/['lsl'] unconditionally as before
        # crashed with "unsupported format string passed to NoneType" the
        # first time a one-sided spec (Trend Analysis's real Thermal
        # conductivity "<=" case) reached this report. Branch the same way
        # the on-screen display already does.
        if capability.get("one_sided"):
            op = capability.get("operator")
            if op == "<=":
                capability_kv = [
                    ("Cpk = Cpu (margin to upper limit)", f"{cpk:.2f}"),
                    ("Upper spec limit", f"{capability['usl']:.3g}"),
                ]
            else:
                capability_kv = [
                    ("Cpk = Cpl (margin to lower limit)", f"{cpk:.2f}"),
                    ("Lower spec limit", f"{capability['lsl']:.3g}"),
                ]
            capability_kv.append((
                "Note",
                "One-sided specification - there is no real opposite limit, so none is shown or invented.",
            ))
        else:
            capability_kv = [
                ("Cpk (overall margin)", f"{cpk:.2f}"),
                ("Cpu (margin to upper limit)", f"{capability['cpu']:.2f}"),
                ("Cpl (margin to lower limit)", f"{capability['cpl']:.2f}"),
                ("Spec range", f"{capability['lsl']:.3g} - {capability['usl']:.3g}"),
            ]

    cusum_categories, cusum_series = [], []
    cusum_line = "Not enough data for a slow-drift check."
    if cusum is not None:
        cusum_df = cusum["chart_df"].copy()
        cusum_df["upper_limit"] = cusum["h"]
        cusum_df["lower_limit"] = -cusum["h"]
        cusum_categories = [str(d) for d in cusum_df["tested_at"]]
        cusum_series = [
            ("CUSUM +", cusum_df["cusum_positive"].tolist()),
            ("CUSUM -", cusum_df["cusum_negative"].tolist()),
            ("Upper limit", cusum_df["upper_limit"].tolist()),
            ("Lower limit", cusum_df["lower_limit"].tolist()),
        ]
        if cusum["breach_index"] is None:
            cusum_line = "No slow drift detected - results have stayed close to target over time."
        else:
            cusum_line = (
                f"A slow {cusum['breach_direction']} drift has been building up, first becoming "
                f"clear at {cusum['breach_tested_at']}."
            )

    trend_line = "Not enough data to test whether this is a real trend."
    if trend is not None:
        if trend["significant"]:
            trend_line = (
                f"Yes - a real, sustained {trend['direction']} trend: the fitted straight line "
                f"changes by {trend['slope_per_run']:+.4g} per run, explains "
                f"{trend['r_squared'] * 100:.0f}% of the run-to-run variation (R²={trend['r_squared']:.2f}), "
                f"with only a {trend['p_value'] * 100:.2g}% chance this slope is random noise "
                f"(p={trend['p_value']:.4f}), across {trend['n']} runs."
            )
        else:
            trend_line = (
                f"No - the apparent {trend['direction']} movement across {trend['n']} runs looks "
                f"like normal run-to-run variation: the fitted straight line explains only "
                f"{trend['r_squared'] * 100:.0f}% of the run-to-run variation (R²={trend['r_squared']:.2f}), "
                f"and there's a {trend['p_value'] * 100:.2g}% chance this slope is random noise "
                f"(p={trend['p_value']:.4f})."
            )
        if trend["mk_significant"] != trend["significant"]:
            trend_line += (
                f" Note: a non-parametric Mann-Kendall cross-check disagrees (tau={trend['mk_tau']:+.2f}, "
                f"p={trend['mk_p_value']:.4f}) - the drift may not be a straight line."
            )

    change_display_rows = [
        {"Date": str(r["Date"]), "Run ID": r["Run ID"], "Change": r["Change"]}
        for r in change_rows
    ]

    conclusions = [control_summary_line, capability_line, cusum_line, trend_line]
    flagged = (
        (chart_result and chart_result.get("ready") and not chart_result.get("in_control", True))
        or (cusum is not None and cusum.get("breach_index") is not None)
        or (trend is not None and trend.get("significant"))
    )
    if change_rows and flagged:
        conclusions.append(
            f"{len(change_rows)} machine change(s)/quality issue(s) recorded on this timeline - "
            "cross-reference their dates against the flags above before drawing conclusions."
        )

    return {
        "property_name": property_name,
        "subject_desc": subject_desc,
        "pooling_grades": pooling_grades,
        "include_trials": include_trials,
        "n_results": len(series),
        "date_range": (
            f"{series['tested_at'].min()} to {series['tested_at'].max()}" if len(series) else "—"
        ),
        "control_categories": control_categories,
        "control_series": control_series,
        "control_flag_rows": control_flag_rows,
        "capability_kv": capability_kv,
        "cusum_categories": cusum_categories,
        "cusum_series": cusum_series,
        "change_rows": change_display_rows,
        "conclusions": conclusions,
        "generated_at": dt.datetime.utcnow(),
    }


def render_trend_analysis_report_pdf(data):
    def build(story):
        _title_block(
            story, "Trend Analysis Report",
            f"{data['property_name']} · {data['subject_desc']}",
        )
        story.append(Paragraph("Context", STYLES["Heading2"]))
        story.append(_key_value_table([
            ("Property", data["property_name"]), ("Subject", data["subject_desc"]),
            ("Results analyzed", data["n_results"]), ("Date range", data["date_range"]),
            ("Lab trial data included", "Yes" if data["include_trials"] else "No"),
            ("Pooled by % of target", "Yes" if data["pooling_grades"] else "No"),
        ]))

        story.append(Paragraph("Analysis", STYLES["Heading2"]))
        _line_chart(
            story, "Control chart (sudden changes)",
            data["control_categories"], data["control_series"],
            note="Actual result vs. center line and control limits." if data["control_categories"] else None,
        )
        if data["control_flag_rows"]:
            _section(story, "Unusual patterns flagged", data["control_flag_rows"])
        if data["capability_kv"]:
            story.append(Spacer(1, 8))
            story.append(Paragraph("Margin to spec", STYLES["Heading3"]))
            story.append(_key_value_table(data["capability_kv"]))
        _line_chart(
            story, "CUSUM (slow drift)",
            data["cusum_categories"], data["cusum_series"],
            note=(
                "Cumulative sum of small deviations from target over time."
                if data["cusum_categories"] else None
            ),
        )
        _section(story, "What else changed on this timeline", data["change_rows"])

        story.append(Paragraph("Conclusions", STYLES["Heading2"]))
        for line in data["conclusions"]:
            story.append(_p(f"• {line}"))
    return _pdf_bytes(build)


def render_trend_analysis_report_docx(data):
    doc = Document()
    _docx_report_header(doc, "Trend Analysis Report", f"{data['property_name']} · {data['subject_desc']}")

    _docx_heading(doc, "Context", size=15)
    _docx_kv_table(doc, [
        ("Property", data["property_name"]), ("Subject", data["subject_desc"]),
        ("Results analyzed", data["n_results"]), ("Date range", data["date_range"]),
        ("Lab trial data included", "Yes" if data["include_trials"] else "No"),
        ("Pooled by % of target", "Yes" if data["pooling_grades"] else "No"),
    ])

    _docx_heading(doc, "Analysis", size=15)
    _docx_line_chart(
        doc, "Control chart (sudden changes)",
        data["control_categories"], data["control_series"],
        note="Actual result vs. center line and control limits." if data["control_categories"] else None,
    )
    if data["control_flag_rows"]:
        _docx_section(doc, "Unusual patterns flagged", data["control_flag_rows"])
    if data["capability_kv"]:
        _docx_heading(doc, "Margin to spec", size=12, color=_HTC_GREY, space_before=10)
        _docx_kv_table(doc, data["capability_kv"])
    _docx_line_chart(
        doc, "CUSUM (slow drift)",
        data["cusum_categories"], data["cusum_series"],
        note=(
            "Cumulative sum of small deviations from target over time."
            if data["cusum_categories"] else None
        ),
    )
    _docx_section(doc, "What else changed on this timeline", data["change_rows"])

    _docx_heading(doc, "Conclusions", size=15)
    for line in data["conclusions"]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(line)
    return _docx_bytes(doc)


# ---------------------------------------------------------------------------
# 12. Process-Property Correlation Report (Context / Analysis / Conclusions)
#
# Added 2026-08-04, same batch and placement logic as the two reports
# above - lives on pages/17_Process_Property_Correlation.py itself (foam
# grade/family plus property: a comprehensive multi-field selection). This
# is the page's own ranked correlation table (analytics.rank_setting_
# correlations), never the PI3 synthesis further down (which has its own
# separate Word download). build_correlation_report_data() takes the exact
# `ranked` DataFrame the page has already computed - never re-derived.
# ---------------------------------------------------------------------------

def build_correlation_report_data(session, unit, property_name, ranked, pooling_grades):
    """unit: the dict from helpers.analysis_unit_picker(). ranked: the
    DataFrame from analytics.rank_setting_correlations() the page has
    already computed for property_name (columns label/n/correlation/
    field)."""
    subject_desc = (
        f"product grade {unit['label']}" if unit["mode"] == "grade"
        else f"product family {unit['label']} (pooling grades: {', '.join(unit['member_grade_names'])})"
    )

    ranked_with_data = ranked.dropna(subset=["correlation"]) if ranked is not None else pd.DataFrame()
    ranking_rows = []
    categories, values = [], []
    if ranked is not None:
        for _, row in ranked.iterrows():
            has_data = pd.notna(row["correlation"])
            ranking_rows.append({
                "Process setting": row["label"], "Runs compared": int(row["n"]),
                "Strength (|r|)": round(abs(row["correlation"]), 3) if has_data else "—",
                "Correlation": round(row["correlation"], 3) if has_data else "—",
            })
            if has_data:
                categories.append(row["label"])
                values.append(round(row["correlation"], 3))

    top_line = "No process setting has enough paired data yet to correlate against this property."
    if not ranked_with_data.empty:
        top = ranked_with_data.iloc[0]
        direction = "positive" if top["correlation"] > 0 else "negative"
        top_line = (
            f"Strongest association: {top['label']} ({direction}, r={top['correlation']:.2f}) "
            f"across {int(top['n'])} runs - a lead to investigate, not a confirmed cause."
        )

    conclusions = [top_line]
    if ranked is not None and len(ranked):
        conclusions.append(
            f"{len(ranked_with_data)} of {len(ranked)} tracked process settings have enough paired "
            "data to compute a correlation with this property."
        )

    return {
        "property_name": property_name,
        "subject_desc": subject_desc,
        "pooling_grades": pooling_grades,
        "ranking_rows": ranking_rows,
        "correlation_categories": categories,
        "correlation_values": values,
        "conclusions": conclusions,
        "generated_at": dt.datetime.utcnow(),
    }


def render_correlation_report_pdf(data):
    def build(story):
        _title_block(
            story, "Process-Property Correlation Report",
            f"{data['property_name']} · {data['subject_desc']}",
        )
        story.append(Paragraph("Context", STYLES["Heading2"]))
        story.append(_key_value_table([
            ("Property", data["property_name"]), ("Subject", data["subject_desc"]),
            ("Pooled by % of target", "Yes" if data["pooling_grades"] else "No"),
            ("Process settings ranked", len(data["ranking_rows"])),
        ]))

        story.append(Paragraph("Analysis", STYLES["Heading2"]))
        _section(story, "All settings, ranked by correlation strength", data["ranking_rows"])
        _bar_chart(
            story, f"Correlation with {data['property_name']}, by process setting",
            data["correlation_categories"], data["correlation_values"],
            note="Positive/negative shows direction; size shows strength of association.",
            zero_floor=False,
        )

        story.append(Paragraph("Conclusions", STYLES["Heading2"]))
        for line in data["conclusions"]:
            story.append(_p(f"• {line}"))
    return _pdf_bytes(build)


def render_correlation_report_docx(data):
    doc = Document()
    _docx_report_header(
        doc, "Process-Property Correlation Report", f"{data['property_name']} · {data['subject_desc']}",
    )
    _docx_heading(doc, "Context", size=15)
    _docx_kv_table(doc, [
        ("Property", data["property_name"]), ("Subject", data["subject_desc"]),
        ("Pooled by % of target", "Yes" if data["pooling_grades"] else "No"),
        ("Process settings ranked", len(data["ranking_rows"])),
    ])

    _docx_heading(doc, "Analysis", size=15)
    _docx_section(doc, "All settings, ranked by correlation strength", data["ranking_rows"])
    _docx_bar_chart(
        doc, f"Correlation with {data['property_name']}, by process setting",
        data["correlation_categories"], data["correlation_values"],
        note="Positive/negative shows direction; size shows strength of association.",
        zero_floor=False,
    )

    _docx_heading(doc, "Conclusions", size=15)
    for line in data["conclusions"]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(line)
    return _docx_bytes(doc)


# ---------------------------------------------------------------------------
# 13. Root-Cause Comparison Report (Context / Analysis / Conclusions)
#
# Added 2026-08-04, part of the same Industrial Intelligence reports batch
# as the three reports above. Lives on pages/18_Root_Cause_Assistant.py
# itself, per the user-approved plan grouping all 5 analysis pages (15-19)
# together as "pages with PI3" that get their deterministic report on the
# page rather than the Report page. This is the page's own deterministic
# run-vs-prior-run diff (recipe/machine/Finalized-phase settings), never
# the PI3 hypothesis further down (which has its own separate Word
# download). build_root_cause_report_data() takes the exact `changes`/
# `setting_shifts` the page has already computed - never re-derived.
# ---------------------------------------------------------------------------

def environment_outcome_context_rows(definitions_by_field, current_values, prior_values):
    """WP7 Phase 4 targeted completion, Item 2 (2026-08-14) - per Charlie's
    Closeout Review Return to JC: 'Environment/Outcome as separate context
    sections (excluded from controllable-setting ranking but visible)'.
    The page's own 'What was different' loop stays scoped to category ==
    'Process Setting' only (the controllable-lever comparison that drives
    both the on-screen change list and the PI3 hypothesis prompt) -
    Environment/Outcome recorded values must never be folded into that
    list or its ranking, but must still be visible to the reviewer as
    context. Reuses the exact values_by_run/definitions_by_field the page
    already computed via analytics.production_run_parameter_dataframe()
    for the Process Setting diff - never re-queried."""
    buckets = {"Environment": [], "Outcome": []}
    for field_key, meta in sorted(definitions_by_field.items(), key=lambda kv: kv[1]["label"] or kv[0]):
        category = meta["parameter_category"]
        if category not in buckets:
            continue
        buckets[category].append({
            "Parameter": meta["label"],
            "Prior (Actual)": prior_values.get(field_key),
            "Current (Actual)": current_values.get(field_key),
            "UOM": meta["unit_symbol"] or "—",
        })
    return buckets


def current_run_process_setting_rows(session, run_id):
    """WP7 Phase 4 Root Cause final targeted completion (2026-08-15, per
    Charlie's Corrected Closeout Review Return to JC): "Add a dedicated
    current-run Process Setting Planned-versus-Actual context to Root
    Cause... For each eligible Process Setting, provide Parameter,
    Planned, Actual, numeric Delta where applicable, and canonical UOM."

    Reuses _process_parameter_report_rows(session, run_id)'s "Process
    Setting" bucket rather than re-deriving - the same definition-driven,
    per-run reader Batch Release already uses (Item 1), so this gets the
    same Parameter/Category/Planned/Actual/Delta/UOM/Limit/Conformance
    shape for free (Limit/Conformance are additive - Charlie's minimum
    required columns are all present; "an equivalent presentation is
    acceptable").

    This is a CURRENT-RUN view (this run's own Planned vs Actual, from
    analytics.production_run_process_parameters), deliberately separate
    from the page's existing current-vs-prior-run shift comparison (which
    reads analytics.production_run_parameter_dataframe's multi-run form
    and only ever carries Actual - Planned never substitutes for a
    missing Actual, so that comparison alone cannot show what THIS run's
    own Planned target was). Charlie's return requires both views kept:
    this one for Planned/Actual/Delta context, the existing one for the
    run-over-run shift."""
    return _process_parameter_report_rows(session, run_id)["Process Setting"]


def _fmt_value(value):
    """Shared formatting for a recorded fact value in prose text (PI3
    prompt lines): None reads as 'not recorded' - NEVER as 0 or blank -
    preserving the NULL-vs-recorded-zero distinction Charlie's Phase 4
    semantics require everywhere else in this app. A recorded zero prints
    as '0', not as 'not recorded'."""
    if value is None:
        return "not recorded"
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, float):
        return f"{value:g}"
    return str(value)


def format_root_cause_facts_for_pi3(
    investigation_facts, env_outcome_rows, current_setting_rows,
):
    """WP7 Phase 4 Root Cause final targeted completion (2026-08-15, per
    Charlie's Corrected Closeout Review Return to JC, Final Targeted
    Completion item 3): "Pass the recorded Root Cause fact content into
    the PI3 investigation payload. Include the actual Environment /
    Outcome values, material metering or usage values, Production Event
    details, relevant QC facts, and current-run Process Setting Planned /
    Actual / Delta context. Counts may remain as summary metadata while
    the fact values carry the investigation context."

    Pure formatting function (no I/O) so it is independently unit-
    testable at the payload level, per Charlie's item 4 ("Add a payload-
    level assertion proving seeded ... fact values reach the PI3
    hypothesis input"), without needing to mock the OpenAI call or drive
    the page through AppTest. Takes the page's own already-computed
    investigation_facts (root_cause_investigation_facts), env_outcome_rows
    (environment_outcome_context_rows), and current_setting_rows
    (current_run_process_setting_rows) dicts/lists - never re-queries.

    Returns a single formatted text block, one bullet line per recorded
    fact with its real value (never just a count), grouped under the same
    section headings shown on screen. A section with no rows contributes
    an explicit 'None recorded' line rather than being silently omitted,
    so PI3 never has to guess whether an empty section means 'not
    checked' versus 'checked, nothing there'."""
    lines = []

    lines.append("Current-run Process Setting (Planned / Actual / Delta):")
    if current_setting_rows:
        for row in current_setting_rows:
            uom = row["UOM"] if row["UOM"] and row["UOM"] != "—" else ""
            delta_text = _fmt_value(row["Delta"]) + (f" {uom}" if uom and row["Delta"] is not None else "")
            lines.append(
                f"- {row['Parameter']}: Planned {_fmt_value(row['Planned'])}{f' {uom}' if uom and row['Planned'] is not None else ''}, "
                f"Actual {_fmt_value(row['Actual'])}{f' {uom}' if uom and row['Actual'] is not None else ''}, "
                f"Delta {delta_text}"
            )
    else:
        lines.append("- None recorded.")

    env_rows = env_outcome_rows.get("Environment", []) if env_outcome_rows else []
    out_rows = env_outcome_rows.get("Outcome", []) if env_outcome_rows else []
    lines.append("Environment context (recorded values, both runs):")
    if env_rows:
        for row in env_rows:
            uom = row["UOM"] if row["UOM"] and row["UOM"] != "—" else ""
            lines.append(
                f"- {row['Parameter']}: prior {_fmt_value(row['Prior (Actual)'])}, "
                f"current {_fmt_value(row['Current (Actual)'])}{f' {uom}' if uom else ''}"
            )
    else:
        lines.append("- None recorded.")
    lines.append("Outcome context (recorded values, both runs):")
    if out_rows:
        for row in out_rows:
            uom = row["UOM"] if row["UOM"] and row["UOM"] != "—" else ""
            lines.append(
                f"- {row['Parameter']}: prior {_fmt_value(row['Prior (Actual)'])}, "
                f"current {_fmt_value(row['Current (Actual)'])}{f' {uom}' if uom else ''}"
            )
    else:
        lines.append("- None recorded.")

    lines.append("Material usage / metering:")
    rows = investigation_facts.get("material_usage_rows", [])
    if rows:
        for r in rows:
            lines.append(
                f"- {r['Stream']}: total delivered {_fmt_value(r['Total delivered'])} {r['Unit']}, "
                f"flow {_fmt_value(r['Flow'])}, pump speed {_fmt_value(r['Pump speed'])}, "
                f"temperature {_fmt_value(r['Temperature (°C)'])}°C, pressure {_fmt_value(r['Pressure (bar)'])} bar, "
                f"calibration {r['Calibration']}"
            )
    else:
        lines.append("- None recorded.")

    lines.append("Production events:")
    rows = investigation_facts.get("production_event_rows", [])
    if rows:
        for r in rows:
            lines.append(f"- {r['Time']} [{r['Type']}, severity {r['Severity']}]: {r['Description']}")
    else:
        lines.append("- None recorded.")

    lines.append("QC context — other quality test results on this run:")
    rows = investigation_facts.get("qc_result_rows", [])
    if rows:
        for r in rows:
            # Rigid grades key results by "Property"/"Target"/"Actual"/"Unit"/"Pass/Fail";
            # non-rigid grades use the same keys (root_cause_investigation_facts
            # normalizes both shapes to this column set) - read defensively either way.
            prop = r.get("Property", "—")
            target = _fmt_value(r.get("Target"))
            actual = _fmt_value(r.get("Actual"))
            unit = r.get("Unit", "") or ""
            verdict = r.get("Pass/Fail", "—")
            lines.append(f"- {prop}: target {target} {unit}, actual {actual} {unit} ({verdict})")
    else:
        lines.append("- None recorded.")

    lines.append("QC context — quality issues logged on this run (including the flagged one):")
    rows = investigation_facts.get("qc_issue_rows", [])
    if rows:
        for r in rows:
            lines.append(
                f"- {r['Issue type']} (severity {r['Severity']}, frequency {r['Frequency']}, "
                f"confidence {r['Confidence']})"
            )
    else:
        lines.append("- None recorded.")

    return "\n".join(lines)


def root_cause_investigation_facts(session, run):
    """WP7 Phase 4 targeted completion, Item 2 (2026-08-14) - per Charlie's
    Closeout Review Return to JC: Root-Cause Assistant must include
    run-linked material usage/metering, Production Events, and QC context
    'as investigation facts separated from inferred hypotheses' - these
    are recorded facts about the flagged run itself (never re-derived,
    never a PI3 output), rendered in their own section above/separate from
    the PI3 hypothesis further down the page.

    Material Metering reads ComponentStreamReading.production_run_id only
    (the 2026-08-14 Item 1 pattern - no ProductionPhase dependency).
    QC context reuses the same rigid-vs-non-rigid conformance resolution
    build_batch_release_record_data() uses (wp3_conformance for rigid
    grades, quality_standards.compute_pass_fail otherwise) so the Pass/
    Fail shown here always matches what Batch Release would show for the
    same run."""
    readings = (
        session.query(ComponentStreamReading)
        .filter(ComponentStreamReading.production_run_id == run.id).all()
    )
    material_usage_rows = [
        {
            "Stream": rd.stream_name, "Flow": rd.flow, "Unit": rd.flow_unit or "",
            "Pump speed": rd.pump_speed, "Total delivered": rd.flow_total_qty,
            "Temperature (°C)": rd.temperature_c, "Pressure (bar)": rd.pressure_bar,
            "Calibration": rd.calibration_status or "—",
        }
        for rd in readings
    ]

    production_event_rows = [
        {
            "Time": e.event_ts, "Type": e.event_type, "Severity": e.severity or "—",
            "Description": e.description or "—",
        }
        for e in session.query(ProductionEvent)
        .filter(ProductionEvent.production_run_id == run.id)
        .order_by(ProductionEvent.event_ts).all()
    ]

    grade = run.foam_grade
    results = (
        session.query(PhysicalPropertyResult)
        .filter(PhysicalPropertyResult.production_run_id == run.id).all()
    )
    if _is_rigid_grade(grade):
        conformance_rows = wp3_conformance.compute_conformance_report(
            session, grade.id, production_run_id=run.id
        )
        results_by_id = {r.id: r for r in results}
        qc_result_rows = _conformance_rows_for_display(session, results_by_id, conformance_rows)
    else:
        qc_result_rows = [
            {
                "Property": r.property_name, "Target": r.target_value, "Actual": r.actual_value,
                "Unit": r.unit or "",
                "Pass/Fail": compute_pass_fail(r.property_name, r.target_value, r.actual_value) or "—",
            }
            for r in results
        ]

    qc_issue_rows = [
        {
            "Issue type": o.observation_type, "Severity": o.severity or "—",
            "Frequency": o.frequency or "—", "Confidence": o.confidence_level or "—",
        }
        for o in session.query(QualityObservation)
        .filter(QualityObservation.production_run_id == run.id).all()
    ]

    return {
        "material_usage_rows": material_usage_rows,
        "production_event_rows": production_event_rows,
        "qc_result_rows": qc_result_rows,
        "qc_issue_rows": qc_issue_rows,
    }


def build_root_cause_report_data(
    session, obs, run, grade, prior, changes, setting_shifts,
    env_outcome_rows=None, investigation_facts=None, current_setting_rows=None,
):
    """obs: QualityObservation. run: its ProductionRun. grade: run.foam_grade.
    prior: the prior-run settings row (a pandas Series from analytics.
    run_settings_dataframe()) the page already selected as the comparison
    baseline. changes: the page's own list of formatted diff strings
    ("What was different"). setting_shifts: a parallel list of dicts
    ({"label", "pct_change"}) for the numeric Finalized-phase shifts only
    (recipe/machine changes aren't percentages, so they're text-only in
    `changes` and don't appear here) - lets the report chart the shift
    magnitudes without re-deriving them.

    WP7 Phase 4 targeted completion, Item 2 (2026-08-14): env_outcome_rows
    (from environment_outcome_context_rows()) and investigation_facts
    (from root_cause_investigation_facts()) are the page's own already-
    computed dicts - never re-derived here. Both default to None/empty so
    existing callers/tests that don't pass them still get a valid report
    (empty context/facts sections).

    WP7 Phase 4 Root Cause final targeted completion (2026-08-15, per
    Charlie's Corrected Closeout Review Return to JC): current_setting_rows
    (from current_run_process_setting_rows()) is this run's own dedicated
    Planned-vs-Actual Process Setting context, kept as a separate report
    section from the "What was different" run-vs-prior-run shift list
    above - Charlie's explicit "keep... as a separate analytical view"
    instruction. Defaults to [] so existing callers/tests unaffected."""
    shift_categories = [s["label"] for s in setting_shifts]
    shift_values = [round(s["pct_change"] * 100, 2) for s in setting_shifts]
    env_outcome_rows = env_outcome_rows or {"Environment": [], "Outcome": []}
    investigation_facts = investigation_facts or {
        "material_usage_rows": [], "production_event_rows": [],
        "qc_result_rows": [], "qc_issue_rows": [],
    }
    current_setting_rows = current_setting_rows or []

    if changes:
        conclusion_lines = [f"{len(changes)} difference(s) found versus the prior run."]
        if setting_shifts:
            largest = max(setting_shifts, key=lambda s: abs(s["pct_change"]))
            conclusion_lines.append(
                f"Largest recorded process-setting shift: {largest['label']} "
                f"({largest['pct_change']:+.2%})."
            )
    else:
        conclusion_lines = [
            "No meaningful difference found in recipe, machine, or recorded process settings "
            "between these two runs - the cause may lie outside what this app currently "
            "captures (raw material lot variation, ambient conditions, downstream handling)."
        ]
    if obs.suspected_cause:
        conclusion_lines.append(f"Logged suspected cause: {obs.suspected_cause}")

    return {
        "observation_type": obs.observation_type,
        "severity": obs.severity,
        "frequency": obs.frequency,
        "suspected_cause": obs.suspected_cause,
        "run_id": run.id,
        "run_date": str(run.run_date),
        "grade_name": grade.grade_name,
        "prior_run_id": int(prior["run_id"]),
        "prior_run_date": str(prior["run_date"]),
        "change_rows": [{"Change": c} for c in changes],
        "current_setting_rows": current_setting_rows,
        "shift_categories": shift_categories,
        "shift_values": shift_values,
        "environment_rows": env_outcome_rows["Environment"],
        "outcome_rows": env_outcome_rows["Outcome"],
        "material_usage_rows": investigation_facts["material_usage_rows"],
        "production_event_rows": investigation_facts["production_event_rows"],
        "qc_result_rows": investigation_facts["qc_result_rows"],
        "qc_issue_rows": investigation_facts["qc_issue_rows"],
        "conclusions": conclusion_lines,
        "generated_at": dt.datetime.utcnow(),
    }


def render_root_cause_report_pdf(data):
    def build(story):
        _title_block(
            story, "Root-Cause Comparison Report",
            f"{data['observation_type']} on run #{data['run_id']} ({data['grade_name']})",
        )
        story.append(Paragraph("Context", STYLES["Heading2"]))
        story.append(_key_value_table([
            ("Quality issue", data["observation_type"]), ("Severity", data["severity"]),
            ("Frequency", data["frequency"]), ("Run", f"#{data['run_id']} ({data['run_date']})"),
            ("Product grade", data["grade_name"]),
            ("Compared against", f"run #{data['prior_run_id']} ({data['prior_run_date']})"),
            ("Logged suspected cause", data["suspected_cause"] or "—"), ("", ""),
        ]))

        story.append(Paragraph("Analysis", STYLES["Heading2"]))
        # WP7 Phase 4 Root Cause final targeted completion (2026-08-15):
        # this run's own Planned-vs-Actual context, kept separate from the
        # run-vs-prior-run shift list below (Charlie's explicit "keep as a
        # separate analytical view" instruction).
        _section(story, "Current run — Process Setting (Planned vs. Actual)", data["current_setting_rows"])
        _section(story, "What was different vs. the prior run", data["change_rows"])
        _bar_chart(
            story, "Process-setting shifts (%)",
            data["shift_categories"], data["shift_values"],
            note="Only settings that shifted at least 2% between the two runs.",
            zero_floor=False,
        )
        # WP7 Phase 4 targeted completion, Item 2 (2026-08-14): context,
        # visible but never folded into "What was different" above.
        _section(story, "Environment — recorded context (both runs)", data["environment_rows"])
        _section(story, "Outcome — recorded context (both runs)", data["outcome_rows"])

        story.append(Paragraph("Investigation facts", STYLES["Heading2"]))
        story.append(_p(
            f"Recorded data for run #{data['run_id']} - facts, not inferred hypotheses."
        ))
        _section(story, "Material usage / metering", data["material_usage_rows"])
        _section(story, "Production events", data["production_event_rows"])
        _section(story, "QC context — other quality test results on this run", data["qc_result_rows"])
        _section(story, "QC context — quality issues logged on this run (including the flagged one)", data["qc_issue_rows"])

        story.append(Paragraph("Conclusions", STYLES["Heading2"]))
        for line in data["conclusions"]:
            story.append(_p(f"• {line}"))
    return _pdf_bytes(build)


def render_root_cause_report_docx(data):
    doc = Document()
    _docx_report_header(
        doc, "Root-Cause Comparison Report",
        f"{data['observation_type']} on run #{data['run_id']} ({data['grade_name']})",
    )
    _docx_heading(doc, "Context", size=15)
    _docx_kv_table(doc, [
        ("Quality issue", data["observation_type"]), ("Severity", data["severity"]),
        ("Frequency", data["frequency"]), ("Run", f"#{data['run_id']} ({data['run_date']})"),
        ("Product grade", data["grade_name"]),
        ("Compared against", f"run #{data['prior_run_id']} ({data['prior_run_date']})"),
        ("Logged suspected cause", data["suspected_cause"] or "—"),
    ])

    _docx_heading(doc, "Analysis", size=15)
    # WP7 Phase 4 Root Cause final targeted completion (2026-08-15): this
    # run's own Planned-vs-Actual context, kept separate from the
    # run-vs-prior-run shift list below (Charlie's explicit "keep as a
    # separate analytical view" instruction).
    _docx_section(doc, "Current run — Process Setting (Planned vs. Actual)", data["current_setting_rows"])
    _docx_section(doc, "What was different vs. the prior run", data["change_rows"])
    _docx_bar_chart(
        doc, "Process-setting shifts (%)",
        data["shift_categories"], data["shift_values"],
        note="Only settings that shifted at least 2% between the two runs.",
        zero_floor=False,
    )
    # WP7 Phase 4 targeted completion, Item 2 (2026-08-14): context, visible
    # but never folded into "What was different" above.
    _docx_section(doc, "Environment — recorded context (both runs)", data["environment_rows"])
    _docx_section(doc, "Outcome — recorded context (both runs)", data["outcome_rows"])

    _docx_heading(doc, "Investigation facts", size=15)
    doc.add_paragraph(f"Recorded data for run #{data['run_id']} - facts, not inferred hypotheses.")
    _docx_section(doc, "Material usage / metering", data["material_usage_rows"])
    _docx_section(doc, "Production events", data["production_event_rows"])
    _docx_section(doc, "QC context — other quality test results on this run", data["qc_result_rows"])
    _docx_section(doc, "QC context — quality issues logged on this run (including the flagged one)", data["qc_issue_rows"])

    _docx_heading(doc, "Conclusions", size=15)
    for line in data["conclusions"]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(line)
    return _docx_bytes(doc)


# ---------------------------------------------------------------------------
# 14. Process Parameter Optimization Report (Context / Analysis / Conclusions)
# (CR-01, 2026-08-10: report title renamed from "Machine Settings
# Optimization Report" - function name/backend unchanged, label-only rename)
#
# Added 2026-08-04, same batch and placement logic as the Process-Property
# Correlation Report above (structurally the closest sibling - both rank
# every process setting against a chosen property) - lives on
# pages/19_Machine_Settings_Optimization.py itself. This is the page's own
# ranked table (analytics.rank_setting_optimization), never the PI3
# synthesis further down (which has its own separate Word download).
# build_machine_settings_report_data() takes the exact `ranked` DataFrame
# the page has already computed - never re-derived.
# ---------------------------------------------------------------------------

def build_machine_settings_report_data(session, unit, property_name, ranked, pooling_grades):
    """unit: the dict from helpers.analysis_unit_picker(). ranked: the
    DataFrame from analytics.rank_setting_optimization() the page has
    already computed for property_name (columns label/n/best_range/
    best_range_setting/best_range_avg_dev_pct/spread_pct/field)."""
    subject_desc = (
        f"product grade {unit['label']}" if unit["mode"] == "grade"
        else f"product family {unit['label']} (pooling grades: {', '.join(unit['member_grade_names'])})"
    )

    ranked_with_data = ranked.dropna(subset=["spread_pct"]) if ranked is not None else pd.DataFrame()
    ranking_rows = []
    categories, values = [], []
    if ranked is not None:
        for _, row in ranked.iterrows():
            has_data = pd.notna(row["spread_pct"])
            ranking_rows.append({
                "Process setting": row["label"], "Runs compared": int(row["n"]),
                "Best range": row["best_range"] if has_data else "—",
                "Best range (values)": row["best_range_setting"] if has_data else "—",
                "Best range avg deviation %": round(row["best_range_avg_dev_pct"], 1) if has_data else "—",
                "Gap vs worst range (pts)": round(row["spread_pct"], 1) if has_data else "—",
            })
            if has_data:
                categories.append(row["label"])
                values.append(round(row["spread_pct"], 1))

    top_line = "No process setting has enough data yet to rank for this property."
    if not ranked_with_data.empty:
        top = ranked_with_data.iloc[0]
        top_line = (
            f"Most actionable: {top['label']}, {top['best_range']} range ({top['best_range_setting']}) "
            f"averages {top['best_range_avg_dev_pct']:.1f}% deviation from target - a "
            f"{top['spread_pct']:.1f} point gap versus this setting's worst-performing range, across "
            f"{int(top['n'])} runs."
        )

    conclusions = [top_line]
    if ranked is not None and len(ranked):
        conclusions.append(
            f"{len(ranked_with_data)} of {len(ranked)} tracked process settings have enough data to "
            "rank for this property."
        )
    conclusions.append(
        "Review applicability against current raw materials and process conditions before "
        "adjusting any setting."
    )

    return {
        "property_name": property_name,
        "subject_desc": subject_desc,
        "pooling_grades": pooling_grades,
        "ranking_rows": ranking_rows,
        "gap_categories": categories,
        "gap_values": values,
        "conclusions": conclusions,
        "generated_at": dt.datetime.utcnow(),
    }


def render_machine_settings_report_pdf(data):
    def build(story):
        _title_block(
            story, "Process Parameter Optimization Report",
            f"{data['property_name']} · {data['subject_desc']}",
        )
        story.append(Paragraph("Context", STYLES["Heading2"]))
        story.append(_key_value_table([
            ("Property", data["property_name"]), ("Subject", data["subject_desc"]),
            ("Pooled by % of target", "Yes" if data["pooling_grades"] else "No"),
            ("Process settings ranked", len(data["ranking_rows"])),
        ]))

        story.append(Paragraph("Analysis", STYLES["Heading2"]))
        _section(story, "All settings, ranked by how clearly they separate outcomes", data["ranking_rows"])
        _bar_chart(
            story, "Gap vs. worst-performing range, by process setting (points)",
            data["gap_categories"], data["gap_values"],
            note="A bigger gap means a more actionable setting.",
        )

        story.append(Paragraph("Conclusions", STYLES["Heading2"]))
        for line in data["conclusions"]:
            story.append(_p(f"• {line}"))
    return _pdf_bytes(build)


def render_machine_settings_report_docx(data):
    doc = Document()
    _docx_report_header(
        doc, "Process Parameter Optimization Report", f"{data['property_name']} · {data['subject_desc']}",
    )
    _docx_heading(doc, "Context", size=15)
    _docx_kv_table(doc, [
        ("Property", data["property_name"]), ("Subject", data["subject_desc"]),
        ("Pooled by % of target", "Yes" if data["pooling_grades"] else "No"),
        ("Process settings ranked", len(data["ranking_rows"])),
    ])

    _docx_heading(doc, "Analysis", size=15)
    _docx_section(doc, "All settings, ranked by how clearly they separate outcomes", data["ranking_rows"])
    _docx_bar_chart(
        doc, "Gap vs. worst-performing range, by process setting (points)",
        data["gap_categories"], data["gap_values"],
        note="A bigger gap means a more actionable setting.",
    )

    _docx_heading(doc, "Conclusions", size=15)
    for line in data["conclusions"]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(line)
    return _docx_bytes(doc)


# ---------------------------------------------------------------------------
# 15. Expert Notes Report (Context / Analysis / Conclusions)
#
# Added 2026-08-04, closing out the Industrial Intelligence reports batch.
# Expert Notes doesn't fit either bucket the other 5 pages fall into (it
# has no PI3-generated recommendation of its own to report on - the
# existing conditional "Download as Word" button on a PI3-sourced note is
# kept as-is, unchanged, per user direction: "The PI3 button stays on is
# only for the PI3 analysis or answer of a question"). Instead this is a
# standing aggregate over the notes the page is already showing - a
# breakdown by confidence level, source (Manual vs PI3), and linked-entity
# type - always visible, not conditional on selecting one note. Lives on
# pages/20_Expert_Notes.py itself. build_expert_notes_report_data() takes
# the exact `notes` list the page has already scoped (tenant) and is
# already showing in its table - never re-derived.
# ---------------------------------------------------------------------------

def build_expert_notes_report_data(session, notes, scope_label):
    """notes: the exact list of db.ExpertNote objects the page has already
    scoped (tenant) and is showing in its own notes table. scope_label: a
    display string describing the company scope in effect, shown in the
    report header."""
    total = len(notes)

    confidence_counts = {}
    source_counts = {}
    link_type_counts = {}
    in_pi3_count = 0
    link_type_labels = {
        "production_run": "Production Run", "foam_grade": "Product Grade", "product_family": "Product Family",
    }
    for n in notes:
        conf = n.confidence_level or "—"
        confidence_counts[conf] = confidence_counts.get(conf, 0) + 1
        source = n.source or "Manual"
        source_counts[source] = source_counts.get(source, 0) + 1
        link_label = link_type_labels.get(n.linked_entity_type, n.linked_entity_type or "—")
        link_type_counts[link_label] = link_type_counts.get(link_label, 0) + 1
        if n.vector_store_file_id:
            in_pi3_count += 1

    confidence_rows = [
        {"Confidence level": k, "Count": v}
        for k, v in sorted(confidence_counts.items(), key=lambda kv: -kv[1])
    ]
    source_rows = [
        {"Source": k, "Count": v} for k, v in sorted(source_counts.items(), key=lambda kv: -kv[1])
    ]
    link_type_rows = [
        {"Linked to": k, "Count": v} for k, v in sorted(link_type_counts.items(), key=lambda kv: -kv[1])
    ]

    conclusions = [f"{total} expert note(s) recorded in this scope."]
    if total:
        top_confidence = max(confidence_counts.items(), key=lambda kv: kv[1])
        top_source = max(source_counts.items(), key=lambda kv: kv[1])
        top_link = max(link_type_counts.items(), key=lambda kv: kv[1])
        conclusions.append(f"Most common confidence level: {top_confidence[0]} ({top_confidence[1]} note(s)).")
        conclusions.append(f"Most common source: {top_source[0]} ({top_source[1]} note(s)).")
        conclusions.append(f"Most common link type: {top_link[0]} ({top_link[1]} note(s)).")
        conclusions.append(
            f"{in_pi3_count} of {total} note(s) ({round(100 * in_pi3_count / total)}%) have been fed "
            "into PI3's knowledge base."
        )
    else:
        conclusions.append("No notes recorded yet in this scope.")

    return {
        "scope_label": scope_label,
        "total": total,
        "confidence_rows": confidence_rows,
        "source_rows": source_rows,
        "link_type_rows": link_type_rows,
        "in_pi3_count": in_pi3_count,
        "conclusions": conclusions,
        "generated_at": dt.datetime.utcnow(),
    }


def render_expert_notes_report_pdf(data):
    def build(story):
        _title_block(story, "Expert Notes Report", data["scope_label"])
        story.append(Paragraph("Context", STYLES["Heading2"]))
        story.append(_key_value_table([
            ("Scope", data["scope_label"]), ("Total notes", data["total"]),
            ("Fed into PI3", f"{data['in_pi3_count']} of {data['total']}"),
        ]))

        story.append(Paragraph("Analysis", STYLES["Heading2"]))
        _bar_chart(
            story, "Notes by confidence level",
            [r["Confidence level"] for r in data["confidence_rows"]],
            [r["Count"] for r in data["confidence_rows"]],
        )
        _bar_chart(
            story, "Notes by source",
            [r["Source"] for r in data["source_rows"]],
            [r["Count"] for r in data["source_rows"]],
        )
        _section(story, "Notes by linked entity type", data["link_type_rows"])

        story.append(Paragraph("Conclusions", STYLES["Heading2"]))
        for line in data["conclusions"]:
            story.append(_p(f"• {line}"))
    return _pdf_bytes(build)


def render_expert_notes_report_docx(data):
    doc = Document()
    _docx_report_header(doc, "Expert Notes Report", data["scope_label"])

    _docx_heading(doc, "Context", size=15)
    _docx_kv_table(doc, [
        ("Scope", data["scope_label"]), ("Total notes", data["total"]),
        ("Fed into PI3", f"{data['in_pi3_count']} of {data['total']}"),
    ])

    _docx_heading(doc, "Analysis", size=15)
    _docx_bar_chart(
        doc, "Notes by confidence level",
        [r["Confidence level"] for r in data["confidence_rows"]],
        [r["Count"] for r in data["confidence_rows"]],
    )
    _docx_bar_chart(
        doc, "Notes by source",
        [r["Source"] for r in data["source_rows"]],
        [r["Count"] for r in data["source_rows"]],
    )
    _docx_section(doc, "Notes by linked entity type", data["link_type_rows"])

    _docx_heading(doc, "Conclusions", size=15)
    for line in data["conclusions"]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(line)
    return _docx_bytes(doc)


# ---------------------------------------------------------------------------
# 12. WP3 Property Conformance Report (Converged Joint Implementation Plan,
# section 7.4, Gate 2 acceptance items A6/A7 - "report displays the result,
# limits, method, unit, condition, sample context and provenance" / "analytics
# receive only comparable records for the selected property context").
#
# Follows the exact same build_data()/render_docx() pattern as every other
# report in this file (see Batch Release Record / Sample Certificate of
# Analysis above) - picked deliberately, per user direction 2026-08-07, to
# reuse the flexible app's established reporting architecture rather than
# invent a new one, since nothing about presenting a conformance verdict is
# materially different here. The one real difference is the underlying
# conformance model: this report's rows come from wp3_conformance.py's
# richer per-spec matching (method/unit/condition/orientation/location-aware,
# with EXCLUDED_CONTEXT/INVALID/NO_RESULT states - see that module) rather
# than quality_standards.compute_pass_fail's simpler target+tolerance model
# used by every other report on this page. Conformance/analytics data is
# never stored (see wp3_conformance.py's module docstring) - both
# compute_conformance_report() and compute_grade_conformance_summary() run
# fresh every time this report is built, so a corrected specification or
# result is reflected immediately with no separate recompute step.
# ---------------------------------------------------------------------------

def build_wp3_conformance_report_data(session, foam_grade_id, production_run_id):
    grade = session.get(FoamGrade, foam_grade_id)
    run = session.get(ProductionRun, production_run_id)
    if grade is None or run is None:
        return None

    rows = wp3_conformance.compute_conformance_report(session, foam_grade_id, production_run_id=production_run_id)
    summary = wp3_conformance.compute_grade_conformance_summary(session, foam_grade_id)

    def _limit_text(spec):
        if spec is None:
            return "—"
        op = spec.target_operator or "<="
        unit = spec.unit or ""
        if op == "between":
            if spec.lower_limit is None or spec.upper_limit is None:
                return "—"
            return f"{spec.lower_limit} – {spec.upper_limit} {unit}".strip()
        if spec.target_value is None:
            return "—"
        return f"{op} {spec.target_value} {unit}".strip()

    conformance_rows = []
    sample_ids_seen = set()
    for row in rows:
        spec = session.get(GradeSpecification, row["spec_id"]) if row.get("spec_id") else None
        result = session.get(PhysicalPropertyResult, row["result_id"]) if row.get("result_id") else None
        sample = result.sample if (result is not None and result.sample_id) else None
        if sample is not None:
            sample_ids_seen.add(sample.id)

        method_code = "—"
        if spec is not None and spec.property_method_id:
            m = session.get(PhysicalPropertyMethod, spec.property_method_id)
            method_code = m.method_code if m else "—"
        elif result is not None and result.test_method:
            method_code = result.test_method

        condition_name = (
            (spec.condition.name if spec is not None and spec.condition else None)
            or (result.condition.name if result is not None and result.condition else None) or "—"
        )
        orientation_name = (
            (spec.orientation.name if spec is not None and spec.orientation else None)
            or (result.orientation.name if result is not None and result.orientation else None) or "—"
        )
        location_name = (
            (spec.location.name if spec is not None and spec.location else None)
            or (result.location.name if result is not None and result.location else None) or "—"
        )
        unit = (spec.unit if spec is not None else None) or (result.unit if result is not None else None) or "—"

        note_parts = []
        if row.get("excluded_reason"):
            note_parts.append(row["excluded_reason"])
        # CR-09 (2026-08-12): translated via customer_presentation, same as
        # the Recipe Optimization achievement table above - see that
        # function's docstring for what's unchanged vs. what's translated.
        release_note = customer_presentation.customer_facing_release_note(row.get("production_release"))
        if release_note:
            note_parts.append(release_note)
        if row.get("unit_converted"):
            note_parts.append(f"converted from {row.get('as_recorded_value')} {row.get('as_recorded_unit')}")

        conformance_rows.append({
            "Property": row["property_name"],
            "Method": method_code,
            "Unit": unit,
            "Condition": condition_name,
            "Orientation": orientation_name,
            "Location": location_name,
            "Limit / target": _limit_text(spec),
            "Actual": row.get("actual_value"),
            "Status": row["status"] or "—",
            "Note": "; ".join(note_parts) or "—",
            "Sample ID": sample.id if sample is not None else "—",
        })

    sample_rows = []
    for sid in sorted(sample_ids_seen):
        s = session.get(Sample, sid)
        sample_rows.append({
            "Sample ID": s.id,
            "Location": s.location.name if s.location else "—",
            "Orientation": s.orientation.name if s.orientation else "—",
            "Thickness (mm)": s.thickness_mm,
            "Age (hours)": s.age_hours,
            "Sample scope": s.sample_scope or "—",
            "Sampled": s.sample_ts,
        })

    summary_rows = [
        {
            "Property": s["property_name"],
            "Pass": s["pass_count"],
            "Fail": s["fail_count"],
            "Total evaluated": s["total_evaluated"],
            "Pass rate (%)": s["pass_rate_pct"],
        }
        for s in summary
    ]

    statuses = [r["Status"] for r in conformance_rows]
    if not statuses:
        overall_verdict = "No specifications recorded for this grade"
    elif "Fail" in statuses:
        overall_verdict = "Non-conforming"
    elif statuses and all(s == "Pass" for s in statuses):
        overall_verdict = "Conforming"
    else:
        overall_verdict = "Incomplete / excluded results present"

    return {
        "foam_grade_id": grade.id,
        "run_id": run.id,
        "grade_name": grade.grade_name,
        "plant": run.plant.name if run.plant else "—",
        "run_date": run.run_date,
        "batch_reference": run.batch_reference or "—",
        "machine": run.machine.name if run.machine else "—",
        "chemistry": grade.chemistry.name if grade.chemistry else "—",
        # Uses the RUN's own immutable Production Method snapshot, not
        # grade.production_method_id (deprecated 2026-08-10 - see db.py's
        # FoamGrade model and helpers.grade_production_methods). This
        # report is generated for one specific run; the run's own
        # snapshot is the only value that can never disagree with which
        # machine/method actually produced it, whereas the grade-level
        # field could silently diverge once a grade's machines span more
        # than one Production Method.
        "production_method": run.production_method.name if run.production_method else "—",
        "application": grade.application.name if grade.application else "—",
        "construction": grade.construction.name if grade.construction else "—",
        "grade_status": grade.status or "—",
        "overall_verdict": overall_verdict,
        "conformance_rows": conformance_rows,
        "sample_rows": sample_rows,
        "summary_rows": summary_rows,
    }


def render_wp3_conformance_report_docx(data):
    doc = Document()
    # CR-09 (2026-08-12): title translated via customer_presentation -
    # "WP3" is a development work-package identifier, not a customer-facing
    # report name. See pages/21_Report.py's matching tab-label fix.
    _docx_report_header(
        doc, f"{customer_presentation.customer_facing_report_title('WP3 Property Conformance Report')} — Run #{data['run_id']}",
        f"{data['grade_name']} · {data['plant']} · Verdict: {data['overall_verdict']}",
    )
    _docx_heading(doc, "Manufacturing scope", size=12, color=_HTC_GREY, space_before=6)
    _docx_kv_table(doc, [
        ("Chemistry", data["chemistry"]), ("Production method", data["production_method"]),
        ("Application", data["application"]), ("Construction", data["construction"]),
        # CR-09 (2026-08-12): FoamGrade.status can be the internal
        # "UAT_ONLY" lifecycle code (db.py) - translated here so it never
        # reaches a customer report verbatim.
        ("Grade status", customer_presentation.customer_facing_grade_status_label(data["grade_status"]) or "—"),
    ])

    _docx_heading(doc, "Run", size=12, color=_HTC_GREY, space_before=10)
    _docx_kv_table(doc, [
        ("Run date", data["run_date"]), ("Batch reference", data["batch_reference"]),
        ("Production Unit or Cell", data["machine"]), ("Product grade", data["grade_name"]),
    ])

    _docx_section(doc, "Conformance results", data["conformance_rows"])
    _docx_section(doc, "Sample provenance", data["sample_rows"])
    _docx_section(doc, "Analytics — pass rate by property (all runs on this grade)", data["summary_rows"])

    note = doc.add_paragraph(
        "Conformance is computed live from the current grade specification and physical property "
        "results each time this report is generated (see wp3_conformance.py) - it is never stored, "
        "so correcting a specification or a result takes effect immediately on the next report."
    )
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(8.5)
    return _docx_bytes(doc)
